# unified_app.py - Sesli Belge Doldurma Sistemi (Birleşik Versiyon)
# Voice.py ve app2.py projelerinin en iyi özelliklerini birleştiren gelişmiş sistem
# EK-15 geliştirmesi: Öğrencinin ağzından özgün, PIŞMAN VE RESMÎ üslupla 4 yanıt üretimi

import io
import os
import re
import json
import sys
import tempfile
import uuid
import traceback
import random
from typing import Dict, List, Optional, Set, Tuple
from datetime import datetime, date, time
from zoneinfo import ZoneInfo

import streamlit as st
import importlib
from docx import Document
import dateparser

# Local session management import
from local_session_manager import get_local_session_manager, merge_extracted_data, detect_conflicts

# Özel form davranışları (Ek bazlı özel prompt ve alan kısıtlama)
SPECIAL_FORMS: Dict[str, Dict[str, object]] = {
    "Ek 15": {
        "expected_placeholders": [
            "{iddia_nedir}",
            "{iddilar_hakkinda_ne_diyorsunuz}",
            "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}",
            "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}",
        ],
        "placeholder_explanations": {
            "{iddia_nedir}": "Hakkındaki iddianın öğrencinin ağzından özeti.",
            "{iddilar_hakkinda_ne_diyorsunuz}": "Ayrıntılı ifade; olayın nasıl geliştiği, itiraf/inkar, pişmanlık.",
            "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}": "Eklemek istediği hususlar, özür vb.",
            "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}": "Tutanağa dair ekleme/çıkarma isteği; yoksa okudum–onayladım beyanı.",
        },
    }
}

# OpenAI import
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Mikrofon kütüphaneleri - birincil ve alternatifler
MIC_IMPORT_ERROR: Optional[str] = None
mic_recorder = None  # from streamlit_mic_recorder
audio_recorder_fn = None  # from audio_recorder_streamlit
st_audiorec_fn = None  # from st_audiorec

try:
    _mic_module = importlib.import_module("streamlit_mic_recorder")
    mic_recorder = getattr(_mic_module, "mic_recorder", None)
except Exception as e:
    MIC_IMPORT_ERROR = f"streamlit-mic-recorder yüklenemedi: {e}"

try:
    _ar_module = importlib.import_module("audio_recorder_streamlit")
    audio_recorder_fn = getattr(_ar_module, "audio_recorder", None)
except Exception:
    pass

try:
    _sar_module = importlib.import_module("st_audiorec")
    st_audiorec_fn = getattr(_sar_module, "st_audiorec", None)
except Exception:
    pass

def render_audio_recorder_ui() -> Optional[bytes]:
    """Tarayıcıdan ses kaydı al ve bytes döndür (mevcut bileşene göre)."""
    if mic_recorder is not None:
        st.write("**Mikrofon ile Kayıt**")
        rec_val = mic_recorder(
            start_prompt="🎙️ Kaydı Başlat",
            stop_prompt="⏹️ Kaydı Durdur",
            just_once=False,
            use_container_width=True,
            key="unified_mic_recorder",
        )
        if isinstance(rec_val, dict) and rec_val.get("error"):
            st.error(f"Mikrofon hatası: {rec_val['error']}")
            return None
        return bytes_from_mic_return(rec_val)

    if audio_recorder_fn is not None:
        st.write("**Mikrofon ile Kayıt (alternatif)**")
        rec_val = audio_recorder_fn()
        return bytes_from_mic_return(rec_val) if rec_val else None

    if st_audiorec_fn is not None:
        st.write("**Mikrofon ile Kayıt (alternatif)**")
        rec_val = st_audiorec_fn()
        return bytes_from_mic_return(rec_val) if rec_val else None

    st.error("Mikrofon kütüphanesi mevcut değil.")
    if MIC_IMPORT_ERROR:
        st.error(MIC_IMPORT_ERROR)
    st.info("Lütfen 'streamlit-mic-recorder' veya 'audio-recorder-streamlit' paketini kurun.")
    return None

# Zaman dilimi
IST = ZoneInfo("Europe/Istanbul")
TR_DAYS = {0:"Pazartesi", 1:"Salı", 2:"Çarşamba", 3:"Perşembe", 4:"Cuma", 5:"Cumartesi", 6:"Pazar"}

# ================== Yardımcı Fonksiyonlar ==================

def ensure_utf8_encoding():
    import locale
    try:
        current_encoding = locale.getpreferredencoding()
        if 'utf-8' not in current_encoding.lower() and 'cp65001' not in current_encoding.lower():
            os.environ['PYTHONIOENCODING'] = 'utf-8'
    except Exception:
        pass

def safe_str(obj) -> str:
    try:
        return str(obj)
    except UnicodeError:
        try:
            return str(obj).encode('utf-8', errors='replace').decode('utf-8')
        except:
            return repr(obj)

def bytes_from_mic_return(value) -> Optional[bytes]:
    if value is None:
        return None
    if isinstance(value, dict) and "bytes" in value:
        return value["bytes"]
    if isinstance(value, (bytes, bytearray)):
        return bytes(value)
    return None

def transcribe_audio_bytes(audio_bytes: bytes, api_key: str, lang: str = "tr") -> Optional[str]:
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil. 'openai' paketini kurun.")
        return None

    tmp_path = None
    try:
        safe_api_key = api_key.strip() if api_key else ""
        if not safe_api_key:
            st.error("API key boş veya geçersiz")
            return None
        client = OpenAI(api_key=safe_api_key)

        safe_filename = f"audio_{uuid.uuid4().hex}.wav"
        tmp_dir = tempfile.gettempdir()
        tmp_path = os.path.join(tmp_dir, safe_filename)
        try:
            tmp_path.encode('ascii')
        except UnicodeEncodeError:
            tmp_path = os.path.join(tempfile.gettempdir(), f"temp_audio_{uuid.uuid4().hex[:8]}.wav")

        with open(tmp_path, "wb") as f:
            f.write(audio_bytes)

        with open(tmp_path, "rb") as f:
            resp = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                language=lang,
                response_format="text"
            )

        if isinstance(resp, str):
            return resp
        else:
            text_result = getattr(resp, "text", None) or (resp.get("text") if isinstance(resp, dict) else None)
            return str(text_result) if text_result else None

    except Exception as e:
        try:
            error_msg = str(e)
            error_msg.encode('ascii')
        except (UnicodeError, UnicodeEncodeError, UnicodeDecodeError):
            try:
                error_msg = repr(str(e))
            except:
                error_msg = "Ses işleme sırasında Unicode karakter hatası oluştu"
        safe_error_msg = safe_str(error_msg)
        st.error(f"Ses metne çevrilemedi: {safe_error_msg}")
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass

def extract_placeholders_from_docx_bytes(file_bytes: bytes) -> Tuple[Set[str], str]:
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text += p.text + " "
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    text += p.text + " "
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    text += p.text + " "
    placeholders = set(re.findall(r"\{[^}]+\}", text))
    return placeholders, text

def replace_placeholders_in_document(doc: Document, placeholder_values: Dict[str, str]) -> int:
    def replace_in_paragraph(paragraph):
        if not paragraph.runs:
            return 0
        original_text = "".join(run.text for run in paragraph.runs)
        replaced_text = original_text
        total_replacements = 0
        for placeholder, value in placeholder_values.items():
            if value is None:
                continue
            count = replaced_text.count(placeholder)
            if count:
                replaced_text = replaced_text.replace(placeholder, str(value))
                total_replacements += count
        if replaced_text != original_text:
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            paragraph.add_run(replaced_text)
        return total_replacements

    replacements_made = 0
    for p in doc.paragraphs:
        replacements_made += replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replacements_made += replace_in_paragraph(p)
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replacements_made += replace_in_paragraph(p)
        if section.footer:
            for p in section.footer.paragraphs:
                replacements_made += replace_in_paragraph(p)
    return replacements_made

def parse_tr_date(text: str) -> Optional[datetime]:
    if not text:
        return None
    return dateparser.parse(text, languages=["tr"])

def split_date(dt: datetime) -> Dict[str, str]:
    return {"gun": f"{dt.day:02d}", "ay": f"{dt.month:02d}", "yil": f"{dt.year}"}

def today_isbu(dt: datetime = None) -> Dict[str, str]:
    now = dt or datetime.now(IST)
    return {
        "isbu_gun": f"{now.day:02d}",
        "isbu_ay": f"{now.month:02d}",
        "isbu_yil": f"{now.year}",
        "isbu_saat": now.strftime("%H:%M")
    }

# ================== AI Analiz Fonksiyonları ==================

def extract_placeholder_contexts_from_docx_bytes(file_bytes: bytes, placeholders: Set[str], window: int = 70) -> Dict[str, List[str]]:
    doc = Document(io.BytesIO(file_bytes))
    blocks = []
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        blocks.append(p.text)
    contexts = {ph: [] for ph in placeholders}
    for block in blocks:
        for ph in placeholders:
            pattern = re.escape(ph)
            for m in re.finditer(pattern, block):
                start, end = m.start(), m.end()
                before = block[max(0, start - window): start]
                after = block[end: end + window]
                snippet = f"{before}{ph}{after}"
                if len(contexts[ph]) < 3:
                    contexts[ph].append(snippet)
    return contexts

def aggregate_contexts_across_templates(templates: List[Tuple[str, bytes]], placeholders: Set[str]) -> Dict[str, List[str]]:
    combined = {ph: [] for ph in placeholders}
    for _, data in templates:
        try:
            local_ctx = extract_placeholder_contexts_from_docx_bytes(data, placeholders)
            for ph, lst in local_ctx.items():
                for s in lst:
                    if len(combined[ph]) < 5:
                        combined[ph].append(s)
        except Exception:
            continue
    return combined

def parse_json_loose(s: str) -> Dict[str, str]:
    try:
        return json.loads(s)
    except Exception:
        pass
    try:
        m = re.search(r"\{[\s\S]*\}", s)
        if m:
            return json.loads(m.group(0))
    except Exception:
        pass
    return {}

def infer_placeholder_values(
    transcript: str,
    placeholders: Set[str],
    contexts: Dict[str, List[str]],
    api_key: str,
    model: str = "gpt-4o-mini",
    only_placeholders: Optional[Set[str]] = None,
    extra_instructions: Optional[str] = None,
    placeholder_explanations: Optional[Dict[str, str]] = None,
) -> Dict[str, str]:
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil.")
        return {}

    client = OpenAI(api_key=api_key)
    ph_list = sorted(list(placeholders))
    if only_placeholders:
        wanted = {p.strip("{}").lower() for p in only_placeholders}
        filtered: List[str] = []
        for ph in ph_list:
            key_nb = ph.strip("{}").lower()
            if key_nb in wanted:
                filtered.append(ph)
        ph_list = filtered or ph_list

    existing_values = {}
    try:
        existing_values = {
            k: v for k, v in (st.session_state.get("current_mapping", {}) or {}).items()
            if k in placeholders and str(v).strip()
        }
    except Exception:
        existing_values = {}

    ph_lower_list = [ph.lower() for ph in placeholders]
    has_fullname_key = any(("ogrenci" in p and ("adi_soyadi" in p or "ad_soyad" in p)) for p in ph_lower_list)
    has_name_key = any(("ogrenci" in p and ("ad" in p or "adi" in p or "isim" in p) and "soyad" not in p) for p in ph_lower_list)
    has_surname_key = any(("ogrenci" in p and ("soyad" in p or "soyadi" in p)) for p in ph_lower_list)
    has_department_key = any(("bolum" in p) or ("bölüm" in p) for p in ph_lower_list)

    name_rules_lines: List[str] = []
    if has_fullname_key:
        name_rules_lines.append("- {ogrenci_adi_soyadi} alanı için öğrencinin tam adını 'Ad Soyad' formatında ver.")
    if has_name_key and has_surname_key:
        name_rules_lines.append("- {ogrenci_ad}/{ogrenci_adi} ve {ogrenci_soyad}/{ogrenci_soyadi} alanlarını ayrı doldur.")
    elif has_name_key or has_surname_key:
        name_rules_lines.append("- Ad veya soyad alanı tekse, transkriptte geçtiği gibi yaz.")

    department_rules_lines: List[str] = []
    if has_department_key:
        department_rules_lines.append("- Bölüm alanlarında tek bir bölüm adı ver (örn: 'İşletme').")

    prompt_text = f"""
SES TRANSKRİPTİ:
"{transcript}"

TEMPLATE PLACEHOLDER'LARI VE BAĞLAMLARI:
"""
    for ph in ph_list:
        if ph in contexts and contexts[ph]:
            context_examples = "\n".join([f"  • {ctx[:200]}" for ctx in contexts[ph][:3]])
            prompt_text += f"\n{ph}:\n{context_examples}\n"
        else:
            prompt_text += f"\n{ph}: (Bağlam bulunamadı)\n"

    prompt_text += """

GÖREV:
1. Ses transkriptini analiz et
2. Her placeholder için template bağlamını incele
3. Bağlama uygun değerleri ses transkriptinden çıkar
4. Çıkaramadığın bilgiler için boş string ("") bırak
5. SADECE JSON formatında cevap ver

MEVCUT DEĞERLER (DEĞİŞTİRME):
""" + json.dumps(existing_values, ensure_ascii=False, indent=2) + """

ÇIKTI KURALLARI:
- JSON anahtarları, placeholder stringleriyle birebir aynı olmalı (örnek: {ogrenci_no})
- Mevcut dolu alanları DEĞİŞTİRME; sadece boş olanları doldur
- Tarih/saat alanları bağlama uygun normalize edilmeli (tarih: YYYY-MM-DD veya {gun,ay,yil}; saat: HH:MM)
- Sayısal alanlar (no, tc vb.) sadece rakam içersin
- Açıklama alanlarında öğrenci ismi geçmesin, olay nesnel anlatılsın

"""
    if name_rules_lines or department_rules_lines:
        prompt_text += "\n".join(name_rules_lines + department_rules_lines) + "\n"

    messages = [
        {"role": "system", "content": "Uzman bir bilgi çıkarım asistanısın. Sadece JSON döndür. Mevcut dolu alanları asla değiştirme."},
        {"role": "user", "content": prompt_text},
    ]

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.0,
        )
        content = resp.choices[0].message.content if resp and resp.choices else "{}"
        data = parse_json_loose(content or "{}")

        result = {}
        for ph in placeholders:
            key_lower = ph.lower()
            raw_val = str(data.get(ph, "")).strip() if isinstance(data, dict) else ""
            val = raw_val
            if ("bolum" in key_lower) or ("bölüm" in key_lower):
                lowered = val.lower()
                for sep in [",", "/", "&", "|", ";"]:
                    if sep in val:
                        val = val.split(sep)[0]
                for conj in [" ve ", " veya "]:
                    if conj in lowered:
                        idx = lowered.index(conj)
                        val = val[:idx]
                        break
                val = str(val).strip()
            result[ph] = val if val else ""

        # İsim alanları için basit senkronizasyon
        try:
            keys_by_lower = {k.lower(): k for k in result.keys()}
            fullname_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("adi_soyadi" in lk or "ad_soyad" in lk))), None)
            name_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("ad" in lk or "adi" in lk or "isim" in lk) and "soyad" not in lk)), None)
            surname_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("soyad" in lk or "soyadi" in lk))), None)
            if fullname_key and (not result.get(fullname_key)) and name_key and surname_key and result.get(name_key) and result.get(surname_key):
                combined = f"{str(result.get(name_key)).strip()} {str(result.get(surname_key)).strip()}".strip()
                result[fullname_key] = combined
            if fullname_key and result.get(fullname_key) and ((name_key and not result.get(name_key)) or (surname_key and not result.get(surname_key))):
                fullname_val = str(result.get(fullname_key)).strip()
                parts = [p for p in fullname_val.split() if p]
                if len(parts) >= 2:
                    first = " ".join(parts[:-1])
                    last = parts[-1]
                    if name_key and not result.get(name_key):
                        result[name_key] = first
                    if surname_key and not result.get(surname_key):
                        result[surname_key] = last
        except Exception:
            pass

        return result
    except Exception as e:
        try:
            error_msg = str(e)
        except UnicodeEncodeError:
            error_msg = "AI analizi sırasında karakter kodlama hatası"
        st.error(f"Ses analizi başarısız: {error_msg}")
        return {ph: "" for ph in placeholders}

# -------- EK-15: Öğrenci ağzından, RESMÎ ve pişmanlık içeren özgün cevap üretici --------

def _build_ek15_base_facts_text(mapping: Dict[str, str], transcript: str) -> str:
    """
    Transkript + mevcut çıkarımlardan olayı toparlayan kısa temel bilgi metni üretir (LLM'e bağlam olarak verilir).
    Uydurma yok; yalnızca bilinenler derlenir.
    """
    parts = []
    def g(keys: List[str]) -> Optional[str]:
        for k in keys:
            v = mapping.get(k) or mapping.get("{"+k.strip("{}")+"}")
            if v and str(v).strip():
                return str(v).strip()
        return None

    ders = g(["{ders_adi}", "{ders}", "{ders_kodu}", "{dersin_adi}"])
    sinav_tarih = g(["{sinav_tarihi}", "{sınav_tarihi}", "{sınav_tarih}", "{sinav_tarih}", "{tarih}"])
    sinav_saat = g(["{sinav_saati}", "{sınav_saati}", "{saat}"])
    yer = g(["{sinav_yeri}", "{salon}", "{sinif}", "{sınıf}"])
    gozetmen = g(["{gozetmen_adi_soyadi}", "{gözetmen_adi_soyadi}", "{gozetmen}", "{gözetmen}"])
    yaklasim = g(["{kopya_yontemi}", "{olay_aciklama}", "{olay_aciklamasi}", "{olay}"])
    cihaz = g(["{cihaz}", "{telefon}", "{elektronik_esya}"])

    if ders: parts.append(f"Ders: {ders}.")
    if sinav_tarih or sinav_saat:
        if sinav_tarih and sinav_saat:
            parts.append(f"Sınav tarihi-saat: {sinav_tarih} {sinav_saat}.")
        elif sinav_tarih:
            parts.append(f"Sınav tarihi: {sinav_tarih}.")
        else:
            parts.append(f"Sınav saati: {sinav_saat}.")
    if yer: parts.append(f"Yer: {yer}.")
    if gozetmen: parts.append(f"Gözetmen: {gozetmen}.")
    if cihaz: parts.append(f"Kullanılan cihaz/araç: {cihaz}.")
    if yaklasim: parts.append(f"Olay özeti (çıkarım): {yaklasim}.")

    if transcript and transcript.strip():
        parts.append(f"Transkript özeti: {transcript.strip()[:900]}")

    return " ".join(parts)

def generate_student_style_response(
    api_key: str,
    question_key: str,
    transcript: str,
    mapping: Dict[str, str],
    user_hint: str = "",
    model: str = "gpt-4o-mini",
) -> str:
    """
    Ek-15 soruları için öğrencinin ağzından ÖZGÜN, PIŞMAN VE **RESMÎ** cevap üretir.
    question_key: "{iddia_nedir}", "{iddilar_hakkinda_ne_diyorsunuz}", "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}", "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}"
    """
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil.")
        return ""

    client = OpenAI(api_key=api_key.strip())
    base_facts = _build_ek15_base_facts_text(mapping, transcript)

    # Soruya göre hedef uzunluk
    if question_key == "{iddilar_hakkinda_ne_diyorsunuz}":
        target = "120-220 kelime, olayın gelişimi, itiraf ve pişmanlık açık; resmî, ölçülü bir üslup."
    elif question_key == "{iddia_nedir}":
        target = "60-120 kelime, iddianın öz ve dürüst özeti; resmî ifade."
    elif question_key == "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}":
        target = "70-130 kelime, pişmanlık ve özür; resmî ve saygılı bir dil, klişe ifadelerden kaçın."
    else:  # {tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}
        target = "50-110 kelime, tutanağın okunduğu ve gerekli görülüyorsa tarih/saatin belirtildiği resmî beyan."

    # Üslup çeşitliliği
    temp = round(random.uniform(0.8, 1.0), 2)
    top_p = round(random.uniform(0.85, 1.0), 2)

    # >>> RESMÎ ÜSLUP TALİMATI (güncellendi)
    system = (
        "Bir üniversite disiplin sürecinde kullanılacak **resmî** bir ifade metni yazarsın. "
        "Metin, öğrencinin ağzından; ciddi, ölçülü, saygılı ve kurum diline uygundur. "
        "Argo, duygusal abartı ve günlük konuşma kalıplarından kaçınırsın. "
        "Gerçeğe dayalı yaz, uydurma detay ekleme. Türkçe yaz. "
        "Çıktı yalnızca tek paragraf doğal metin olmalı; madde işaretleri ve başlık kullanma."
    )

    hint_line = f"İpucu/Ton isteği: {user_hint}" if user_hint else "İpucu/Ton isteği: (yok)"

    question_prompt_map = {
        "{iddia_nedir}": "Soru: Hakkınızda yöneltilen iddia nedir? Açıklar mısınız?",
        "{iddilar_hakkinda_ne_diyorsunuz}": "Soru: İddia/iddialar hakkında ne diyorsunuz? Ayrıntılı açıklayarak anlatınız.",
        "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}": "Soru: Konu hakkında eklemek istediğiniz başka bir şey var mı?",
        "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}": "Soru: Ekleyeceğiniz başka bir şey yoksa tutanağı okuyunuz, eklenmesini/düzeltilmesini istediğiniz yer var mı?",
    }
    qtext = question_prompt_map.get(question_key, "Soru: (Ek-15)")

    extra_rules = []
    if question_key == "{iddia_nedir}":
        extra_rules.append("İddianın kapsamını tarih/ders/ihlâl yönüyle özlü biçimde belirt; 'kopya çekmeye teşebbüs' gibi net bir ifade kullanılabilir.")
    if question_key == "{iddilar_hakkinda_ne_diyorsunuz}":
        extra_rules.append("Olay akışını mantıklı sırayla aktar: hazırlık durumu → sınav anı → hatalı eylem → tespit ve müdahale.")
        extra_rules.append("İsim/yer/üçüncü kişi bilgisi uydurma; yalnızca bağlamda verilenleri kullan.")
    if question_key == "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}":
        extra_rules.append("Pişmanlık ve özür ifadesi içersin; aynı kalıp cümleleri tekrar etme.")
    if question_key == "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}":
        extra_rules.append("Uygunsa ifade alma tarih-saatini belirt; değilse 'okudum, ek talebim yoktur' beyanı yeterlidir.")

    user = f"""
EK-15 CEVAP ÜRETİMİ (RESMÎ DİL)
Bağlam (bilinen gerçekler): {base_facts}

{qtext}
{hint_line}

Hedef uzunluk: {target}
Biçim: tek paragraf, resmî ve ölçülü doğal metin.
TON KURALI:
- Öğrencinin ağzından, resmî, saygılı, ölçülü; ceza hafifletici bir çerçeve
- Klişe ve tekrarlayan kalıplardan kaçın; özgün cümle yapıları
- Gerçekçi ve tutarlı; uydurma yok

Ek kurallar:
- Üçüncü kişileri/özel adları yalnızca bağlamda açıkça varsa kullan
- Aşırı duygusal söylem veya argoya yer verme

Ek özel talimatlar:
- {" ".join(extra_rules) if extra_rules else "(yok)"}
"""

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=temp,
            top_p=top_p,
            presence_penalty=0.2,
            frequency_penalty=0.3,
        )
        text = (resp.choices[0].message.content or "").strip()
        if not text or len(text.split()) < 25:
            # Resmî kısa fallback
            text = (
                "Hakkımdaki sürece ilişkin beyanımı saygıyla sunarım. Sınav anında tarafımdan gerçekleştirilen hatalı davranışın "
                "ciddiyetini idrak etmiş bulunmaktayım. Meydana gelen durum nedeniyle üzüntü duyuyor, benzer bir hatayı "
                "tekrarlamamak için gerekli özeni göstereceğimi bildiriyorum."
            )
        return text
    except Exception as e:
        st.error(f"Ek-15 cevabı üretilemedi: {e}")
        return ""

# ================== Öğrenci Yönetimi ==================

def extract_student_info(session_data):
    if not session_data or not session_data.get('extracted_data'):
        return None, None
    extracted = session_data['extracted_data']
    student_no = None
    student_name = None
    for key, value in extracted.items():
        if value and str(value).strip():
            key_lower = key.lower().replace('{', '').replace('}', '')
            value_str = str(value).strip()
            skip_person_keywords = [
                'gozetmen', 'gözetmen', 'ogretim', 'öğretim', 'elemani', 'elemanı',
                'gorevli', 'görevli', 'bolum_baskanligi', 'bölüm başkanlığı', 'baskan', 'başkan',
                'danisman', 'danışman', 'sifre', 'şifre', 'yetkili', 'imza'
            ]
            if any(k in key_lower for k in skip_person_keywords):
                continue
            if 'ogrencino' in key_lower or 'ogrenci_no' in key_lower:
                student_no = value_str
            elif 'no' in key_lower and not student_no:
                student_no = value_str
            elif ('ogrenci' in key_lower) and ('ad' in key_lower or 'adi' in key_lower or 'isim' in key_lower) and 'soyad' not in key_lower:
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{value_str} {student_name}"
                    else:
                        student_name = value_str
            elif ('ogrenci' in key_lower) and ('soyad' in key_lower or 'soyadi' in key_lower):
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{student_name} {value_str}"
                    else:
                        student_name = value_str
            elif 'ogrenci' in key_lower and any(keyword in key_lower for keyword in ['adi_soyadi', 'ad_soyad']):
                if not value_str.isdigit():
                    student_name = value_str
    return student_no, student_name

def update_session_name_if_needed(session_id, session_data):
    try:
        def _format_standard(no: str, name: str) -> str:
            safe_no = (no or "").strip()
            safe_name = " ".join((name or "").split())
            return f"{safe_no} - {safe_name}" if safe_no and safe_name else ""
        student_no, student_name = extract_student_info(session_data)
        current_name = session_data.get('session_name', '')
        if student_no and student_name:
            new_name = _format_standard(student_no, student_name)
            if new_name and new_name != current_name:
                sm = get_local_session_manager()
                session_data['session_name'] = new_name
                return sm.save_session(session_id, session_data)
        return False
    except Exception as e:
        st.error(f"Session ismi güncellenirken hata: {e}")
        return False

# ================== Ana Uygulama ==================

def main():
    ensure_utf8_encoding()
    st.set_page_config(page_title="🎯 Sesli Belge Doldurma Sistemi", page_icon="🎯", layout="wide")

    if "page" not in st.session_state:
        st.session_state["page"] = "session_manager"
    if "current_session_id" not in st.session_state:
        st.session_state["current_session_id"] = None
    if "current_session_name" not in st.session_state:
        st.session_state["current_session_name"] = ""
    if "api_key" not in st.session_state:
        st.session_state["api_key"] = ""
    if "selected_form_group" not in st.session_state:
        st.session_state["selected_form_group"] = None
    if "form_group_applied" not in st.session_state:
        st.session_state["form_group_applied"] = None
    if "templates_initialized_for" not in st.session_state:
        st.session_state["templates_initialized_for"] = None
    if "selected_templates" not in st.session_state:
        st.session_state["selected_templates"] = []
    if "ek15_hints" not in st.session_state:
        st.session_state["ek15_hints"] = {}
    if "ek15_inputs" not in st.session_state:
        st.session_state["ek15_inputs"] = {}

    if st.session_state["page"] == "session_manager":
        show_session_manager()
    elif st.session_state["page"] == "form_selector":
        show_form_selector()
    elif st.session_state["page"] == "voice_app":
        show_voice_app()
    else:
        st.session_state["page"] = "session_manager"
        st.rerun()

def show_session_manager():
    st.title("🎯 Sesli Belge Doldurma Sistemi")
    st.caption("Ses girdi ile Word şablonlarını otomatik dolduran akıllı sistem")

    sm = get_local_session_manager()
    search_term = st.text_input("🔍 Öğrenci Ara", placeholder="Öğrenci adı veya öğrenci numarası...")
    sessions = sm.get_all_sessions()

    if search_term:
        q = search_term.lower().strip()
        filtered_sessions = []
        for session in sessions:
            student_no, student_name = extract_student_info(session)
            match_no = bool(student_no and q in student_no.lower())
            match_name = bool(student_name and q in student_name.lower())
            if match_no or match_name:
                filtered_sessions.append(session)
        sessions = filtered_sessions

    col1, col2 = st.columns([3, 1])
    with col1:
        st.subheader("📁 Öğrenci Session'ları")
        if not sessions:
            if search_term:
                st.info("Arama kriterinize uygun session bulunamadı.")
            else:
                st.info("Henüz session oluşturulmamış. 'Yeni Session' butonuna tıklayın.")
        else:
            for session in sessions:
                student_no, student_name = extract_student_info(session)
                if student_no and student_name:
                    display_title = f"👤 {student_no} - {student_name}"
                else:
                    fallback_name = session.get('session_name', '')
                    display_title = f"👤 {fallback_name}"
                with st.expander(display_title, expanded=False):
                    col_info, col_actions = st.columns([2, 1])
                    with col_info:
                        st.write(f"**Oluşturma:** {session['created_date'][:10]}")
                        if student_no:
                            st.write(f"**Öğrenci No:** {student_no}")
                        if student_name:
                            st.write(f"**Öğrenci Adı:** {student_name}")
                        data_count = len([v for v in session.get('extracted_data', {}).values() if v])
                        st.write(f"**Dolu Alanlar:** {data_count}")
                    with col_actions:
                        if st.button(f"🚀 Aç", key=f"open_{session['session_id']}"):
                            st.session_state["current_session_id"] = session['session_id']
                            st.session_state["current_session_name"] = session['session_name']
                            st.session_state["selected_form_group"] = None
                            st.session_state["form_group_applied"] = None
                            st.session_state["templates_initialized_for"] = None
                            st.session_state["selected_templates"] = []
                            st.session_state["page"] = "form_selector"
                            st.rerun()
                        if st.button(f"🗑️ Sil", key=f"delete_{session['session_id']}"):
                            st.session_state[f"confirm_delete_{session['session_id']}"] = True
                            st.rerun()
                        if st.session_state.get(f"confirm_delete_{session['session_id']}", False):
                            st.warning("⚠️ Silmek istediğinizden emin misiniz?")
                            col_yes, col_no = st.columns(2)
                            with col_yes:
                                if st.button("✅ Evet", key=f"confirm_yes_{session['session_id']}"):
                                    if sm.delete_session(session['session_id']):
                                        st.success("Session silindi!")
                                        del st.session_state[f"confirm_delete_{session['session_id']}"]
                                        st.rerun()
                            with col_no:
                                if st.button("❌ İptal", key=f"confirm_no_{session['session_id']}"):
                                    del st.session_state[f"confirm_delete_{session['session_id']}"]
                                    st.rerun()
    with col2:
        st.subheader("🚀 Yeni Session")
        st.write("Yeni bir öğrenci için session başlatın.")
        if st.button("📝 Yeni Session Başlat", type="primary", use_container_width=True):
            from datetime import datetime as _dt
            timestamp = _dt.now().strftime("%Y%m%d_%H%M%S")
            session_name = f"Yeni Session - {timestamp}"
            sm = get_local_session_manager()
            session_id = sm.create_session(session_name)
            if session_id:
                st.success("Yeni session başlatıldı!")
                st.session_state["current_session_id"] = session_id
                st.session_state["current_session_name"] = session_name
                st.session_state["page"] = "form_selector"
                st.session_state["current_transcript"] = ""
                st.session_state["transcript_loaded_for"] = session_id
                st.session_state["current_mapping"] = {}
                st.session_state["mapping_loaded_for"] = session_id
                st.session_state["results"] = None
                st.session_state["selected_form_group"] = None
                st.session_state["form_group_applied"] = None
                st.session_state["templates_initialized_for"] = None
                st.session_state["selected_templates"] = []
                st.rerun()
            else:
                st.error("Session oluşturulamadı!")
        st.info("💡 **İpucu:** Session başlattıktan sonra öğrenci bilgilerini sesli girdi ile kaydedin.")

def show_form_selector():
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return

    st.title("🧩 Hangi Ek doldurulacak?")
    st.caption(f"{current_session_name}")
    st.markdown("Seçiminiz bu session için şablonları otomatik işaretler. İstediğiniz zaman değiştirebilirsiniz.")

    options = ["Ek 1-2-3", "Ek 4", "Ek 6", "Ek 8", "Ek 9", "Ek 11", "Ek 15"]
    default_idx = options.index(st.session_state.get("selected_form_group")) if st.session_state.get("selected_form_group") in options else 0
    selected = st.radio("Form seti", options=options, index=default_idx, horizontal=True)

    col_go, col_back = st.columns([1, 1])
    with col_go:
        if st.button("Devam et ➜", type="primary", use_container_width=True):
            st.session_state["selected_form_group"] = selected
            st.session_state["form_group_applied"] = None
            st.session_state["templates_initialized_for"] = None
            st.session_state["selected_templates"] = []
            st.session_state["page"] = "voice_app"
            st.rerun()
    with col_back:
        if st.button("↩️ Session listesine dön", use_container_width=True):
            st.session_state["page"] = "session_manager"
            st.rerun()

def show_voice_app():
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return

    sm = get_local_session_manager()
    session_data = sm.get_session(current_session_id)
    if not session_data:
        st.error("Session verisi yüklenemedi!")
        return

    if st.session_state.get("mapping_loaded_for") != current_session_id:
        st.session_state["current_mapping"] = {}
        st.session_state["mapping_loaded_for"] = current_session_id
        st.session_state["results"] = None
    if st.session_state.get("transcript_loaded_for") != current_session_id:
        st.session_state["current_transcript"] = session_data.get('transcript', "")
        st.session_state["transcript_loaded_for"] = current_session_id

    col_title, col_actions = st.columns([4, 2])
    with col_title:
        st.title(f"🎯 {current_session_name}")
        st.caption(f"Session ID: {current_session_id[:12]}...")
        active_group = st.session_state.get("selected_form_group") or "Seçilmedi"
        st.info(f"Aktif Form Seti: {active_group}")
    with col_actions:
        if st.button("🧩 Form setini değiştir"):
            st.session_state["page"] = "form_selector"
            st.rerun()
        if st.button("🏠 Session listesi"):
            st.session_state["page"] = "session_manager"
            st.rerun()

    st.markdown("---")

    col_api1, col_api2 = st.columns([3, 1])
    with col_api1:
        api_key_input = st.text_input(
            "🔑 OpenAI API Key",
            value=st.session_state.get("api_key", ""),
            type="password",
            help="Whisper ve AI analizi için gerekli"
        )
    with col_api2:
        if st.checkbox("Hatırla", value=bool(st.session_state.get("api_key"))):
            st.session_state["api_key"] = api_key_input
        else:
            st.session_state["api_key"] = ""

    st.markdown("---")

    # Şablon seçimi
    st.subheader("📝 Şablon Belgeleri")
    default_dir = os.path.join(os.getcwd(), "templates")
    selected_names = []
    available = []
    try:
        if os.path.isdir(default_dir):
            available = sorted([f for f in os.listdir(default_dir) if f.lower().endswith(".docx")])
            if available:
                group = st.session_state.get("selected_form_group")
                def _match_group_files(group_label: Optional[str], files: List[str]) -> List[str]:
                    if not group_label:
                        return []
                    prefixes_map = {
                        "Ek 1-2-3": ["Ek-1", "Ek-2", "Ek-3"],
                        "Ek 4": ["Ek-4"],
                        "Ek 6": ["Ek-6"],
                        "Ek 8": ["Ek-8"],
                        "Ek 9": ["Ek-9"],
                        "Ek 11": ["Ek-11"],
                        "Ek 15": ["Ek-15"],
                    }
                    prefixes = prefixes_map.get(group_label, [])
                    return [f for f in files if any(f.startswith(pfx) for pfx in prefixes)]
                should_apply_preselection = (
                    st.session_state.get("templates_initialized_for") != current_session_id or
                    st.session_state.get("form_group_applied") != group
                )
                if should_apply_preselection:
                    preselected = _match_group_files(group, available)
                    st.session_state["selected_templates"] = preselected
                    st.session_state["templates_initialized_for"] = current_session_id
                    st.session_state["form_group_applied"] = group
                selected_names = st.multiselect(
                    "Kullanılacak şablonları seçin",
                    options=available,
                    default=st.session_state.get("selected_templates", []),
                    help="Seçtiğiniz şablonların tam önizlemesi aşağıda görüntülenecek",
                    key="selected_templates"
                )
            else:
                st.info("Templates klasöründe .docx şablon bulunamadı.")
        else:
            st.info("Templates klasörü bulunamadı.")
    except Exception as e:
        st.error(f"Templates klasörü okunamadı: {e}")

    template_items = []
    for name in selected_names:
        try:
            full = os.path.join(default_dir, name)
            with open(full, "rb") as fh:
                template_items.append((name, fh.read()))
        except Exception as e:
            st.error(f"{name} okunamadı: {e}")

    union_placeholders = set()
    if template_items:
        for name, data in template_items:
            try:
                placeholders, _ = extract_placeholders_from_docx_bytes(data)
                union_placeholders |= placeholders
            except Exception as e:
                st.error(f"{name} analiz edilemedi: {e}")

    st.markdown("---")

    # Ses kaydı bölümü
    st.subheader("🎤 Ses Kaydı ve Analiz")
    col_mic, col_btn = st.columns([3, 1])
    with col_mic:
        audio_bytes = render_audio_recorder_ui()
        special_text_input = None
        if st.session_state.get("selected_form_group") == "Ek 15":
            special_text_input = st.text_area(
                "📝 (Opsiyonel) Ek 15 İçerik/Notlar",
                value="",
                height=120,
                help="Bu alan opsiyoneldir; per-soru üretim zaten var. İsterseniz bağlam olarak ek metin yazabilirsiniz."
            )
    with col_btn:
        if st.button("🧠 Analiz Et", use_container_width=True, type="primary"):
            effective_key = (api_key_input or st.session_state.get("api_key", "")).strip()
            if not template_items:
                st.warning("Önce şablon seçin.")
                return
            if not union_placeholders:
                st.warning("Şablonlarda placeholder bulunamadı.")
                return
            if not effective_key:
                st.warning("OpenAI API anahtarı girin.")
                return

            existing_transcript = (st.session_state.get("current_transcript", "")).strip()
            merged_transcript = existing_transcript
            if audio_bytes:
                with st.spinner("Ses metne çevriliyor..."):
                    text = transcribe_audio_bytes(audio_bytes, effective_key)
                if not text:
                    st.error("Ses metne çevrilemedi.")
                    return
                merged_transcript = (existing_transcript + " " + text.strip()).strip() if existing_transcript else text.strip()
                st.session_state["current_transcript"] = merged_transcript
                sm.update_session_transcript(current_session_id, merged_transcript)

            with st.spinner("Bilgiler çıkarılıyor..."):
                ctx = aggregate_contexts_across_templates(template_items, union_placeholders)
                suggested: Dict[str, str] = {}
                selected_group = st.session_state.get("selected_form_group")
                ek15_conf = SPECIAL_FORMS.get("Ek 15", {}) if selected_group == "Ek 15" else {}
                ek15_set: Set[str] = set(ek15_conf.get("expected_placeholders", []) or [])

                # Genel çıkarım (Ek 15 dışındaki alanlar)
                general_placeholders: Set[str] = set(union_placeholders)
                if ek15_set:
                    general_placeholders = set(ph for ph in union_placeholders if ph not in ek15_set)
                if general_placeholders:
                    general_suggested = infer_placeholder_values(
                        merged_transcript,
                        general_placeholders,
                        ctx,
                        effective_key,
                    )
                    suggested.update(general_suggested or {})

                # Mevcut verilerle birleştir
                existing_data = st.session_state.get("current_mapping", {})
                conflicts = detect_conflicts(existing_data, suggested)
                if conflicts:
                    st.warning(f"⚠️ {len(conflicts)} çakışma tespit edildi: {', '.join(conflicts)}")

                merged_data = merge_extracted_data(existing_data, suggested)
                st.session_state["current_mapping"] = merged_data
                try:
                    if sm.update_session_data(current_session_id, suggested, merge=True):
                        filled_count = len([v for v in suggested.values() if v.strip()])
                        st.success(f"✅ {filled_count} yeni bilgi eklendi ve kaydedildi!")
                        updated_session = sm.get_session(current_session_id)
                        if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                            st.session_state["current_session_name"] = updated_session['session_name']
                            st.info("📝 Session ismi güncellendi!")
                except Exception as e:
                    st.warning(f"Veriler çıkarıldı ama kaydetme sırasında hata: {e}")
                st.rerun()

    # Transkript gösterimi
    if st.session_state.get("current_transcript"):
        col_transcript, col_clear = st.columns([4, 1])
        with col_transcript:
            st.text_area(
                "📜 Birleşik Transkript",
                value=st.session_state.get("current_transcript", ""),
                height=120,
                disabled=True,
                help="Bu transkript session bazında saklanır ve tüm Ek formlarında kullanılabilir"
            )
        with col_clear:
            st.write("")
            if st.button("🗑️ Temizle"):
                st.session_state["current_transcript"] = ""
                sm.update_session_transcript(current_session_id, "")
                st.rerun()

    # -------------------- EK-15 ÖZEL ÜRETİM ARAYÜZÜ --------------------
    if st.session_state.get("selected_form_group") == "Ek 15":
        st.markdown("---")
        st.subheader("🧠 Ek-15 — Öğrenci Ağzından (RESMÎ DİL) Cevap Üretimi")

        qmap = [
            ("{iddia_nedir}", "1) Hakkınızda yöneltilen iddia nedir? Açıklar mısınız?"),
            ("{iddilar_hakkinda_ne_diyorsunuz}", "2) İddia/iddialar hakkında ne diyorsunuz? Ayrıntılı açıklayarak anlatınız."),
            ("{konu_hk_eklemek_istediginiz_bir_sey_var_mi}", "3) Konu hakkında eklemek istediğiniz başka bir şey var mı?"),
            ("{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}", "4) Tutanak – eklemek/düzeltmek istediğiniz yer var mı?"),
        ]

        if "current_mapping" not in st.session_state:
            st.session_state["current_mapping"] = {}

        cols = st.columns(2)
        for idx, (ph, title) in enumerate(qmap):
            col = cols[idx % 2]
            with col:
                st.markdown(f"**{title}**")
                current_val = (st.session_state.get("current_mapping", {}) or {}).get(ph, "")
                new_val = st.text_area(
                    f"{ph}__textarea",
                    value=current_val,
                    height=180,
                    placeholder="Bu sorunun cevabını buraya yazabilir veya aşağıdaki butonla AI'a yazdırabilirsiniz.",
                    label_visibility="collapsed",
                    key=f"ek15_textarea_{ph}",
                )
                hint_val = st.text_input(
                    "İpucu (opsiyonel, örn: 'öğrenci pişman ve suçu kabul ediyor')",
                    value=st.session_state["ek15_hints"].get(ph, ""),
                    key=f"ek15_hint_{ph}",
                )

                colb1, colb2 = st.columns([1,1])
                with colb1:
                    if st.button("🔮 Bu soru için cevap üret", key=f"gen_{ph}", use_container_width=True):
                        api_key = (st.session_state.get("api_key") or "").strip() or (api_key_input or "").strip()
                        if not api_key:
                            st.warning("Önce OpenAI API anahtarını girin.")
                        else:
                            with st.spinner("Cevap üretiliyor..."):
                                gen = generate_student_style_response(
                                    api_key=api_key,
                                    question_key=ph,
                                    transcript=st.session_state.get("current_transcript", ""),
                                    mapping=st.session_state.get("current_mapping", {}),
                                    user_hint=hint_val,
                                )
                            if gen:
                                st.session_state["current_mapping"][ph] = gen
                                st.session_state["ek15_hints"][ph] = hint_val
                                sm.update_session_data(st.session_state["current_session_id"], {ph: gen}, merge=True)
                                st.success("Cevap üretildi ve kaydedildi.")
                                st.rerun()
                with colb2:
                    if st.button("🧹 Temizle", key=f"clr_{ph}", use_container_width=True):
                        st.session_state["current_mapping"][ph] = ""
                        sm.update_session_data(st.session_state["current_session_id"], {ph: ""}, merge=True)
                        st.rerun()

        if st.button("✨ Hepsini üret (4 cevap)", type="primary", use_container_width=True):
            api_key = (st.session_state.get("api_key") or "").strip() or (api_key_input or "").strip()
            if not api_key:
                st.warning("Önce OpenAI API anahtarını girin.")
            else:
                made = 0
                with st.spinner("Dört cevap birden üretiliyor..."):
                    for ph, _title in qmap:
                        hint_val = st.session_state["ek15_hints"].get(ph, "")
                        gen = generate_student_style_response(
                            api_key=api_key,
                            question_key=ph,
                            transcript=st.session_state.get("current_transcript", ""),
                            mapping=st.session_state.get("current_mapping", {}),
                            user_hint=hint_val,
                        )
                        if gen:
                            st.session_state["current_mapping"][ph] = gen
                            sm.update_session_data(st.session_state["current_session_id"], {ph: gen}, merge=True)
                            made += 1
                st.success(f"✅ {made}/4 cevap üretildi ve kaydedildi.")
                st.rerun()

    # -------------------- DÜZENLEME / UYGULAMA --------------------
    if union_placeholders:
        st.markdown("---")
        st.subheader("✏️ Bilgi Düzenleme")

        col_apply, col_clear = st.columns([2, 1])
        with col_apply:
            if st.button("🔄 Session Verilerini Uygula"):
                session_data = sm.get_session(current_session_id)
                if session_data and session_data.get('extracted_data'):
                    current_mapping = st.session_state.get("current_mapping", {})
                    applied_count = 0
                    for ph in union_placeholders:
                        if ph in session_data['extracted_data'] and session_data['extracted_data'][ph]:
                            if ph not in current_mapping or not current_mapping.get(ph):
                                current_mapping[ph] = session_data['extracted_data'][ph]
                                applied_count += 1
                    st.session_state["current_mapping"] = current_mapping
                    if applied_count > 0:
                        st.success(f"✅ {applied_count} alan dolduruldu!")
                        st.rerun()

        with col_clear:
            if st.button("🧹 Temizle"):
                st.session_state["current_mapping"] = {}
                st.rerun()

        edit_cols = st.columns(2)
        for idx, ph in enumerate(sorted(list(union_placeholders))):
            with edit_cols[idx % 2]:
                display_name = ph.replace("{", "").replace("}", "")
                st.markdown(f"**{display_name}**")
                cur_val = st.session_state.get("current_mapping", {}).get(ph, "")
                new_val = st.text_input(
                    "Değer",
                    value=cur_val,
                    key=f"edit_{idx}_{ph}",
                    placeholder="Değer girin...",
                    label_visibility="collapsed"
                )
                if new_val != cur_val:
                    st.session_state["current_mapping"][ph] = new_val
                    sm.update_session_data(current_session_id, {ph: new_val}, merge=True)
                    updated_session = sm.get_session(current_session_id)
                    if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                        st.session_state["current_session_name"] = updated_session['session_name']
                    st.rerun()
                st.markdown("---")

    # -------------------- ŞABLON ÖNİZLEME --------------------
    if template_items:
        st.markdown("---")
        st.subheader("👁️ Seçilen Şablonların Önizlemesi")
        for template_name, template_data in template_items:
            with st.expander(f"📄 {template_name}", expanded=False):
                try:
                    doc = Document(io.BytesIO(template_data))
                    parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            parts.append(paragraph.text.strip())
                    for table in doc.tables:
                        for row in table.rows:
                            cells = []
                            for cell in row.cells:
                                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                                if cell_text:
                                    cells.append(cell_text)
                            if cells:
                                parts.append(" | ".join(cells))
                    for section in doc.sections:
                        if section.header:
                            for p in section.header.paragraphs:
                                if p.text.strip():
                                    parts.insert(0, f"[BAŞLIK: {p.text.strip()}]")
                        if section.footer:
                            for p in section.footer.paragraphs:
                                if p.text.strip():
                                    parts.append(f"[ALT BİLGİ: {p.text.strip()}]")

                    full_text = "\n\n".join(parts).strip()
                    if not full_text:
                        st.info("Bu şablonda görüntülenebilir metin bulunamadı.")
                    else:
                        current_mapping = st.session_state.get("current_mapping", {}) or {}
                        mapping_with_isbu = {**current_mapping, **today_isbu(datetime.now(IST))}
                        import html as _html
                        pattern = re.compile(r"\{[^}]+\}")
                        def _replace_placeholder(match: re.Match) -> str:
                            ph = match.group(0)
                            raw_val = str(mapping_with_isbu.get(ph, "")).strip()
                            if not raw_val:
                                key_nobraces = ph.strip('{}')
                                for k, v in mapping_with_isbu.items():
                                    if isinstance(k, str) and k.strip('{}').lower() == key_nobraces.lower():
                                        raw_val = str(v).strip()
                                        if raw_val:
                                            break
                            if raw_val:
                                return _html.escape(raw_val)
                            return f"<span style=\"color:#ff4d4f;font-weight:700;\">{_html.escape(ph)}</span>"
                        highlighted_text = pattern.sub(_replace_placeholder, full_text)
                        st.markdown(
                            f"""
                            <div style="white-space: pre-wrap; word-wrap: break-word; line-height: 1.75; font-size: 16px; font-weight: 500; color: #374151;">{highlighted_text}</div>
                            """,
                            unsafe_allow_html=True
                        )
                except Exception as e:
                    st.error(f"Şablon önizlemesi oluşturulamadı: {e}")

    # -------------------- BELGE OLUŞTURMA --------------------
    if template_items:
        st.subheader("📄 Belge Oluşturma")
        if st.button("📄 Tüm Belgeleri Oluştur", type="primary", use_container_width=True):
            if not st.session_state.get("current_mapping"):
                st.warning("Önce bilgileri doldurun.")
            else:
                try:
                    results = []
                    current_mapping = st.session_state["current_mapping"]
                    for idx, (name, data) in enumerate(template_items):
                        doc = Document(io.BytesIO(data))
                        mapping = {k: v for k, v in current_mapping.items() if str(v).strip()}
                        mapping = {**mapping, **today_isbu(datetime.now(IST))}
                        replaced = replace_placeholders_in_document(doc, mapping)
                        buf = io.BytesIO()
                        doc.save(buf)
                        out_bytes = buf.getvalue()
                        safe_session_name = re.sub(r'[^\w\s-]', '', current_session_name).strip()[:20]
                        out_name = f"{safe_session_name}_{os.path.splitext(name)[0]}.docx"
                        results.append({"name": out_name, "replaced": replaced, "data": out_bytes, "key": f"dl_{idx}_{out_name}"})
                    st.session_state["results"] = results
                    st.success("✅ Belgeler hazırlandı!")
                except Exception as e:
                    st.error(f"Belge oluşturma hatası: {e}")

        if st.session_state.get("results"):
            st.markdown("---")
            st.subheader("📥 İndirilecek Belgeler")
            for r in st.session_state["results"]:
                col_info, col_download = st.columns([3, 1])
                with col_info:
                    st.write(f"**{r['name']}** → {r['replaced']} değişiklik")
                with col_download:
                    st.download_button(
                        label="📥 İndir",
                        data=r["data"],
                        file_name=r["name"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=r["key"],
                        use_container_width=True
                    )

if __name__ == "__main__":
    main()

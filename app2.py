# app.py — Kopya Tutanak Sistemi (Streamlit + Session)
# Akış: 1) Ses  2) Transkript  3) Belge seçimi  4) Eksik alan formu  5) İşbu + İndir
# Ek: Öğrenci oturum kaydı (ad/no/TC/bölüm/hikaye), isim-no-TC ile arama, alanlara uygula
# Düzeltme: Sınav tarihini LLM'e bağımlı olmadan sağlam çözümle (tüm formatlar ve tüm belgeler)

import os, io, json
from datetime import datetime
from zoneinfo import ZoneInfo

import streamlit as st
from st_audiorec import st_audiorec


import dateparser
from docx import Document
from openai import OpenAI

# ================== Genel ==================
IST = ZoneInfo("Europe/Istanbul")
st.set_page_config(page_title="Kopya Tutanak Sistemi", page_icon="🧩", layout="wide")
st.title("🧩 Kopya Tutanak Sistemi (Session) ")

# ---- Şablon dizini (sağlam) ----
def detect_template_dir() -> str:
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    except Exception:
        base_dir = os.getcwd()
    cand1 = os.path.join(base_dir, "templates")
    cand2 = os.path.join(os.getcwd(), "templates")
    if os.path.isdir(cand1): return cand1
    if os.path.isdir(cand2): return cand2
    return cand1

TEMPLATE_DIR = detect_template_dir()
TEMPLATE_FILES = {
    "iybf": os.path.join(TEMPLATE_DIR, "İYBF-F029-_1-OLAY-TESPİT-TUTANAĞI kopyası.docx"),
    "ek1":  os.path.join(TEMPLATE_DIR, "Ek-1-Tutanak (Disiplin) kopyası.docx"),
    "ek2":  os.path.join(TEMPLATE_DIR, "Ek-2-Dilekçe (Disiplin) kopyası.docx"),
    "ek3":  os.path.join(TEMPLATE_DIR, "Ek-3-Bölüm başkanlığı (Disiplin) kopyası.docx"),
}

# ================== Session init ==================
DEFAULT_SESSION = {
    "api_key": "",
    "audio": None,
    "transcript": "",
    "extracted": None,             # dict (fields)
    "missing_questions": [],
    "extra_free_text": "",
    "selected_docs": {"iybf"},     # önce bir belge seçili görünsün

    # >>> Yeni: Öğrenci kayıtları (oturum boyunca)
    "students": {},                # key -> öğrenci dict
    "__search_results": [],        # son arama sonuçları (UI için)
}
for k, v in DEFAULT_SESSION.items():
    st.session_state.setdefault(k, v)

# ================== Yardımcılar ==================
def _keep_bytes(slot: str, file_or_bytes, default_name: str):
    if file_or_bytes is None: return
    if hasattr(file_or_bytes, "read"):
        data = file_or_bytes.read()
    else:
        data = file_or_bytes
    st.session_state[slot] = {"bytes": data, "name": default_name}

def today_isbu(dt: datetime | None = None):
    now = dt or datetime.now(IST)
    return {"isbu_gun": f"{now.day:02d}", "isbu_ay": f"{now.month:02d}", "isbu_yil": f"{now.year}", "isbu_saat": now.strftime("%H:%M")}

TR_DAYS = {0:"Pazartesi",1:"Salı",2:"Çarşamba",3:"Perşembe",4:"Cuma",5:"Cumartesi",6:"Pazar"}

def parse_tr_date(text):
    if not text: return None
    return dateparser.parse(text, languages=["tr"])

def split_date(dt: datetime):
    return {"gun": f"{dt.day:02d}", "ay": f"{dt.month:02d}", "yil": f"{dt.year}"}

def replace_placeholders_bytes(template_bytes: bytes, mapping: dict) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    def _rp(p):
        text = p.text
        changed = False
        for k, v in mapping.items():
            ph = "{"+k+"}"
            if ph in text:
                text = text.replace(ph, v if v is not None else "")
                changed = True
        if changed:
            for i in range(len(p.runs)-1, -1, -1):
                p.runs[i].text = ""
            (p.runs[0] if p.runs else p.add_run("")).text = text
    for p in doc.paragraphs: _rp(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _rp(p)
    out = io.BytesIO(); doc.save(out); out.seek(0); return out.read()

# ================== OpenAI ==================
def _client():
    key = st.session_state.get("api_key") or ""
    if not key: raise RuntimeError("API anahtarı girilmedi.")
    return OpenAI(api_key=key)

def transcribe_audio_bytes(audio_bytes: bytes, lang="tr") -> str:
    client = _client()
    f = io.BytesIO(audio_bytes); f.name = "audio.wav"
    tr = client.audio.transcriptions.create(model="whisper-1", file=f, language=lang, response_format="text")
    return tr

SYSTEM_PROMPT = """Sen bir bilgi çıkarım uzmanısın. Kullanıcıdan gelen Türkçe serbest anlatımdan alanları JSON olarak çıkarırsın.
- Türkçe ekleri doğru yorumla; ad-soyadları yalınlaştır.
- Tek tarih ifadesinden hem ISO (YYYY-MM-DD) hem {gun, ay, yil} hem hafta_gunu_tr üret.
- Saat bilgilerini hh:mm olarak ver.
- Sadece JSON döndür.
"""

def build_user_prompt(transkript: str, ek_bilgi: str = "") -> str:
    extra = f"\nEk bilgiler (kullanıcı düzeltmeleri/eklemeler):\n{ek_bilgi}\n" if ek_bilgi.strip() else ""
    return f"""
Aşağıdaki belgeleri dolduracağız ve bu placeholder alanlarını istiyoruz:
1) IYBF-F029: olay_tarihi, olay_saati, derslik, ders_adi, sinav_turu, ogrenci_adi, ogrenci_no, aciklama, kanit_1, kanit_2, kanit_3
2) Ek-1: sinav_gun, sinav_ay, sinav_yil, hafta_gunu, sinav_saati_baslangic, sinav_saati_bitis, ders_adi, ogrenci_no, ogrenci_adi_soyadi, aciklama, olay_saati, gozetmen_adi_1, gozetmen_adi_2, blok, derslik, bolum, bolum_baskanligi
3) Ek-2: sinav_gun, sinav_ay, sinav_yil, hafta_gunu, blok, derslik, ders_adi, ogrenci_no, ogrenci_tc, ogrenci_adi_soyadi, bolum
4) Ek-3: bolum_baskanligi, sinav_gun, sinav_ay, sinav_yil, hafta_gunu, blok, derslik, ders_adi, bolum, ogrenci_no, ogrenci_adi_soyadi, sinav_saati_baslangic, sinav_saati_bitis

Şema:
{{
  "fields": {{
    "ogrenci": {{"ad_soyad": "", "no": "", "tc": ""}},
    "ders": {{"adi": "", "sinav_turu": "", "blok": "", "derslik": ""}},
    "sinav": {{"tarih_iso": "", "gun": "", "ay": "", "yil": "", "hafta_gunu_tr": "", "saat_baslangic": "", "saat_bitis": ""}},
    "olay": {{"saat": "", "aciklama": "", "kanit_1": "", "kanit_2": "", "kanit_3": ""}},
    "gorevliler": {{"gozetmenler": ["", ""], "ogretim_elemani": "", "bolum": "", "bolum_baskanligi": ""}}
  }},
  "missing_questions": []
}}

Kullanıcı anlatımı:
\"\"\"{transkript}\"\"\"
{extra}
"""

def _find_first_json_object(text: str) -> str | None:
    if not text: return None
    text = text.strip().removeprefix("```json").removesuffix("```").strip("` \n\r\t")
    start = text.find("{");  depth = 0
    if start == -1: return None
    for i in range(start, len(text)):
        c = text[i]
        if c == "{": depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0: return text[start:i+1]
    return None

def _coerce_schema(obj: dict) -> dict:
    safe = {
        "fields": {
            "ogrenci": {"ad_soyad": "", "no": "", "tc": ""},
            "ders": {"adi": "", "sinav_turu": "", "blok": "", "derslik": ""},
            "sinav": {"tarih_iso": "", "gun": "", "ay": "", "yil": "", "hafta_gunu_tr": "", "saat_baslangic": "", "saat_bitis": ""},
            "olay": {"saat": "", "aciklama": "", "kanit_1": "", "kanit_2": "", "kanit_3": ""},
            "gorevliler": {"gozetmenler": ["", ""], "ogretim_elemani": "", "bolum": "", "bolum_baskanligi": ""}
        },
        "missing_questions": []
    }
    try:
        f = obj.get("fields", {})
        g = f.get("gorevliler", {})
        safe["fields"]["ogrenci"].update({k: "" if v is None else str(v) for k, v in f.get("ogrenci", {}).items()})
        safe["fields"]["ders"].update({k: "" if v is None else str(v) for k, v in f.get("ders", {}).items()})
        safe["fields"]["sinav"].update({k: "" if v is None else str(v) for k, v in f.get("sinav", {}).items()})
        safe["fields"]["olay"].update({k: "" if v is None else str(v) for k, v in f.get("olay", {}).items()})
        gl = g.get("gozetmenler", [])
        if not isinstance(gl, list): gl = [str(gl)]
        while len(gl) < 2: gl.append("")
        safe["fields"]["gorevliler"].update({
            "gozetmenler": [str(x) for x in gl[:2]],
            "ogretim_elemani": str(g.get("ogretim_elemani", "") or ""),
            "bolum": str(g.get("bolum", "") or ""),
            "bolum_baskanligi": str(g.get("bolum_baskanligi", "") or "")
        })
        mq = obj.get("missing_questions", [])
        if not isinstance(mq, list): mq = [str(mq)]
        safe["missing_questions"] = [str(x) for x in mq]
    except Exception:
        pass
    return safe

def extract_fields_with_llm(transkript: str, ek_bilgi: str = "", model="gpt-4o-mini") -> dict:
    client = _client()
    try:
        msg = client.chat.completions.create(
            model=model, temperature=0.1,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT + "\nSadece geçerli JSON üret."},
                {"role": "user", "content": build_user_prompt(transkript, ek_bilgi)}
            ]
        )
        content = msg.choices[0].message.content or ""
    except Exception:
        msg = client.chat.completions.create(
            model=model, temperature=0.1,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT + "\nSadece JSON döndür."},
                {"role": "user", "content": build_user_prompt(transkript, ek_bilgi)}
            ]
        )
        content = msg.choices[0].message.content or ""
    try:
        data = json.loads(content)
    except Exception:
        snip = _find_first_json_object(content)
        if not snip: raise ValueError("LLM yanıtı JSON değil/boş.")
        data = json.loads(snip)
    return _coerce_schema(data)

# ================== (Yeni) Öğrenci kayıt/arama yardımcıları ==================
def _student_key(rec: dict) -> str | None:
    tc = (rec or {}).get("tc", "").strip()
    no = (rec or {}).get("no", "").strip()
    ad = (rec or {}).get("ad_soyad", "").strip().lower()
    if tc: return f"tc:{tc}"
    if no: return f"no:{no}"
    if ad: return f"ad:{ad}"
    return None

def _normalize_student_from_fields(fields: dict, transcript: str) -> dict | None:
    if not fields: return None
    ogr = (fields.get("ogrenci") or {})
    gorev = (fields.get("gorevliler") or {})
    olay  = (fields.get("olay") or {})
    rec = {
        "ad_soyad": (ogr.get("ad_soyad") or "").strip(),
        "no": (ogr.get("no") or "").strip(),
        "tc": (ogr.get("tc") or "").strip(),
        "bolum": (gorev.get("bolum") or "").strip(),
        # hikaye: önce olay.aciklama, yoksa tüm transkript
        "story": (olay.get("aciklama") or transcript or "").strip(),
    }
    if not (rec["ad_soyad"] or rec["no"] or rec["tc"]):
        return None
    return rec

def upsert_student(rec: dict, *, also_index_aliases: bool = True) -> str | None:
    if not rec: return None
    key = _student_key(rec)
    if not key: return None
    students = st.session_state.setdefault("students", {})
    now = datetime.now(IST).isoformat(timespec="seconds")
    base = students.get(key, {})
    base.update({
        "ad_soyad": rec.get("ad_soyad", base.get("ad_soyad","")),
        "no": rec.get("no", base.get("no","")),
        "tc": rec.get("tc", base.get("tc","")),
        "bolum": rec.get("bolum", base.get("bolum","")),
        "story": rec.get("story", base.get("story","")),
        "last_seen_at": now,
    })
    students[key] = base
    if also_index_aliases:
        for alias_key in [
            f"tc:{base['tc']}" if base.get("tc") else None,
            f"no:{base['no']}" if base.get("no") else None,
            f"ad:{base['ad_soyad'].strip().lower()}" if base.get("ad_soyad") else None,
        ]:
            if alias_key and alias_key not in students:
                students[alias_key] = base
    return key

def search_students(query: str) -> list[tuple[str, dict]]:
    q = (query or "").strip().lower()
    if not q: return []
    seen = set()
    out = []
    for k, s in (st.session_state.get("students") or {}).items():
        if id(s) in seen:  # alias anahtarlar aynı objeyi gösterebilir
            continue
        haystack = " ".join([str(s.get("ad_soyad","")).lower(),
                             str(s.get("no","")).lower(),
                             str(s.get("tc","")).lower()])
        if q in haystack:
            out.append((k, s))
            seen.add(id(s))
    return out

def apply_student_to_extracted(student: dict):
    if not student: return
    st.session_state.setdefault("extracted", {})
    f = st.session_state["extracted"]
    f.setdefault("ogrenci", {})
    f.setdefault("gorevliler", {})
    f["ogrenci"]["ad_soyad"] = student.get("ad_soyad","")
    f["ogrenci"]["no"] = student.get("no","")
    f["ogrenci"]["tc"] = student.get("tc","")
    if student.get("bolum"):
        f["gorevliler"]["bolum"] = student["bolum"]

# ================== Sidebar ==================
with st.sidebar:
    st.markdown("### 🔑 OpenAI API Key")
    key_in = st.text_input("API Key", type="password", value=st.session_state["api_key"])
    if key_in != st.session_state["api_key"]: st.session_state["api_key"] = key_in

    st.markdown("### 📄 Şablonlar")
    missing = [k for k, p in TEMPLATE_FILES.items() if not os.path.exists(p)]
    if missing:
        st.error("Eksik şablonlar: " + ", ".join(missing))
        st.caption(f"templates/ klasörünü kontrol et: {TEMPLATE_DIR}")
        try:
            st.caption("Bulunan dosyalar:")
            for fname in os.listdir(TEMPLATE_DIR): st.caption(f"• {fname}")
        except Exception: st.caption("(templates klasörü bulunamadı)")
    else:
        for k, p in TEMPLATE_FILES.items(): st.caption(f"✅ {k.upper()} — {os.path.basename(p)}")

    # ===== Yeni: Öğrenci Kayıt / Arama UI =====
    st.markdown("---")
    with st.expander("👩‍🎓 Öğrenci Kayıt / Arama", expanded=True):
        st.caption("Öğrenciler oturum boyunca saklanır. No/TC unique olarak güncellenir.")

        flds_pref = (st.session_state.get("extracted") or {})
        ogr0 = (flds_pref.get("ogrenci") or {})
        gor0 = (flds_pref.get("gorevliler") or {})
        olay0 = (flds_pref.get("olay") or {})

        name_in = st.text_input("Ad Soyad", value=ogr0.get("ad_soyad",""))
        no_in   = st.text_input("Öğrenci No", value=ogr0.get("no",""))
        tc_in   = st.text_input("TC Kimlik No", value=ogr0.get("tc",""))
        bolum_in= st.text_input("Bölüm", value=gor0.get("bolum",""))
        story_in= st.text_area("Hikaye / Olay Açıklaması", value=(olay0.get("aciklama") or st.session_state.get("transcript","")), height=80)

        cols_sb = st.columns(2)
        if cols_sb[0].button("Kaydet / Güncelle", use_container_width=True):
            key = upsert_student({
                "ad_soyad": name_in, "no": no_in, "tc": tc_in, "bolum": bolum_in, "story": story_in
            })
            if key: st.success("Öğrenci kaydedildi/güncellendi.")
            else:   st.warning("En azından ad veya no/TC girilmelidir.")

        # Arama
        st.markdown("---")
        q = st.text_input("🔎 Ara (İsim / No / TC)")
        if st.button("Ara", use_container_width=True):
            st.session_state["__search_results"] = search_students(q)

        results = st.session_state.get("__search_results", [])
        if results:
            st.caption(f"Bulunan kayıtlar: {len(results)}")
            for idx, (k, s) in enumerate(results, start=1):
                with st.container(border=True):
                    st.markdown(f"**{idx}. {s.get('ad_soyad','(adsız)')}**")
                    st.write(f"- Öğrenci No: {s.get('no','')}")
                    st.write(f"- TC: {s.get('tc','')}")
                    if s.get("bolum"): st.write(f"- Bölüm: {s.get('bolum')}")
                    if s.get("last_seen_at"): st.write(f"- Son güncelleme: {s.get('last_seen_at')}")
                    if s.get("story"):
                        with st.expander("Hikaye"):
                            st.write(s["story"])
                    if st.button("Bu öğrenciyi alanlara uygula", key=f"apply_{idx}", use_container_width=True):
                        apply_student_to_extracted(s)
                        st.success("Öğrenci bilgileri alanlara uygulandı.")

# ================== 1) Ses ==================
st.header("1) Ses Kaydı")
colA, colB = st.columns([2,1], vertical_alignment="center")
with colA:
    wav_bytes = st_audiorec()
    if wav_bytes:
        _keep_bytes("audio", wav_bytes, "recording.wav")
        st.success("Mikrofon kaydı alındı (session'a yazıldı).")
with colB:
    uploaded_audio = st.file_uploader("Ses dosyası yükle (mp3, m4a, wav)", type=["mp3","m4a","wav"])
    if uploaded_audio:
        _keep_bytes("audio", uploaded_audio, uploaded_audio.name)
        st.success(f"Ses dosyası yüklendi: {uploaded_audio.name}")
if st.session_state["audio"]:
    st.audio(st.session_state["audio"]["bytes"], format="audio/wav")

# ================== 2) Transkript ==================
st.header("2) Transkript")
c1, c2 = st.columns(2)
with c1:
    if st.button("🎧 Sesi Metne Çevir"):
        if not st.session_state["api_key"]: st.error("API anahtarı gerekli."); st.stop()
        if not st.session_state["audio"]: st.error("Ses kaydı/yüklemesi yok."); st.stop()
        with st.spinner("Whisper transcribe..."):
            st.session_state["transcript"] = transcribe_audio_bytes(st.session_state["audio"]["bytes"])
        st.success("Transkript alındı.")
with c2:
    if st.button("🧠 Alanları Çıkar (LLM)"):
        if not st.session_state["api_key"]: st.error("API anahtarı gerekli."); st.stop()
        if not st.session_state["transcript"].strip(): st.error("Önce transkript üretin veya aşağıya yazın."); st.stop()
        with st.spinner("LLM alan çıkarımı..."):
            data = extract_fields_with_llm(st.session_state["transcript"], st.session_state["extra_free_text"])
            st.session_state["extracted"] = data["fields"]
            st.session_state["missing_questions"] = data.get("missing_questions", [])
        # >>> LLM'den sonra öğrenci otomatik kaydı
        stu = _normalize_student_from_fields(st.session_state["extracted"], st.session_state.get("transcript",""))
        if stu:
            upsert_student(stu)

st.session_state["transcript"] = st.text_area(
    "✍️ Transkript (düzenlenebilir)", st.session_state.get("transcript",""), height=160
)
st.session_state["extra_free_text"] = st.text_area(
    "➕ Serbest Ek Bilgi (opsiyonel)", st.session_state.get("extra_free_text",""), height=110
)

# ================== 3) Belge Seçimi ==================
st.header("3) Belge Seçimi")
b1, b2, b3, b4 = st.columns(4)
sel_iybf = b1.toggle("İYBF-F029", value=("iybf" in st.session_state["selected_docs"]))
sel_ek1  = b2.toggle("Ek-1", value=("ek1" in st.session_state["selected_docs"]))
sel_ek2  = b3.toggle("Ek-2", value=("ek2" in st.session_state["selected_docs"]))
sel_ek3  = b4.toggle("Ek-3", value=("ek3" in st.session_state["selected_docs"]))
st.session_state["selected_docs"] = {k for k, v in {"iybf":sel_iybf,"ek1":sel_ek1,"ek2":sel_ek2,"ek3":sel_ek3}.items() if v}

# ================== 4) Sadece Seçilen Belge(ler) için Eksik Alan Formu ==================
def required_fields_for(doc_key: str):
    if doc_key == "iybf":
        return {
            "ogrenci_adi": ("ogrenci", "ad_soyad"),
            "ogrenci_no": ("ogrenci", "no"),
            "ders_adi": ("ders", "adi"),
            "sinav_turu": ("ders", "sinav_turu"),
            "derslik": ("ders", "derslik"),
            "olay_tarihi": ("sinav", "tarih_iso"),
            "olay_saati": ("olay", "saat"),
            "aciklama": ("olay", "aciklama"),
            "kanit_1": ("olay", "kanit_1"),
            "kanit_2": ("olay", "kanit_2"),
            "kanit_3": ("olay", "kanit_3"),
        }
    if doc_key == "ek1":
        return {
            "ogrenci_adi_soyadi": ("ogrenci", "ad_soyad"),
            "ogrenci_no": ("ogrenci", "no"),
            "ders_adi": ("ders", "adi"),
            "blok": ("ders", "blok"),
            "derslik": ("ders", "derslik"),
            "sinav_saati_baslangic": ("sinav", "saat_baslangic"),
            "sinav_saati_bitis": ("sinav", "saat_bitis"),
            "olay_saati": ("olay", "saat"),
            "aciklama": ("olay", "aciklama"),
            "gozetmen_adi_1": ("gorevliler", "gozetmenler", 0),
            "gozetmen_adi_2": ("gorevliler", "gozetmenler", 1),
            "bolum": ("gorevliler", "bolum"),
            "bolum_baskanligi": ("gorevliler", "bolum_baskanligi"),
            "sinav_tarih_iso": ("sinav", "tarih_iso"),
        }
    if doc_key == "ek2":
        return {
            "ogrenci_adi_soyadi": ("ogrenci", "ad_soyad"),
            "ogrenci_no": ("ogrenci", "no"),
            "ogrenci_tc": ("ogrenci", "tc"),
            "ders_adi": ("ders", "adi"),
            "blok": ("ders", "blok"),
            "derslik": ("ders", "derslik"),
            "bolum": ("gorevliler", "bolum"),
            "sinav_tarih_iso": ("sinav", "tarih_iso"),
        }
    if doc_key == "ek3":
        return {
            "ogrenci_adi_soyadi": ("ogrenci", "ad_soyad"),
            "ogrenci_no": ("ogrenci", "no"),
            "ogrenci_tc": ("ogrenci", "tc"),
            "ders_adi": ("ders", "adi"),
            "blok": ("ders", "blok"),
            "derslik": ("ders", "derslik"),
            "bolum": ("gorevliler", "bolum"),
            "bolum_baskanligi": ("gorevliler", "bolum_baskanligi"),
            "sinav_saati_baslangic": ("sinav", "saat_baslangic"),
            "sinav_saati_bitis": ("sinav", "saat_bitis"),
            "sinav_tarih_iso": ("sinav", "tarih_iso"),
        }
    return {}

def read_value(flds: dict, path: tuple):
    cur = flds
    for p in path:
        if isinstance(p, int):
            if not isinstance(cur, list) or len(cur) <= p: return ""
            cur = cur[p]
        else:
            cur = (cur or {}).get(p, "")
    return cur or ""

def render_missing_form_for(doc_key: str, col):
    flds = st.session_state.get("extracted")
    if not flds:
        col.info("Önce 2. adımda **Alanları Çıkar (LLM)** düğmesine basın.")
        return None
    req = required_fields_for(doc_key)
    empties = {}
    with col:
        st.markdown(f"**{doc_key.upper()} için eksik alanlar**")
        for nice_name, path in req.items():
            if nice_name in {"sinav_gun","sinav_ay","sinav_yil"}:
                continue
            val = read_value(flds, path)
            if str(val).strip() == "":
                empties[nice_name] = st.text_input(
                    nice_name.replace("_"," ").title(),
                    key=f"{doc_key}_{nice_name}",
                    placeholder=("24 Ağustos 2024 / 24.08.2024 / 24/08/2024" if "tarih" in nice_name else None)
                )
        if not empties:
            st.success("Bu belge için eksik alan bulunmuyor. İndirebilirsiniz.")
    return empties

st.header("4) Eksik Alanlar (Sadece Seçilen Belge)")
selected = list(st.session_state["selected_docs"])
cols = st.columns(len(selected) or 1)
empties_all = {}
for i, key in enumerate(["iybf","ek1","ek2","ek3"]):
    if key in st.session_state["selected_docs"]:
        empties = render_missing_form_for(key, cols[selected.index(key) if key in selected and len(selected)>1 else i % len(cols)])
        if empties is not None:
            empties_all[key] = empties

if empties_all and st.button("➕ Eksik Alan Yanıtlarını Uygula ve Yeniden Çıkar"):
    ek_bilgi = st.session_state["extra_free_text"].strip()
    for doc_key, fields in empties_all.items():
        for name in fields.keys():
            ans = st.session_state.get(f"{doc_key}_{name}", "").strip()
            if ans:
                ek_bilgi += f"\n{doc_key}.{name}: {ans}"
    if not ek_bilgi.strip():
        st.warning("Eksik alan girdisi yok.")
    else:
        with st.spinner("Yanıtlarla tekrar çıkarım..."):
            data2 = extract_fields_with_llm(st.session_state["transcript"], ek_bilgi=ek_bilgi)
            st.session_state["extracted"] = data2["fields"]
            st.session_state["missing_questions"] = data2.get("missing_questions", [])
            st.success("Alanlar güncellendi ✅")
        # tekrar çıkarımdan sonra da öğrenciyi güncelle
        stu2 = _normalize_student_from_fields(st.session_state["extracted"], st.session_state.get("transcript",""))
        if stu2:
            upsert_student(stu2)

# --------- YENİ: LLM'den bağımsız sınav tarihi çözücü ----------
def resolve_exam_date(extracted_fields: dict) -> tuple[dict, str]:
    """
    Sınav tarihini sağlam şekilde çöz:
    1) extracted.sinav.tarih_iso
    2) extracted.sinav.{gun,ay,yil}
    3) Kullanıcı formu: *sinav_tarih_iso alanları (ek1_/ek2_/ek3_)
    Dönen: (split_dict, hafta_gunu)  -> {"gun":"dd","ay":"MM","yil":"YYYY"}, "Pazartesi"...
    """
    sinav_iso = (extracted_fields.get("sinav") or {}).get("tarih_iso","").strip()
    dt = parse_tr_date(sinav_iso) if sinav_iso else None

    if dt is None:
        s = extracted_fields.get("sinav") or {}
        g, a, y = (s.get("gun","").strip(), s.get("ay","").strip(), s.get("yil","").strip())
        if g and a and y and g.isdigit() and a.isdigit() and y.isdigit():
            try:
                dt = datetime(int(y), int(a), int(g), tzinfo=IST)
            except Exception:
                dt = None

    if dt is None:
        for key in st.session_state.keys():
            if key.endswith("_sinav_tarih_iso"):
                val = str(st.session_state.get(key) or "").strip()
                if val:
                    dt = parse_tr_date(val)
                    if dt: break

    if dt:
        split = split_date(dt)
        hafta = TR_DAYS[dt.weekday()]
        return split, hafta
    return {"gun":"", "ay":"", "yil":""}, ""

# ================== 5) İşbu Tarih/Saat ve İndir ==================
st.header("5) İşbu Tarih/Saat ve İndir")
left, right = st.columns([1,3])
with left:
    with st.expander("İşbu tarih/saat (opsiyonel)"):
        use_custom = st.checkbox("Özel tarih/saat kullan", value=False, key="use_custom_isbu")
        c_date = st.date_input("Tarih", value=datetime.now(IST).date(), key="isbu_date")
        c_time = st.time_input("Saat", value=datetime.now(IST).time().replace(second=0, microsecond=0), key="isbu_time")

with right:
    st.caption("İndirme anında 'işbu gün/ay/yıl/saat' hesaplanır. Manuel seçerseniz o kullanılır.")

    def _make_docx(which_key: str) -> bytes | None:
        if not st.session_state.get("extracted"): return None
        p = TEMPLATE_FILES.get(which_key)
        if not p or not os.path.exists(p): return None
        with open(p, "rb") as f: tpl = f.read()

        # İşbu
        dt = None
        if st.session_state.get("use_custom_isbu"):
            d = st.session_state.get("isbu_date"); t = st.session_state.get("isbu_time")
            dt = datetime.combine(d, t, tzinfo=IST)
        isbu = today_isbu(dt)

        flds = st.session_state["extracted"]

        # Tarih çözümleme (LLM bağımsız)
        sinav_split, hafta = resolve_exam_date(flds)

        ortak = {
            "sinav_gun": sinav_split["gun"], "sinav_ay": sinav_split["ay"], "sinav_yil": sinav_split["yil"], "hafta_gunu": hafta,
            "ders_adi": (flds.get("ders") or {}).get("adi",""),
            "ogrenci_no": (flds.get("ogrenci") or {}).get("no",""),
            "ogrenci_adi_soyadi": (flds.get("ogrenci") or {}).get("ad_soyad",""),
            "blok": (flds.get("ders") or {}).get("blok",""),
            "derslik": (flds.get("ders") or {}).get("derslik",""),
            "sinav_saati_baslangic": (flds.get("sinav") or {}).get("saat_baslangic",""),
            "sinav_saati_bitis": (flds.get("sinav") or {}).get("saat_bitis",""),
            "bolum": (flds.get("gorevliler") or {}).get("bolum",""),
            "bolum_baskanligi": (flds.get("gorevliler") or {}).get("bolum_baskanligi",""),
        }

        # İYBF
        map_iybf = {
            "olay_tarihi": (flds.get("sinav") or {}).get("tarih_iso","") or "/".join([sinav_split["yil"], sinav_split["ay"], sinav_split["gun"]]),
            "olay_saati": (flds.get("olay") or {}).get("saat",""),
            "derslik": (flds.get("ders") or {}).get("derslik",""),
            "ders_adi": (flds.get("ders") or {}).get("adi",""),
            "sinav_turu": (flds.get("ders") or {}).get("sinav_turu",""),
            "ogrenci_adi": (flds.get("ogrenci") or {}).get("ad_soyad",""),
            "ogrenci_no": (flds.get("ogrenci") or {}).get("no",""),
            "aciklama": (flds.get("olay") or {}).get("aciklama",""),
            "kanit_1": (flds.get("olay") or {}).get("kanit_1",""),
            "kanit_2": (flds.get("olay") or {}).get("kanit_2",""),
            "kanit_3": (flds.get("olay") or {}).get("kanit_3",""),
        }

        gl = (flds.get("gorevliler") or {}).get("gozetmenler", []) or []
        map_ek1 = {**ortak, **isbu, "olay_saati": (flds.get("olay") or {}).get("saat",""),
                   "aciklama": (flds.get("olay") or {}).get("aciklama",""),
                   "gozetmen_adi_1": (gl[0] if len(gl)>0 else ""), "gozetmen_adi_2": (gl[1] if len(gl)>1 else "")}
        map_ek2 = {**ortak, **isbu, "ogrenci_tc": (flds.get("ogrenci") or {}).get("tc","")}
        map_ek3 = {**ortak, **isbu , "ogrenci_tc": (flds.get("ogrenci") or {}).get("tc","")}

        mapping = {"iybf": map_iybf, "ek1": map_ek1, "ek2": map_ek2, "ek3": map_ek3}[which_key]
        return replace_placeholders_bytes(tpl, mapping)

    labels = {
        "iybf": ("⬇️ İYBF-F029", "IYBF-F029_doldurulmus.docx"),
        "ek1":  ("⬇️ Ek-1 Tutanak", "Ek-1_Tutanak_doldurulmus.docx"),
        "ek2":  ("⬇️ Ek-2 Dilekçe", "Ek-2_Dilekce_doldurulmus.docx"),
        "ek3":  ("⬇️ Ek-3 Bölüm Başkanlığı", "Ek-3_BolumBaskanligi_doldurulmus.docx"),
    }

    cols_dl = st.columns(4)
    j = 0
    for key in ["iybf","ek1","ek2","ek3"]:
        if key not in st.session_state["selected_docs"]: continue
        with cols_dl[j % 4]:
            data = _make_docx(key)
            if data:
                st.download_button(labels[key][0], data=data,
                    file_name=labels[key][1],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.caption(f"⚠️ {key.upper()} için veri/şablon eksik.")
        j += 1
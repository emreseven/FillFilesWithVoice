import io
import os
import re
import json
import sys
import tempfile
import importlib.util
import traceback
from typing import Dict, List, Optional, Set, Tuple

import streamlit as st
from docx import Document

# Local session management import
from local_session_manager import get_local_session_manager, merge_extracted_data, detect_conflicts


MIC_IMPORT_ERROR: Optional[str] = None
try:
    # Optional mic recorder; if unavailable, show diagnostics in UI
    from streamlit_mic_recorder import mic_recorder  # type: ignore
except Exception as _mic_e:  # pragma: no cover - optional dependency
    mic_recorder = None  # type: ignore
    MIC_IMPORT_ERROR = f"{type(_mic_e).__name__}: {str(_mic_e)}\n{traceback.format_exc()}"

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None  # type: ignore


def bytes_from_mic_return(value) -> Optional[bytes]:
    """
    Normalize various return types from mic_recorder into raw WAV/PCM bytes.
    Known shapes: dict with 'bytes' key, bytes directly, or list/np arrays (not supported here).
    """
    if value is None:
        return None
    if isinstance(value, dict) and "bytes" in value:
        return value["bytes"]
    if isinstance(value, (bytes, bytearray)):
        return bytes(value)
    return None


def transcribe_audio_bytes(audio_bytes: bytes, api_key: str, file_ext: str = "wav") -> Optional[str]:
    """
    Transcribe raw audio bytes using OpenAI's transcription API.
    Writes to a temporary file and calls the API. Returns text or None on failure.
    """
    if OpenAI is None:
        st.error("OpenAI SDK is not available. Please ensure 'openai' is installed.")
        return None

    tmp_path = None
    try:
        client = OpenAI(api_key=api_key)
        
        # Güvenli temp dosya oluşturma - ASCII safe
        import uuid
        safe_filename = f"audio_{uuid.uuid4().hex}.{file_ext}"
        tmp_dir = tempfile.gettempdir()
        tmp_path = os.path.join(tmp_dir, safe_filename)
        
        # Dosyayı binary modda yaz
        with open(tmp_path, "wb") as f:
            f.write(audio_bytes)

        # OpenAI API çağrısı
        with open(tmp_path, "rb") as f:
            resp = client.audio.transcriptions.create(
                model="whisper-1", 
                file=f,
                language="tr",  # Türkçe dil kodu
                response_format="text"  # Sadece text döndür
            )

        # Cevabı işle
        if isinstance(resp, str):
            text = resp
        else:
            text = getattr(resp, "text", None) or (resp.get("text") if isinstance(resp, dict) else None)
        
        return text if text else None
        
    except Exception as e:
        error_msg = str(e)
        # Unicode karakterleri güvenli göster
        try:
            error_msg = error_msg.encode('ascii', errors='replace').decode('ascii')
        except:
            error_msg = "Encoding error occurred"
        st.error(f"Ses metne çevrilemedi: {error_msg}")
        return None
    finally:
        # Temp dosyayı temizle
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass


def replace_placeholders_in_paragraph(paragraph, placeholder_values: Dict[str, str]) -> int:
    if not paragraph.runs:
        return 0
    original_text = "".join(run.text for run in paragraph.runs)
    replaced_text = original_text
    total_replacements = 0
    for placeholder, value in placeholder_values.items():
        if value is None:
            continue
        count = replaced_text.count(placeholder)
        if count:
            replaced_text = replaced_text.replace(placeholder, str(value))
            total_replacements += count
    if replaced_text != original_text:
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        paragraph.add_run(replaced_text)
    return total_replacements


def replace_placeholders_in_document(doc: Document, placeholder_values: Dict[str, str]) -> int:
    replacements_made = 0
    # Body paragraphs
    for p in doc.paragraphs:
        replacements_made += replace_placeholders_in_paragraph(p, placeholder_values)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replacements_made += replace_placeholders_in_paragraph(p, placeholder_values)
    # Headers/Footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replacements_made += replace_placeholders_in_paragraph(p, placeholder_values)
        if section.footer:
            for p in section.footer.paragraphs:
                replacements_made += replace_placeholders_in_paragraph(p, placeholder_values)
    return replacements_made


def extract_placeholders_from_docx_bytes(file_bytes: bytes) -> Tuple[Set[str], str]:
    """
    Return (placeholders, concatenated_text) from a .docx file in bytes.
    """
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text += p.text + " "
    placeholders = set(re.findall(r"\{[^}]+\}", text))
    return placeholders, text


def collect_text_blocks(doc: Document) -> List[str]:
    blocks: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        blocks.append(p.text)
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    blocks.append(p.text)
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    blocks.append(p.text)
    return blocks


def generate_preview_text(file_bytes: bytes, placeholder_values: Dict[str, str]) -> str:
    """
    Şablon dosyasını placeholder değerleriyle doldurarak önizleme metni üretir.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        preview_text = ""
        
        # Paragrafları işle
        for p in doc.paragraphs:
            if p.text.strip():
                text = p.text
                # Placeholder'ları değerlerle değiştir
                for placeholder, value in placeholder_values.items():
                    if value and str(value).strip():
                        text = text.replace(placeholder, str(value))
                preview_text += text + "\n\n"
        
        # Tabloları işle
        for table in doc.tables:
            preview_text += "\n[TABLO]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_content = ""
                    for p in cell.paragraphs:
                        if p.text.strip():
                            text = p.text
                            # Placeholder'ları değerlerle değiştir
                            for placeholder, value in placeholder_values.items():
                                if value and str(value).strip():
                                    text = text.replace(placeholder, str(value))
                            cell_content += text + " "
                    row_text.append(cell_content.strip())
                preview_text += " | ".join(row_text) + "\n"
            preview_text += "[/TABLO]\n\n"
        
        # Header/Footer işle
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        text = p.text
                        for placeholder, value in placeholder_values.items():
                            if value and str(value).strip():
                                text = text.replace(placeholder, str(value))
                        preview_text += f"[HEADER] {text}\n"
            if section.footer:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        text = p.text
                        for placeholder, value in placeholder_values.items():
                            if value and str(value).strip():
                                text = text.replace(placeholder, str(value))
                        preview_text += f"[FOOTER] {text}\n"
        
        return preview_text.strip()
    except Exception as e:
        return f"Önizleme oluşturulamadı: {e}"


def extract_placeholder_contexts_from_docx_bytes(file_bytes: bytes, placeholders: Set[str], window: int = 70) -> Dict[str, List[str]]:
    """
    Return mapping of placeholder -> list of short context snippets around each occurrence
    in the document, to help the model infer context-aware values.
    """
    doc = Document(io.BytesIO(file_bytes))
    blocks = collect_text_blocks(doc)
    contexts: Dict[str, List[str]] = {ph: [] for ph in placeholders}
    for block in blocks:
        for ph in placeholders:
            pattern = re.escape(ph)
            for m in re.finditer(pattern, block):
                start, end = m.start(), m.end()
                before = block[max(0, start - window): start]
                after = block[end: end + window]
                snippet = f"{before}{ph}{after}"
                lst = contexts.setdefault(ph, [])
                if len(lst) < 3:
                    lst.append(snippet)
    return contexts


def aggregate_contexts_across_templates(templates: List[Tuple[str, bytes]], placeholders: Set[str]) -> Dict[str, List[str]]:
    combined: Dict[str, List[str]] = {ph: [] for ph in placeholders}
    for _, data in templates:
        try:
            local_ctx = extract_placeholder_contexts_from_docx_bytes(data, placeholders)
            for ph, lst in local_ctx.items():
                for s in lst:
                    if len(combined[ph]) < 5:
                        combined[ph].append(s)
        except Exception:
            continue
    return combined


def parse_json_loose(s: str) -> Dict[str, str]:
    try:
        return json.loads(s)
    except Exception:
        pass
    try:
        m = re.search(r"\{[\s\S]*\}", s)
        if m:
            return json.loads(m.group(0))
    except Exception:
        pass
    return {}


def infer_placeholder_values(
    transcript: str,
    placeholders: Set[str],
    contexts: Dict[str, List[str]],
    api_key: str,
    model: str = "gpt-4o-mini",
) -> Dict[str, str]:
    if OpenAI is None:
        st.error("OpenAI SDK is not available. Lütfen 'openai' paketini kurun.")
        return {}
    client = OpenAI(api_key=api_key)
    ph_list = sorted(list(placeholders))
    
    # Gelişmiş prompt - template bağlamını ve ses transkriptini analiz et
    prompt_text = f"""
SES TRANSKRİPTİ:
"{transcript}"

TEMPLATE PLACEHOLDER'LARI VE BAĞLAMLARI:
"""
    for ph in ph_list:
        if ph in contexts and contexts[ph]:
            # Tüm bağlam örneklerini göster
            context_examples = "\n".join([f"  • {ctx[:200]}" for ctx in contexts[ph][:3]])
            prompt_text += f"\n{ph}:\n{context_examples}\n"
        else:
            prompt_text += f"\n{ph}: (Bağlam bulunamadı)\n"
    
    prompt_text += """

GÖREV:
1. Ses transkriptini analiz et
2. Her placeholder için template bağlamını incele (placeholder'ın etrafındaki kelimeler ne anlama geliyor?)
3. Bağlama uygun değerleri ses transkriptinden çıkar
4. Template'teki placeholder olmayan kısımları da dikkate al (örn: "Sayın {{isim}}, {{tarih}} tarihinde..." → isim ve tarih arasındaki ilişkiyi anla)

KURALLAR:
- Sadece ses transkriptinden çıkarabileceğin bilgileri kullan
- Bağlama tam uygun değerler ver
- Çıkaramadığın bilgiler için boş string ("") bırak
- SADECE JSON formatında cevap ver

JSON formatı örneği:
""" + "{" + ", ".join([f'"{ph}": "değer_veya_boş_string"' for ph in ph_list[:3]]) + "...}"

    messages = [
        {"role": "system", "content": "Sen uzman bir belge analiz asistanısın. Ses transkriptini ve belge bağlamını analiz ederek doğru bilgileri çıkarırsın. Sadece JSON döndür."},
        {"role": "user", "content": prompt_text},
    ]
    
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.1,  # Daha tutarlı sonuçlar için düşük
        )
        content = resp.choices[0].message.content if resp and resp.choices else "{}"
        data = parse_json_loose(content or "{}")
        
        # Sadece gerçek değerleri döndür, varsayılan yok
        result = {}
        for ph in placeholders:
            if ph in data and str(data[ph]).strip():
                result[ph] = str(data[ph]).strip()
            else:
                result[ph] = ""  # Boş bırak, varsayılan değer yok
        
        return result
    except Exception as e:
        st.error(f"Ses analizi başarısız: {e}")
        return {ph: "" for ph in placeholders}  # Hata durumunda boş değerler


def extract_student_info(session_data):
    """Session verisinden öğrenci bilgilerini çıkar"""
    if not session_data or not session_data.get('extracted_data'):
        return None, None
    
    extracted = session_data['extracted_data']
    student_no = None
    student_name = None
    
    # Öğrenci numarası ve isim için olası alanları kontrol et
    for key, value in extracted.items():
        if value and str(value).strip():
            key_lower = key.lower().replace('{', '').replace('}', '')
            value_str = str(value).strip()
            
            # Öğrenci numarası - öncelik sırası
            if 'ogrencino' in key_lower or 'ogrenci_no' in key_lower:
                # Öğrenci numarası alanı (en yüksek öncelik)
                student_no = value_str
            elif 'tcno' in key_lower and len(value_str) <= 6:
                # TC no alanı ama kısa ise (muhtemelen öğrenci numarası olarak kullanılmış)
                student_no = value_str
            elif any(keyword in key_lower for keyword in ['sinifno', 'sınıfno', 'class_no']):
                # Sınıf numarası da öğrenci numarası olabilir
                if not student_no:  # Daha önce öğrenci numarası bulunmamışsa
                    student_no = value_str
            
            # Öğrenci adı - ad ve soyad birleştir
            elif 'ad' in key_lower and 'soyad' not in key_lower:
                # Ad alanı
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{value_str} {student_name}"
                    else:
                        student_name = value_str
            elif 'soyad' in key_lower:
                # Soyad alanı
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{student_name} {value_str}"
                    else:
                        student_name = value_str
            elif any(keyword in key_lower for keyword in ['isim', 'name']):
                # Genel isim alanı
                if not value_str.isdigit():
                    student_name = value_str
    
    return student_no, student_name

def update_session_name_if_needed(session_id, session_data):
    """Öğrenci bilgileri varsa session ismini otomatik güncelle"""
    try:
        student_no, student_name = extract_student_info(session_data)
        
        # Eğer öğrenci bilgileri varsa ve session ismi hala otomatik isimse
        current_name = session_data.get('session_name', '')
        if (student_no or student_name) and ('Yeni Session' in current_name or current_name.startswith('Session')):
            # Yeni session ismi oluştur - Öncelik: Öğrenci No + İsim
            if student_no and student_name:
                new_name = f"{student_no} - {student_name}"
            elif student_name:
                new_name = f"{student_name}"
            elif student_no:
                new_name = f"Öğrenci No: {student_no}"
            else:
                return False
            
            # Session ismini güncelle
            sm = get_local_session_manager()
            if hasattr(sm, 'update_session_name'):
                return sm.update_session_name(session_id, new_name)
            else:
                # Eğer update_session_name metodu yoksa, session verisini manuel güncelle
                session_data['session_name'] = new_name
                return sm.save_session(session_id, session_data)
                
    except Exception as e:
        st.error(f"Session ismi güncellenirken hata: {e}")
        return False
    
    return False

def show_session_manager():
    """Session seçim ve yönetim arayüzü"""
    st.title("🎯 Sesli Belge Doldurma - Session Yöneticisi")
    st.caption("Öğrenci bilgilerini sesli girdi ile otomatik kaydeden akıllı sistem")
    
    sm = get_local_session_manager()
    
    # Arama çubuğu
    search_term = st.text_input("🔍 Öğrenci Ara (İsim veya Numara)", placeholder="Öğrenci adı veya numarası yazın...")
    
    # Session listesi
    sessions = sm.get_all_sessions()
    
    # Arama filtresi uygula
    if search_term:
        filtered_sessions = []
        for session in sessions:
            student_no, student_name = extract_student_info(session)
            if (student_no and search_term.lower() in student_no.lower()) or \
               (student_name and search_term.lower() in student_name.lower()) or \
               search_term.lower() in session['session_name'].lower():
                filtered_sessions.append(session)
        sessions = filtered_sessions
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.subheader("📁 Öğrenci Session'ları")
        
        if not sessions:
            if search_term:
                st.info("Arama kriterinize uygun session bulunamadı.")
            else:
                st.info("Henüz hiç session oluşturulmamış. 'Session Başlat' butonuna tıklayın.")
        else:
            for session in sessions:
                student_no, student_name = extract_student_info(session)
                
                # Session başlığını öğrenci bilgilerine göre oluştur
                if student_no and student_name:
                    display_title = f"👤 {student_no} - {student_name}"
                elif student_name:
                    display_title = f"👤 {student_name}"
                elif student_no:
                    display_title = f"👤 Öğrenci No: {student_no}"
                else:
                    display_title = f"📄 {session['session_name']}"
                
                with st.expander(display_title, expanded=False):
                    col_info, col_actions = st.columns([2, 1])
                    
                    with col_info:
                        st.write(f"**Oluşturma:** {session['created_date'][:10]}")
                        
                        # Öğrenci bilgilerini göster
                        if student_no:
                            st.write(f"**Öğrenci No:** {student_no}")
                        if student_name:
                            st.write(f"**Öğrenci Adı:** {student_name}")
                    
                    with col_actions:
                        if st.button(f"🚀 Aç", key=f"open_{session['session_id']}"):
                            st.session_state["current_session_id"] = session['session_id']
                            st.session_state["current_session_name"] = session['session_name']
                            st.session_state["page"] = "voice_app"
                            st.rerun()
                        
                        # Silme onay sistemi
                        if st.button(f"🗑️ Sil", key=f"delete_{session['session_id']}"):
                            st.session_state[f"confirm_delete_{session['session_id']}"] = True
                            st.rerun()
                        
                        # Onay mesajı göster
                        if st.session_state.get(f"confirm_delete_{session['session_id']}", False):
                            st.warning("⚠️ Bu session'ı silmek istediğinizden emin misiniz?")
                            col_yes, col_no = st.columns(2)
                            
                            with col_yes:
                                if st.button("✅ Evet, Sil", key=f"confirm_yes_{session['session_id']}", type="primary"):
                                    if sm.delete_session(session['session_id']):
                                        st.success("Session silindi!")
                                        # Onay state'ini temizle
                                        if f"confirm_delete_{session['session_id']}" in st.session_state:
                                            del st.session_state[f"confirm_delete_{session['session_id']}"]
                                        st.rerun()
                                    else:
                                        st.error("Session silinemedi!")
                            
                            with col_no:
                                if st.button("❌ İptal", key=f"confirm_no_{session['session_id']}"):
                                    # Onay state'ini temizle
                                    if f"confirm_delete_{session['session_id']}" in st.session_state:
                                        del st.session_state[f"confirm_delete_{session['session_id']}"]
                                    st.rerun()
    
    with col2:
        st.subheader("🚀 Yeni Session")
        st.write("Yeni bir öğrenci için session başlatın. Öğrenci bilgileri sesli girdi ile otomatik kaydedilecek.")
        
        if st.button("📝 Session Başlat", type="primary", use_container_width=True):
            # Otomatik session oluştur
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            session_name = f"Yeni Session - {timestamp}"
            
            session_id = sm.create_session(session_name)
            if session_id:
                st.success("Yeni session başlatıldı! Öğrenci bilgilerini sesli girdi ile ekleyebilirsiniz.")
                st.session_state["current_session_id"] = session_id
                st.session_state["current_session_name"] = session_name
                st.session_state["page"] = "voice_app"
                st.rerun()
            else:
                st.error("Session oluşturulamadı!")
        
        st.info("💡 **İpucu:** Session başlattıktan sonra öğrencinin adını ve numarasını söyleyerek kaydedin. Session otomatik olarak öğrenci bilgileriyle adlandırılacak.")


def show_voice_app():
    """Ana voice app arayüzü (session context'inde)"""
    # Session bilgilerini al
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return
    
    # Session verilerini yükle
    sm = get_local_session_manager()
    session_data = sm.get_session(current_session_id)
    
    if not session_data:
        st.error("Session verisi yüklenemedi!")
        return
    
    # Session state'leri initialize et - session verilerini yükle
    if "single_mapping" not in st.session_state:
        st.session_state["single_mapping"] = session_data['extracted_data'].copy()
    else:
        # Mevcut session verilerini kontrol et ve eksik olanları ekle
        session_extracted = session_data.get('extracted_data', {})
        current_mapping = st.session_state["single_mapping"]
        
        # Session'dan eksik verileri al
        for key, value in session_extracted.items():
            if key not in current_mapping and value:
                current_mapping[key] = value
        
        st.session_state["single_mapping"] = current_mapping
    if "single_transcript" not in st.session_state:
        st.session_state["single_transcript"] = ""
    if "single_results" not in st.session_state:
        st.session_state["single_results"] = []

    # Header
    col_title, col_back = st.columns([4, 1])
    with col_title:
        st.title(f"🎯 {current_session_name}")
        st.caption(f"Session ID: {current_session_id}")
    
    with col_back:
        if st.button("🏠 Session Yöneticisi"):
            st.session_state["page"] = "session_manager"
            st.rerun()
    
    st.markdown("---")


def main() -> None:
    st.set_page_config(page_title="🎯 Sesli Belge Doldurma", page_icon="🎯", layout="wide")

    # Session state initialization
    if "page" not in st.session_state:
        st.session_state["page"] = "session_manager"
    if "current_session_id" not in st.session_state:
        st.session_state["current_session_id"] = None
    if "current_session_name" not in st.session_state:
        st.session_state["current_session_name"] = ""
    if "api_key" not in st.session_state:
        st.session_state["api_key"] = ""

    # Page routing
    if st.session_state["page"] == "session_manager":
        show_session_manager()
        return
    elif st.session_state["page"] == "voice_app":
        show_voice_app()
        # Ana uygulama devamı aşağıda
    else:
        st.session_state["page"] = "session_manager"
        st.rerun()



    api_c1, api_c2 = st.columns([2, 1])
    with api_c1:
        api_key_input = st.text_input(
            "OpenAI API Key",
            value=st.session_state.get("api_key", ""),
            type="password",
            help="Whisper ve LLM için gereklidir.",
        )
    with api_c2:
        if st.checkbox("Bu oturumda hatırla", value=bool(st.session_state.get("api_key"))):
            st.session_state["api_key"] = api_key_input
        else:
            st.session_state["api_key"] = ""

    st.markdown("---")

    st.subheader("📝 Şablon Belgeleri")
    
    # Sabit templates klasöründen şablonları listele
    default_dir = os.path.join(os.getcwd(), "templates")
    selected_names: List[str] = []
    available: List[str] = []
    try:
        if os.path.isdir(default_dir):
            available = sorted([f for f in os.listdir(default_dir) if f.lower().endswith(".docx")])
            if available:
                selected_names = st.multiselect("Kullanılacak şablonları seçin", options=available, default=[])
            else:
                st.info("Templates klasöründe .docx şablon bulunamadı.")
        else:
            st.info("Templates klasörü bulunamadı. Lütfen 'templates' klasörü oluşturup .docx dosyalarınızı ekleyin.")
    except Exception as e:
        st.error(f"Templates klasörü okunamadı: {e}")

    template_items: List[Tuple[str, bytes]] = []
    for name in selected_names:
        try:
            full = os.path.join(default_dir, name)
            with open(full, "rb") as fh:
                template_items.append((name, fh.read()))
        except Exception as e:
            st.error(f"{name} okunamadı: {e}")

    summaries: List[Dict[str, object]] = []
    union_placeholders: Set[str] = set()
    if template_items:
        for name, data in template_items:
            try:
                placeholders, _ = extract_placeholders_from_docx_bytes(data)
                union_placeholders |= placeholders
                summaries.append({"name": name, "placeholders": sorted(list(placeholders))})
            except Exception as e:
                st.error(f"{name} okunamadı: {e}")



    st.markdown("---")

    st.subheader("🎤 Tek Seferde Kaydet → Çözümle → Öneri Üret")
    col_mic, col_btn = st.columns([3, 1])

    with col_mic:
        st.write("Konuşma kaydı")
        audio_bytes: Optional[bytes] = None
        if mic_recorder is not None:
            rec_val = mic_recorder(
                start_prompt="Kaydı Başlat",
                stop_prompt="Kaydı Durdur",
                just_once=False,
                use_container_width=True,
                key="single_shot_mic_standalone",
            )
            if isinstance(rec_val, dict) and rec_val.get("error"):
                st.error(f"Mikrofon hatası: {rec_val['error']}")
            audio_bytes = bytes_from_mic_return(rec_val)
        else:
            st.error("Mikrofon bileşeni mevcut değil (import başarısız).")

    with col_btn:
        if st.button("🧠 Transcribe + Anlamlandır", use_container_width=True):
            effective_key = (api_key_input or st.session_state.get("api_key", "")).strip()
            current_session_id = st.session_state.get("current_session_id")
            
            if not template_items:
                st.warning("Önce en az bir şablon seçin veya yükleyin.")
            elif not union_placeholders:
                st.warning("Şablonlarda placeholder bulunamadı.")
            elif not effective_key:
                st.warning("Lütfen OpenAI API anahtarınızı girin.")
            elif not audio_bytes:
                st.warning("Ses kaydı yapın.")
            elif not current_session_id:
                st.warning("Session bilgisi bulunamadı!")
            else:
                with st.spinner("Ses metne çevriliyor..."):
                    text = transcribe_audio_bytes(audio_bytes, effective_key, file_ext="wav")
                if text:
                    # Yeni transkripti mevcut transkriptle birleştir
                    existing_transcript = st.session_state.get("single_transcript", "")
                    new_text = text.strip()
                    
                    if existing_transcript:
                        # Mevcut transkript varsa, araya boşluk koyarak birleştir
                        merged_transcript = f"{existing_transcript} {new_text}"
                    else:
                        # İlk transkript ise direkt kullan
                        merged_transcript = new_text
                    
                    st.session_state["single_transcript"] = merged_transcript
                    with st.spinner("Placeholder değerleri çıkarılıyor..."):
                        ctx = aggregate_contexts_across_templates(template_items, union_placeholders)
                        suggested = infer_placeholder_values(
                            st.session_state["single_transcript"],
                            union_placeholders,
                            ctx,
                            effective_key,
                        )
                        
                        # Mevcut session verisini al
                        sm = get_local_session_manager()
                        existing_data = st.session_state.get("single_mapping", {})
                        
                        # Çakışmaları kontrol et
                        conflicts = detect_conflicts(existing_data, suggested)
                        
                        if conflicts:
                            st.warning(f"⚠️ {len(conflicts)} çakışma tespit edildi: {', '.join(conflicts)}")
                            st.info("Yeni değerler mevcut verilerle birleştirildi. Aşağıdan kontrol edip düzenleyebilirsiniz.")
                        
                        # Veriyi birleştir
                        merged_data = merge_extracted_data(existing_data, suggested)
                        st.session_state["single_mapping"] = merged_data
                        
                        # Session'a kaydet
                        if sm.update_session_data(current_session_id, suggested, merge=True):
                            filled_count = len([v for v in suggested.values() if v.strip()])
                            st.success(f"✅ Ses analizi tamamlandı! {filled_count} yeni bilgi eklendi ve session'a kaydedildi.")
                            
                            # Öğrenci bilgileri eklendiyse session ismini güncelle
                            updated_session = sm.get_session(current_session_id)
                            if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                                st.session_state["current_session_name"] = updated_session['session_name']
                                st.info("📝 Session ismi öğrenci bilgilerine göre güncellendi!")
                        else:
                            st.warning("Veriler çıkarıldı ama session'a kaydedilemedi.")
                        
                        st.rerun()  # Sayfayı yenile ki input kutuları güncellensin
                else:
                    st.error("Metne çeviri başarısız oldu.")

    if st.session_state.get("single_transcript"):
        # Transkript alanı ve temizleme butonu
        col_transcript, col_clear_transcript = st.columns([4, 1])
        
        with col_transcript:
            st.text_area(
                "📜 Birleşik Transkript",
                value=st.session_state.get("single_transcript", ""),
                height=160,
                disabled=True,
                help="Tüm ses kayıtlarınız burada birleşik olarak görünür. AI analizi bu metni kullanır."
            )
        
        with col_clear_transcript:
            st.write("")  # Boşluk için
            st.write("")  # Boşluk için
            if st.button("🗑️ Temizle", help="Transkripti temizler, placeholder değerleri korunur"):
                st.session_state["single_transcript"] = ""
                st.success("Transkript temizlendi!")
                st.rerun()

    if union_placeholders:
        st.subheader("✏️ Placeholder Değerleri - API Önerileri ve Manuel Düzenleme")
        

        
        # Bağlam bilgilerini al
        if template_items:
            ctx = aggregate_contexts_across_templates(template_items, union_placeholders)
        else:
            ctx = {}
            
        # Session verilerini şablonlara uygula butonu
        col_apply, col_clear = st.columns([2, 1])
        with col_apply:
            if st.button("🔄 Session Verilerini Tüm Şablonlara Uygula", help="Mevcut session'daki tüm verileri yeni seçilen şablonlara otomatik uygular"):
                current_session_id = st.session_state.get("current_session_id")
                if current_session_id:
                    sm = get_local_session_manager()
                    session_data = sm.get_session(current_session_id)
                    if session_data and session_data.get('extracted_data'):
                        # Tüm session verilerini current mapping'e aktar
                        session_extracted = session_data['extracted_data']
                        current_mapping = st.session_state.get("single_mapping", {})
                        
                        # Yeni placeholder'lar için session'daki verileri kullan
                        applied_count = 0
                        for ph in union_placeholders:
                            if ph in session_extracted and session_extracted[ph]:
                                if ph not in current_mapping or not current_mapping.get(ph):
                                    current_mapping[ph] = session_extracted[ph]
                                    applied_count += 1
                        
                        st.session_state["single_mapping"] = current_mapping
                        if applied_count > 0:
                            st.success(f"✅ {applied_count} placeholder session verilerinden dolduruldu!")
                            st.rerun()
                        else:
                            st.info("Uygulanacak yeni veri bulunamadı.")
                    else:
                        st.warning("Session'da kullanılabilir veri bulunamadı.")
        
        with col_clear:
            if st.button("🧹 Placeholder'ları Temizle", help="Tüm placeholder değerlerini temizler"):
                st.session_state["single_mapping"] = {}
                st.success("Placeholder'lar temizlendi!")
                st.rerun()
        
        st.markdown("---")
        
        edit_cols = st.columns(2)
        for idx, ph in enumerate(sorted(list(union_placeholders))):
            with edit_cols[idx % 2]:
                # Placeholder başlığı - süslü parantez olmadan göster
                display_name = ph.replace("{", "").replace("}", "")
                st.markdown(f"**{display_name}**")
                

                
                # Input kutusu
                cur_val = st.session_state.get("single_mapping", {}).get(ph, "")
                new_val = st.text_input(
                    f"Değer", 
                    value=cur_val, 
                    key=f"edit_standalone_{idx}_{ph}",
                    placeholder="API önerisi bekleyin veya manuel girin...",
                    label_visibility="collapsed"
                )
                
                # Manuel değişiklik varsa güncelle
                if new_val != cur_val:
                    st.session_state["single_mapping"][ph] = new_val
                    
                    # Session'a da kaydet
                    current_session_id = st.session_state.get("current_session_id")
                    if current_session_id:
                        sm = get_local_session_manager()
                        update_data = {ph: new_val}
                        sm.update_session_data(current_session_id, update_data, merge=True)
                        
                        # Öğrenci bilgileri değiştiyse session ismini güncelle
                        updated_session = sm.get_session(current_session_id)
                        if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                            st.session_state["current_session_name"] = updated_session['session_name']
                    
                    st.rerun()  # Bağlam önizlemelerini güncelle
                
                st.markdown("---")

    # Canlı Önizleme Bölümü
    if template_items:
        st.markdown("---")
        st.subheader("👁️ Canlı Önizleme - Şablon İçerikleri")
        st.caption("Placeholder'ları girdiğiniz değerlerle değiştirilerek gösteriliyor (boş olanlar olduğu gibi kalır)")
        
        # Session verilerini al ve single_mapping ile birleştir
        current_session_id = st.session_state.get("current_session_id")
        session_mapping = {}
        
        if current_session_id:
            sm = get_local_session_manager()
            session_data = sm.get_session(current_session_id)
            if session_data and session_data.get('extracted_data'):
                session_mapping = session_data['extracted_data']
        
        # Session verileri ile current mapping'i birleştir
        current_mapping = st.session_state.get("single_mapping", {})
        merged_mapping = {**session_mapping, **current_mapping}  # current_mapping öncelikli
        
        for name, data in template_items:
            with st.expander(f"📄 {name} - Önizleme", expanded=False):
                # Bu şablona özel placeholder'ları al
                template_placeholders, _ = extract_placeholders_from_docx_bytes(data)
                
                preview_text = generate_preview_text(data, merged_mapping)
                
                # Sadece bu şablondaki doldurulmamış placeholder'ları göster
                unfilled = []
                for ph in template_placeholders:
                    if not merged_mapping.get(ph, "").strip():
                        unfilled.append(ph.replace("{", "").replace("}", ""))
                
                if unfilled:
                    st.info(f"Bu şablonda henüz doldurulmamış: {', '.join(unfilled)}")
                else:
                    st.success("Bu şablondaki tüm placeholder'lar dolduruldu! ✅")
                
                # Tüm içeriği doğrudan göster - kaydırma yok
                st.text(preview_text)

    # Session bilgilerini sidebar'da göster
    with st.sidebar:
        st.subheader(f"📊 Session Bilgileri")
        current_session_id = st.session_state.get("current_session_id")
        if current_session_id:
            sm = get_local_session_manager()
            session_data = sm.get_session(current_session_id)
            if session_data:
                st.write(f"**Ad:** {session_data['session_name']}")
                st.write(f"**ID:** {current_session_id[:12]}...")
                st.write(f"**Dolu Alanlar:** {len([v for v in session_data['extracted_data'].values() if v])}")
                
                if st.button("🔄 Session'ı Yenile"):
                    # Session verilerini yeniden yükle
                    fresh_data = sm.get_session(current_session_id)
                    if fresh_data:
                        st.session_state["single_mapping"] = fresh_data['extracted_data'].copy()
                        st.success("Session verileri yenilendi!")
                        st.rerun()

    st.markdown("---")

    a1, a2, a3 = st.columns([1, 1, 1])
    with a1:
        if st.button("📄 Tüm Şablonları Doldur", use_container_width=True):
            if not template_items:
                st.warning("Önce şablon seçin veya yükleyin.")
            elif not st.session_state.get("single_mapping"):
                st.warning("Kullanılacak değer bulunamadı. Ses girişi yapın.")
            else:
                try:
                    results: List[Dict[str, object]] = []
                    current_session_name = st.session_state.get("current_session_name", "Session")
                    
                    for idx, (name, data) in enumerate(template_items):
                        doc = Document(io.BytesIO(data))
                        mapping = {k: v for k, v in st.session_state["single_mapping"].items() if str(v).strip()}
                        replaced = replace_placeholders_in_document(doc, mapping)
                        buf = io.BytesIO()
                        doc.save(buf)
                        out_bytes = buf.getvalue()
                        # Session adını dosya adına ekle
                        safe_session_name = re.sub(r'[^\w\s-]', '', current_session_name).strip()[:20]
                        out_name = f"{safe_session_name}_{os.path.splitext(name)[0]}.docx"
                        results.append({
                            "name": out_name,
                            "replaced": replaced,
                            "data": out_bytes,
                            "key": f"dl_session_{idx}_{out_name}",
                            "values": mapping,
                        })
                    st.session_state["single_results"] = results
                    if results:
                        st.success("✅ Doldurma tamamlandı! İndirilebilir dosyalar hazır.")
                except Exception as e:
                    st.error(f"İşlem başarısız: {e}")

    with a2:
        if st.button("🧹 Bu Oturumu Temizle", use_container_width=True):
            st.session_state["single_results"] = []
            st.session_state["single_transcript"] = ""
            st.success("Oturum verileri temizlendi!")
    
    with a3:
        if st.button("🗑️ Session Verilerini Sıfırla", use_container_width=True, type="secondary"):
            current_session_id = st.session_state.get("current_session_id")
            if current_session_id:
                # Session'daki verileri sıfırla
                sm = get_local_session_manager()
                if sm.update_session_data(current_session_id, {}, merge=False):
                    st.session_state["single_mapping"] = {}
                    st.session_state["single_results"] = []
                    st.session_state["single_transcript"] = ""
                    st.success("Session verileri sıfırlandı!")
                    st.rerun()
                else:
                    st.error("Session sıfırlanamadı!")

    if st.session_state.get("single_results"):
        st.markdown("---")
        st.subheader("📥 İndirilecek Dosyalar")
        for r in st.session_state["single_results"]:
            st.write(f"**{r['name']}** → {r['replaced']} yer değiştirme")
            
            # Eksik değerleri kontrol et
            try:
                # Orijinal şablon dosyasını analiz et (hangi template'ten geldiğini bul)
                template_name = r['name'].split('_', 1)[-1].replace('.docx', '.docx')
                original_template = None
                
                for name, data in template_items:
                    if name in r['name'] or template_name == name:
                        original_template = data
                        break
                
                if original_template:
                    # Şablondaki tüm placeholder'ları bul
                    placeholders, _ = extract_placeholders_from_docx_bytes(original_template)
                    used_values = r.get("values", {})
                    
                    # Eksik placeholder'ları tespit et
                    missing_placeholders = []
                    for ph in placeholders:
                        if ph not in used_values or not used_values.get(ph, "").strip():
                            missing_placeholders.append(ph)
                    
                    # Durum göstergesi
                    if missing_placeholders:
                        st.warning(f"⚠️ **{len(missing_placeholders)} eksik değer:** {', '.join(missing_placeholders)}")
                    else:
                        st.success("✅ **Tüm değerler dolduruldu!**")
                        
                else:
                    st.info("Şablon analizi yapılamadı.")
                    
            except Exception as e:
                st.info("Eksik değer analizi yapılamadı.")
            
            # Kullanılan değerler (daha kompakt)
            with st.expander(f"🔧 Kullanılan Değerler ({r['name']})"):
                if r["values"]:
                    cols = st.columns(2)
                    items = list(r["values"].items())
                    for i, (k, v) in enumerate(items):
                        with cols[i % 2]:
                            st.write(f"**{k}:** {v}")
                else:
                    st.write("Hiç değer kullanılmadı.")
            
            # İndirme butonu
            st.download_button(
                label=f"📥 İndir: {r['name']}",
                data=r["data"],
                file_name=r["name"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                key=r["key"],
                use_container_width=True
            )
            st.markdown("---")


if __name__ == "__main__":
    main()



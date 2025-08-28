# unified_app.py - Sesli Belge Doldurma Sistemi (Birleşik Versiyon)
# Voice.py ve app2.py projelerinin en iyi özelliklerini birleştiren gelişmiş sistem

import io
import os
import re
import json
import sys
import tempfile
import uuid
import traceback
from typing import Dict, List, Optional, Set, Tuple
from datetime import datetime, date, time
from zoneinfo import ZoneInfo

import streamlit as st
import importlib
from docx import Document
import dateparser

# Local session management import
from local_session_manager import get_local_session_manager, merge_extracted_data, detect_conflicts

# OpenAI import
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Mikrofon kütüphaneleri - birincil ve alternatifler
MIC_IMPORT_ERROR: Optional[str] = None
mic_recorder = None  # from streamlit_mic_recorder
audio_recorder_fn = None  # from audio_recorder_streamlit
st_audiorec_fn = None  # from st_audiorec

try:
    _mic_module = importlib.import_module("streamlit_mic_recorder")
    mic_recorder = getattr(_mic_module, "mic_recorder", None)
except Exception as e:
    MIC_IMPORT_ERROR = f"streamlit-mic-recorder yüklenemedi: {e}"

try:
    _ar_module = importlib.import_module("audio_recorder_streamlit")
    audio_recorder_fn = getattr(_ar_module, "audio_recorder", None)
except Exception:
    pass

try:
    _sar_module = importlib.import_module("st_audiorec")
    st_audiorec_fn = getattr(_sar_module, "st_audiorec", None)
except Exception:
    pass

def render_audio_recorder_ui() -> Optional[bytes]:
    """Tarayıcıdan ses kaydı al ve bytes döndür (mevcut bileşene göre)."""
    # Öncelik: streamlit-mic-recorder
    if mic_recorder is not None:
        st.write("**Mikrofon ile Kayıt**")
        rec_val = mic_recorder(
            start_prompt="🎙️ Kaydı Başlat",
            stop_prompt="⏹️ Kaydı Durdur",
            just_once=False,
            use_container_width=True,
            key="unified_mic_recorder",
        )
        if isinstance(rec_val, dict) and rec_val.get("error"):
            st.error(f"Mikrofon hatası: {rec_val['error']}")
            return None
        return bytes_from_mic_return(rec_val)

    # Alternatif 1: audio-recorder-streamlit
    if audio_recorder_fn is not None:
        st.write("**Mikrofon ile Kayıt (alternatif)**")
        rec_val = audio_recorder_fn()
        return bytes_from_mic_return(rec_val) if rec_val else None

    # Alternatif 2: streamlit-audiorec
    if st_audiorec_fn is not None:
        st.write("**Mikrofon ile Kayıt (alternatif)**")
        rec_val = st_audiorec_fn()
        return bytes_from_mic_return(rec_val) if rec_val else None

    st.error("Mikrofon kütüphanesi mevcut değil.")
    if MIC_IMPORT_ERROR:
        st.error(MIC_IMPORT_ERROR)
    st.info("Lütfen 'streamlit-mic-recorder' veya 'audio-recorder-streamlit' paketini kurun.")
    return None

# Zaman dilimi
IST = ZoneInfo("Europe/Istanbul")
TR_DAYS = {0:"Pazartesi", 1:"Salı", 2:"Çarşamba", 3:"Perşembe", 4:"Cuma", 5:"Cumartesi", 6:"Pazar"}

# ================== Yardımcı Fonksiyonlar ==================

def ensure_utf8_encoding():
    """Sistem encoding'ini kontrol et ve gerekirse UTF-8'e ayarla"""
    import locale
    try:
        # Sistem locale'ini kontrol et
        current_encoding = locale.getpreferredencoding()
        if 'utf-8' not in current_encoding.lower() and 'cp65001' not in current_encoding.lower():
            # Windows'ta UTF-8 desteği için
            import os
            os.environ['PYTHONIOENCODING'] = 'utf-8'
    except Exception:
        pass  # Encoding ayarlanamadıysa sessizce devam et

def safe_str(obj) -> str:
    """Herhangi bir objeyi güvenli şekilde string'e çevir"""
    try:
        return str(obj)
    except UnicodeError:
        try:
            return str(obj).encode('utf-8', errors='replace').decode('utf-8')
        except:
            return repr(obj)

def bytes_from_mic_return(value) -> Optional[bytes]:
    """Mikrofon dönüş değerini normalize et"""
    if value is None:
        return None
    if isinstance(value, dict) and "bytes" in value:
        return value["bytes"]
    if isinstance(value, (bytes, bytearray)):
        return bytes(value)
    return None

def transcribe_audio_bytes(audio_bytes: bytes, api_key: str, lang: str = "tr") -> Optional[str]:
    """Ses dosyasını metne çevir"""
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil. 'openai' paketini kurun.")
        return None

    tmp_path = None
    try:
        # API key'i güvenli şekilde işle
        safe_api_key = api_key.strip() if api_key else ""
        if not safe_api_key:
            st.error("API key boş veya geçersiz")
            return None
            
        client = OpenAI(api_key=safe_api_key)
        
        # Güvenli temp dosya oluştur - Unicode güvenli
        safe_filename = f"audio_{uuid.uuid4().hex}.wav"
        tmp_dir = tempfile.gettempdir()
        tmp_path = os.path.join(tmp_dir, safe_filename)
        
        # Dosya yolunun Unicode karakterler içerip içermediğini kontrol et
        try:
            tmp_path.encode('ascii')
        except UnicodeEncodeError:
            # ASCII olmayan karakterler varsa, farklı bir yol kullan
            tmp_path = os.path.join(tempfile.gettempdir(), f"temp_audio_{uuid.uuid4().hex[:8]}.wav")
        
        with open(tmp_path, "wb") as f:
            f.write(audio_bytes)

        with open(tmp_path, "rb") as f:
            resp = client.audio.transcriptions.create(
                model="whisper-1", 
                file=f,
                language=lang,
                response_format="text"
            )

        # Response'u güvenli şekilde işle
        if isinstance(resp, str):
            # Türkçe karakterleri güvenli şekilde handle et
            try:
                return resp
            except UnicodeError:
                return resp.encode('utf-8', errors='replace').decode('utf-8')
        else:
            text_result = getattr(resp, "text", None) or (resp.get("text") if isinstance(resp, dict) else None)
            if text_result:
                try:
                    return str(text_result)
                except UnicodeError:
                    return str(text_result).encode('utf-8', errors='replace').decode('utf-8')
            return None
        
    except Exception as e:
        # Unicode karakterleri tamamen güvenli şekilde işle
        try:
            # Exception mesajını güvenli şekilde al
            error_msg = str(e)
            # ASCII dışı karakterleri kontrol et ve temizle
            error_msg.encode('ascii')
        except (UnicodeError, UnicodeEncodeError, UnicodeDecodeError):
            # Unicode sorunu varsa, güvenli bir mesaj kullan
            try:
                error_msg = repr(str(e))  # repr() kullanarak güvenli gösterim
            except:
                error_msg = "Ses işleme sırasında Unicode karakter hatası oluştu"
        
        # Streamlit error mesajını da güvenli şekilde göster
        safe_error_msg = safe_str(error_msg)
        st.error(f"Ses metne çevrilemedi: {safe_error_msg}")
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass

def extract_placeholders_from_docx_bytes(file_bytes: bytes) -> Tuple[Set[str], str]:
    """Word dosyasından placeholder'ları çıkar"""
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    
    # Tüm metinleri topla
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + " "
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text += p.text + " "
    
    # Header/Footer
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    text += p.text + " "
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    text += p.text + " "
    
    placeholders = set(re.findall(r"\{[^}]+\}", text))
    return placeholders, text

def replace_placeholders_in_document(doc: Document, placeholder_values: Dict[str, str]) -> int:
    """Word belgesindeki placeholder'ları değiştir"""
    def replace_in_paragraph(paragraph):
        if not paragraph.runs:
            return 0
        original_text = "".join(run.text for run in paragraph.runs)
        replaced_text = original_text
        total_replacements = 0
        
        for placeholder, value in placeholder_values.items():
            if value is None:
                continue
            count = replaced_text.count(placeholder)
            if count:
                replaced_text = replaced_text.replace(placeholder, str(value))
                total_replacements += count
        
        if replaced_text != original_text:
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            paragraph.add_run(replaced_text)
        
        return total_replacements

    replacements_made = 0
    
    # Body paragraphs
    for p in doc.paragraphs:
        replacements_made += replace_in_paragraph(p)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replacements_made += replace_in_paragraph(p)
    
    # Headers/Footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replacements_made += replace_in_paragraph(p)
        if section.footer:
            for p in section.footer.paragraphs:
                replacements_made += replace_in_paragraph(p)
    
    return replacements_made

def parse_tr_date(text: str) -> Optional[datetime]:
    """Türkçe tarih formatlarını çözümle"""
    if not text:
        return None
    return dateparser.parse(text, languages=["tr"])

def split_date(dt: datetime) -> Dict[str, str]:
    """Datetime'ı gün/ay/yıl olarak böl"""
    return {
        "gun": f"{dt.day:02d}",
        "ay": f"{dt.month:02d}",
        "yil": f"{dt.year}"
    }

def today_isbu(dt: datetime = None) -> Dict[str, str]:
    """İşbu tarih/saat bilgilerini al"""
    now = dt or datetime.now(IST)
    return {
        "isbu_gun": f"{now.day:02d}",
        "isbu_ay": f"{now.month:02d}",
        "isbu_yil": f"{now.year}",
        "isbu_saat": now.strftime("%H:%M")
    }

# ================== AI Analiz Fonksiyonları ==================

def extract_placeholder_contexts_from_docx_bytes(file_bytes: bytes, placeholders: Set[str], window: int = 70) -> Dict[str, List[str]]:
    """Placeholder'ların bağlamlarını çıkar"""
    doc = Document(io.BytesIO(file_bytes))
    blocks = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        blocks.append(p.text)
    
    contexts = {ph: [] for ph in placeholders}
    
    for block in blocks:
        for ph in placeholders:
            pattern = re.escape(ph)
            for m in re.finditer(pattern, block):
                start, end = m.start(), m.end()
                before = block[max(0, start - window): start]
                after = block[end: end + window]
                snippet = f"{before}{ph}{after}"
                if len(contexts[ph]) < 3:
                    contexts[ph].append(snippet)
    
    return contexts

def aggregate_contexts_across_templates(templates: List[Tuple[str, bytes]], placeholders: Set[str]) -> Dict[str, List[str]]:
    """Tüm şablonlardan bağlamları topla"""
    combined = {ph: [] for ph in placeholders}
    
    for _, data in templates:
        try:
            local_ctx = extract_placeholder_contexts_from_docx_bytes(data, placeholders)
            for ph, lst in local_ctx.items():
                for s in lst:
                    if len(combined[ph]) < 5:
                        combined[ph].append(s)
        except Exception:
            continue
    
    return combined

def parse_json_loose(s: str) -> Dict[str, str]:
    """JSON'u esnek şekilde parse et"""
    try:
        return json.loads(s)
    except Exception:
        pass
    
    try:
        m = re.search(r"\{[\s\S]*\}", s)
        if m:
            return json.loads(m.group(0))
    except Exception:
        pass
    
    return {}

# Regex tabanlı fallback kaldırıldı (tüm çıkarım LLM ile yapılacak)

def infer_placeholder_values(
    transcript: str,
    placeholders: Set[str],
    contexts: Dict[str, List[str]],
    api_key: str,
    model: str = "gpt-4o-mini",
) -> Dict[str, str]:
    """AI ile placeholder değerlerini çıkar"""
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil.")
        return {}
    
    client = OpenAI(api_key=api_key)
    ph_list = sorted(list(placeholders))
    
    # Mevcut (kullanıcı tarafından girilmiş) değerleri al ve prompt'a ekle (değiştirme)
    existing_values = {}
    try:
        existing_values = {
            k: v for k, v in (st.session_state.get("current_mapping", {}) or {}).items()
            if k in placeholders and str(v).strip()
        }
    except Exception:
        existing_values = {}

    # Gelişmiş prompt
    # İsim ve bölüm alanları için kuralları dinamik üret
    ph_lower_list = [ph.lower() for ph in placeholders]
    has_fullname_key = any(("ogrenci" in p and ("adi_soyadi" in p or "ad_soyad" in p)) for p in ph_lower_list)
    has_name_key = any(("ogrenci" in p and ("ad" in p or "adi" in p or "isim" in p) and "soyad" not in p) for p in ph_lower_list)
    has_surname_key = any(("ogrenci" in p and ("soyad" in p or "soyadi" in p)) for p in ph_lower_list)
    has_department_key = any(("bolum" in p) or ("bölüm" in p) for p in ph_lower_list)

    name_rules_lines: List[str] = []
    if has_fullname_key:
        name_rules_lines.append("- {ogrenci_adi_soyadi} alanı için öğrencinin tam adını 'Ad Soyad' formatında ver (örn: 'Ecem Nalbantoğlu').")
    if has_name_key and has_surname_key:
        name_rules_lines.append("- Ayrı alanlar varsa {ogrenci_ad}/{ogrenci_adi} ve {ogrenci_soyad}/{ogrenci_soyadi} alanlarını ayrı ayrı doldur (örn: 'Ecem' ve 'Nalbantoğlu').")
    elif has_name_key or has_surname_key:
        name_rules_lines.append("- Öğrenci adı/soyadı alanları varsa transkriptte geçtiği şekliyle doldur.")

    department_rules_lines: List[str] = []
    if has_department_key:
        department_rules_lines.append("- {bolum}/{bolum_adi} gibi bölüm alanlarında tek bir bölüm adı ver. Birden fazla bölüm yazma (örn: 'İşletme' veya 'Bilgisayar Mühendisliği'; 'İşletme Muhasebe' yazma).")

    prompt_text = f"""
SES TRANSKRİPTİ:
"{transcript}"

TEMPLATE PLACEHOLDER'LARI VE BAĞLAMLARI:
"""
    
    for ph in ph_list:
        if ph in contexts and contexts[ph]:
            context_examples = "\n".join([f"  • {ctx[:200]}" for ctx in contexts[ph][:3]])
            prompt_text += f"\n{ph}:\n{context_examples}\n"
        else:
            prompt_text += f"\n{ph}: (Bağlam bulunamadı)\n"
    
    prompt_text += """

GÖREV:
1. Ses transkriptini analiz et
2. Her placeholder için template bağlamını incele
3. Bağlama uygun değerleri ses transkriptinden çıkar
4. Çıkaramadığın bilgiler için boş string ("") bırak
5. SADECE JSON formatında cevap ver

MEVCUT DEĞERLER (DEĞİŞTİRME):
""" + json.dumps(existing_values, ensure_ascii=False, indent=2) + """

ÇIKTI KURALLARI:
- JSON anahtarları, placeholder stringleriyle birebir aynı olmalı (örnek: {ogrenci_no})
- Mevcut dolu alanları DEĞİŞTİRME; sadece boş olanları doldur
- Tarih ve saat alanları bağlama uygun normalize edilmeli (tarih: YYYY-MM-DD veya {gun,ay,yil}; saat: HH:MM)
- Sayısal alanlar (no, tc vb.) sadece rakam içersin
- İsim alanlarında gereksiz ekleri çıkar; açıklama alanlarında öğrenci ismi geçmesin
 - {blok} alanı sadece tek büyük harf (A-Z) olmalı (ör: A, B, C). Tahmin etme; transkriptte yoksa boş bırak.
 - {ogrenci_adi_soyadi} alanına öğrencinin tam adı ve soyadı gelmeli (örn: "Emre Yılmaz").

EK ÖZEL KURALLAR:
"""
    # Dinamik ek kuralları prompt'a ekle
    if name_rules_lines or department_rules_lines:
        extra_rules = "\n".join(name_rules_lines + department_rules_lines)
        prompt_text += extra_rules + "\n"
    prompt_text += """

ÖZEL İSTEK:
- Açıklama alanlarında sadece olayın kendisini yaz
- Öğrencinin adı ve soyadını açıklama alanlarına ekleme
- Sadece ne olduğunu objektif şekilde açıkla
 - Verilmeyen bilgileri uydurma; emin değilsen boş string ver

JSON formatı örneği:
""" + "{" + ", ".join([f'"{ph}": "değer_veya_boş_string"' for ph in ph_list[:3]]) + "...}"

    messages = [
        {"role": "system", "content": "Uzman bir bilgi çıkarım asistanısın. Kullanıcı transkriptini ve template bağlamlarını analiz ederek, placeholder anahtarlarıyla birebir eşleşen JSON üretirsin. Mevcut dolu değerleri asla değiştirme; sadece eksik (boş) alanları doldur. Tarih/saat ve sayısal alanları normalize et. Açıklama alanında öğrenci ismi geçmesin. Sadece transkriptte açıkça geçen bilgileri kullan; emin olmadığın durumda boş string ver. Sadece JSON döndür."},
        {"role": "user", "content": prompt_text},
    ]
    
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.0,
        )
        content = resp.choices[0].message.content if resp and resp.choices else "{}"
        data = parse_json_loose(content or "{}")
        
        result = {}
        for ph in placeholders:
            key_lower = ph.lower()
            raw_val = str(data.get(ph, "")).strip() if isinstance(data, dict) else ""
            val = raw_val
            # Özel kural: {blok} sadece tek harf (A-Z)
            if "blok" in key_lower:
                import re as _re
                only_letters = "".join(ch for ch in val if ch.isalpha())
                val = only_letters[:1].upper() if only_letters else ""
            # Özel kural: bölüm alanlarında tek bölüm adı döndür
            if ("bolum" in key_lower) or ("bölüm" in key_lower):
                lowered = val.lower()
                # Önce ayırıcılarla kes
                for sep in [",", "/", "&", "|", ";"]:
                    if sep in val:
                        val = val.split(sep)[0]
                # Bağlaçlara göre kes (ve/veya)
                for conj in [" ve ", " veya "]:
                    if conj in lowered:
                        idx = lowered.index(conj)
                        val = val[:idx]
                        break
                val = str(val).strip()
            # Genel: güvenli string
            if val:
                try:
                    result[ph] = str(val).strip()
                except UnicodeEncodeError:
                    result[ph] = val.encode('utf-8', errors='replace').decode('utf-8')
            else:
                result[ph] = ""

        # İsim alanları için ek post-processing: fullname <-> ad/soyad senkronizasyonu
        try:
            # Anahtarların küçük harf normalize edilmiş haritasını oluştur
            keys_by_lower = {k.lower(): k for k in result.keys()}
            # Varyantları bul
            fullname_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("adi_soyadi" in lk or "ad_soyad" in lk))), None)
            name_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("ad" in lk or "adi" in lk or "isim" in lk) and "soyad" not in lk)), None)
            surname_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("soyad" in lk or "soyadi" in lk))), None)

            # Eğer fullname boş ama ad ve soyad doluysa, birleştir
            if fullname_key and (not result.get(fullname_key)) and name_key and surname_key and result.get(name_key) and result.get(surname_key):
                combined = f"{str(result.get(name_key)).strip()} {str(result.get(surname_key)).strip()}".strip()
                result[fullname_key] = combined
            # Eğer ad/soyad boş ama fullname doluysa, basit böl
            if fullname_key and result.get(fullname_key) and ((name_key and not result.get(name_key)) or (surname_key and not result.get(surname_key))):
                fullname_val = str(result.get(fullname_key)).strip()
                parts = [p for p in fullname_val.split() if p]
                if len(parts) >= 2:
                    first = " ".join(parts[:-1])
                    last = parts[-1]
                    if name_key and not result.get(name_key):
                        result[name_key] = first
                    if surname_key and not result.get(surname_key):
                        result[surname_key] = last
        except Exception:
            pass
        
        return result
    except Exception as e:
        # Unicode güvenli hata mesajı
        try:
            error_msg = str(e)
        except UnicodeEncodeError:
            error_msg = "AI analizi sırasında karakter kodlama hatası"
        st.error(f"Ses analizi başarısız: {error_msg}")
        return {ph: "" for ph in placeholders}

# ================== Öğrenci Yönetimi ==================

def extract_student_info(session_data):
    """Session'dan öğrenci bilgilerini çıkar"""
    if not session_data or not session_data.get('extracted_data'):
        return None, None
    
    extracted = session_data['extracted_data']
    student_no = None
    student_name = None
    
    for key, value in extracted.items():
        if value and str(value).strip():
            key_lower = key.lower().replace('{', '').replace('}', '')
            value_str = str(value).strip()
            
            # Öğrenci dışı kişi alanlarını atla
            skip_person_keywords = [
                'gozetmen', 'gözetmen', 'ogretim', 'öğretim', 'elemani', 'elemanı',
                'gorevli', 'görevli', 'bolum_baskanligi', 'bölüm başkanlığı', 'baskan', 'başkan',
                'danisman', 'danışman', 'sifre', 'şifre', 'yetkili', 'imza'
            ]
            if any(k in key_lower for k in skip_person_keywords):
                continue

            # Öğrenci numarası
            if 'ogrencino' in key_lower or 'ogrenci_no' in key_lower:
                student_no = value_str
            elif 'no' in key_lower and not student_no:
                student_no = value_str
            
            # Öğrenci adı
            elif ('ogrenci' in key_lower) and ('ad' in key_lower or 'adi' in key_lower or 'isim' in key_lower) and 'soyad' not in key_lower:
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{value_str} {student_name}"
                    else:
                        student_name = value_str
            elif ('ogrenci' in key_lower) and ('soyad' in key_lower or 'soyadi' in key_lower):
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{student_name} {value_str}"
                    else:
                        student_name = value_str
            elif 'ogrenci' in key_lower and any(keyword in key_lower for keyword in ['adi_soyadi', 'ad_soyad']):
                if not value_str.isdigit():
                    student_name = value_str
    
    return student_no, student_name

def update_session_name_if_needed(session_id, session_data):
    """Öğrenci bilgileri varsa session ismini standart formata güncelle."""
    try:
        def _format_standard(no: str, name: str) -> str:
            safe_no = (no or "").strip()
            safe_name = " ".join((name or "").split())
            return f"{safe_no} - {safe_name}" if safe_no and safe_name else ""

        student_no, student_name = extract_student_info(session_data)
        current_name = session_data.get('session_name', '')

        # Sadece her ikisi de varsa standart isim uygula
        if student_no and student_name:
            new_name = _format_standard(student_no, student_name)
            if new_name and new_name != current_name:
                sm = get_local_session_manager()
                session_data['session_name'] = new_name
                return sm.save_session(session_id, session_data)
        return False
    except Exception as e:
        st.error(f"Session ismi güncellenirken hata: {e}")
        return False

# ================== Ana Uygulama ==================

def main():
    # UTF-8 encoding'i kontrol et ve ayarla
    ensure_utf8_encoding()
    
    st.set_page_config(
        page_title="🎯 Sesli Belge Doldurma Sistemi", 
        page_icon="🎯", 
        layout="wide"
    )

    # Session state initialization
    if "page" not in st.session_state:
        st.session_state["page"] = "session_manager"
    if "current_session_id" not in st.session_state:
        st.session_state["current_session_id"] = None
    if "current_session_name" not in st.session_state:
        st.session_state["current_session_name"] = ""
    if "api_key" not in st.session_state:
        st.session_state["api_key"] = ""
    # Form selection related state
    if "selected_form_group" not in st.session_state:
        st.session_state["selected_form_group"] = None  # Örn: "Ek 1-2-3", "Ek 4", "Ek 6", "Ek 8"
    if "form_group_applied" not in st.session_state:
        st.session_state["form_group_applied"] = None
    if "templates_initialized_for" not in st.session_state:
        st.session_state["templates_initialized_for"] = None
    if "selected_templates" not in st.session_state:
        st.session_state["selected_templates"] = []

    # Page routing
    if st.session_state["page"] == "session_manager":
        show_session_manager()
    elif st.session_state["page"] == "form_selector":
        show_form_selector()
    elif st.session_state["page"] == "voice_app":
        show_voice_app()
    else:
        st.session_state["page"] = "session_manager"
        st.rerun()

def show_session_manager():
    """Session yönetim arayüzü"""
    st.title("🎯 Sesli Belge Doldurma Sistemi")
    st.caption("Ses girdi ile Word şablonlarını otomatik dolduran akıllı sistem")
    
    sm = get_local_session_manager()
    
    # Arama çubuğu (yalnızca öğrenci adı veya numarasına göre)
    search_term = st.text_input("🔍 Öğrenci Ara", placeholder="Öğrenci adı veya öğrenci numarası...")
    
    # Session listesi
    sessions = sm.get_all_sessions()
    
    # Arama filtresi (yalnızca ad veya numara)
    if search_term:
        q = search_term.lower().strip()
        filtered_sessions = []
        for session in sessions:
            student_no, student_name = extract_student_info(session)
            match_no = bool(student_no and q in student_no.lower())
            match_name = bool(student_name and q in student_name.lower())
            if match_no or match_name:
                filtered_sessions.append(session)
        sessions = filtered_sessions
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.subheader("📁 Öğrenci Session'ları")
        
        if not sessions:
            if search_term:
                st.info("Arama kriterinize uygun session bulunamadı.")
            else:
                st.info("Henüz session oluşturulmamış. 'Yeni Session' butonuna tıklayın.")
        else:
            for session in sessions:
                student_no, student_name = extract_student_info(session)
                
                if student_no and student_name:
                    display_title = f"👤 {student_no} - {student_name}"
                else:
                    # Standart dışı isim varsa da aynı formatla göster
                    fallback_name = session.get('session_name', '')
                    display_title = f"👤 {fallback_name}"
                
                with st.expander(display_title, expanded=False):
                    col_info, col_actions = st.columns([2, 1])
                    
                    with col_info:
                        st.write(f"**Oluşturma:** {session['created_date'][:10]}")
                        if student_no:
                            st.write(f"**Öğrenci No:** {student_no}")
                        if student_name:
                            st.write(f"**Öğrenci Adı:** {student_name}")
                        
                        data_count = len([v for v in session.get('extracted_data', {}).values() if v])
                        st.write(f"**Dolu Alanlar:** {data_count}")
                    
                    with col_actions:
                        if st.button(f"🚀 Aç", key=f"open_{session['session_id']}"):
                            st.session_state["current_session_id"] = session['session_id']
                            st.session_state["current_session_name"] = session['session_name']
                            # Form seçim sayfasına yönlendir ve önceki seçimleri sıfırla
                            st.session_state["selected_form_group"] = None
                            st.session_state["form_group_applied"] = None
                            st.session_state["templates_initialized_for"] = None
                            st.session_state["selected_templates"] = []
                            st.session_state["page"] = "form_selector"
                            st.rerun()
                        
                        if st.button(f"🗑️ Sil", key=f"delete_{session['session_id']}"):
                            st.session_state[f"confirm_delete_{session['session_id']}"] = True
                            st.rerun()
                        
                        if st.session_state.get(f"confirm_delete_{session['session_id']}", False):
                            st.warning("⚠️ Silmek istediğinizden emin misiniz?")
                            col_yes, col_no = st.columns(2)
                            
                            with col_yes:
                                if st.button("✅ Evet", key=f"confirm_yes_{session['session_id']}"):
                                    if sm.delete_session(session['session_id']):
                                        st.success("Session silindi!")
                                        del st.session_state[f"confirm_delete_{session['session_id']}"]
                                        st.rerun()
                            
                            with col_no:
                                if st.button("❌ İptal", key=f"confirm_no_{session['session_id']}"):
                                    del st.session_state[f"confirm_delete_{session['session_id']}"]
                                    st.rerun()
    
    with col2:
        st.subheader("🚀 Yeni Session")
        st.write("Yeni bir öğrenci için session başlatın.")
        
        if st.button("📝 Yeni Session Başlat", type="primary", use_container_width=True):
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            session_name = f"Yeni Session - {timestamp}"
            
            session_id = sm.create_session(session_name)
            if session_id:
                st.success("Yeni session başlatıldı!")
                st.session_state["current_session_id"] = session_id
                st.session_state["current_session_name"] = session_name
                # Yeni session sonrası form seçim ekranına git
                st.session_state["page"] = "form_selector"
                # Yeni session'da transkript ve mapping boşlansın
                st.session_state["current_transcript"] = ""
                st.session_state["transcript_loaded_for"] = session_id
                st.session_state["current_mapping"] = {}
                st.session_state["mapping_loaded_for"] = session_id
                st.session_state["results"] = None
                # Form seçim state'leri
                st.session_state["selected_form_group"] = None
                st.session_state["form_group_applied"] = None
                st.session_state["templates_initialized_for"] = None
                st.session_state["selected_templates"] = []
                st.rerun()
            else:
                st.error("Session oluşturulamadı!")
        
        st.info("💡 **İpucu:** Session başlattıktan sonra öğrenci bilgilerini sesli girdi ile kaydedin.")

def show_form_selector():
    """Form (Ek) seçim ekranı"""
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return

    st.title("🧩 Hangi Ek doldurulacak?")
    st.caption(f"{current_session_name}")
    st.markdown("Seçiminiz bu session için şablonları otomatik işaretler. İstediğiniz zaman değiştirebilirsiniz.")

    options = ["Ek 1-2-3", "Ek 4", "Ek 6", "Ek 8", "Ek 9", "Ek 11"]
    default_idx = options.index(st.session_state.get("selected_form_group")) if st.session_state.get("selected_form_group") in options else 0
    selected = st.radio("Form seti", options=options, index=default_idx, horizontal=True)

    col_go, col_back = st.columns([1, 1])
    with col_go:
        if st.button("Devam et ➜", type="primary", use_container_width=True):
            st.session_state["selected_form_group"] = selected
            st.session_state["form_group_applied"] = None  # Voice sayfasında yeniden uygula
            st.session_state["templates_initialized_for"] = None
            st.session_state["selected_templates"] = []
            st.session_state["page"] = "voice_app"
            st.rerun()
    with col_back:
        if st.button("↩️ Session listesine dön", use_container_width=True):
            st.session_state["page"] = "session_manager"
            st.rerun()

def show_voice_app():
    """Ana ses uygulama arayüzü"""
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return
    
    # Session verilerini yükle
    sm = get_local_session_manager()
    session_data = sm.get_session(current_session_id)
    
    if not session_data:
        st.error("Session verisi yüklenemedi!")
        return
    
    # Session state'leri initialize et (mapping ve transcript session bazlı yüklensin)
    if st.session_state.get("mapping_loaded_for") != current_session_id:
        st.session_state["current_mapping"] = {}
        st.session_state["mapping_loaded_for"] = current_session_id
        st.session_state["results"] = None
    # Transkript, session bazlı yüklensin (diğer session'dan taşınmasın)
    if st.session_state.get("transcript_loaded_for") != current_session_id:
        st.session_state["current_transcript"] = session_data.get('transcript', "")
        st.session_state["transcript_loaded_for"] = current_session_id
    
    # Header
    col_title, col_actions = st.columns([4, 2])
    with col_title:
        st.title(f"🎯 {current_session_name}")
        st.caption(f"Session ID: {current_session_id[:12]}...")
        active_group = st.session_state.get("selected_form_group") or "Seçilmedi"
        st.info(f"Aktif Form Seti: {active_group}")
    
    with col_actions:
        if st.button("🧩 Form setini değiştir"):
            st.session_state["page"] = "form_selector"
            st.rerun()
        if st.button("🏠 Session listesi"):
            st.session_state["page"] = "session_manager"
            st.rerun()
    
    st.markdown("---")
    
    # API Key
    col_api1, col_api2 = st.columns([3, 1])
    with col_api1:
        api_key_input = st.text_input(
            "🔑 OpenAI API Key",
            value=st.session_state.get("api_key", ""),
            type="password",
            help="Whisper ve AI analizi için gerekli"
        )
    with col_api2:
        if st.checkbox("Hatırla", value=bool(st.session_state.get("api_key"))):
            st.session_state["api_key"] = api_key_input
        else:
            st.session_state["api_key"] = ""
    
    st.markdown("---")
    
    # Şablon seçimi
    st.subheader("📝 Şablon Belgeleri")
    
    default_dir = os.path.join(os.getcwd(), "templates")
    selected_names = []
    available = []
    
    try:
        if os.path.isdir(default_dir):
            available = sorted([f for f in os.listdir(default_dir) if f.lower().endswith(".docx")])
            if available:
                # Form setine göre ön seçim hazırla (sadece ilk girişte uygula veya grup değiştiyse)
                group = st.session_state.get("selected_form_group")
                def _match_group_files(group_label: Optional[str], files: List[str]) -> List[str]:
                    if not group_label:
                        return []
                    prefixes_map = {
                        "Ek 1-2-3": ["Ek-1", "Ek-2", "Ek-3"],
                        "Ek 4": ["Ek-4"],
                        "Ek 6": ["Ek-6"],
                        "Ek 8": ["Ek-8"],
                        "Ek 9": ["Ek-9"],
                        "Ek 11": ["Ek-11"],
                    }
                    prefixes = prefixes_map.get(group_label, [])
                    return [f for f in files if any(f.startswith(pfx) for pfx in prefixes)]

                should_apply_preselection = (
                    st.session_state.get("templates_initialized_for") != current_session_id or
                    st.session_state.get("form_group_applied") != group
                )
                if should_apply_preselection:
                    preselected = _match_group_files(group, available)
                    st.session_state["selected_templates"] = preselected
                    st.session_state["templates_initialized_for"] = current_session_id
                    st.session_state["form_group_applied"] = group

                selected_names = st.multiselect(
                    "Kullanılacak şablonları seçin",
                    options=available,
                    default=st.session_state.get("selected_templates", []),
                    help="Seçtiğiniz şablonların tam önizlemesi aşağıda görüntülenecek",
                    key="selected_templates"
                )
            else:
                st.info("Templates klasöründe .docx şablon bulunamadı.")
        else:
            st.info("Templates klasörü bulunamadı.")
    except Exception as e:
        st.error(f"Templates klasörü okunamadı: {e}")
    
    template_items = []
    for name in selected_names:
        try:
            full = os.path.join(default_dir, name)
            with open(full, "rb") as fh:
                template_items.append((name, fh.read()))
        except Exception as e:
            st.error(f"{name} okunamadı: {e}")
    
    # Placeholder'ları topla
    union_placeholders = set()
    if template_items:
        for name, data in template_items:
            try:
                placeholders, _ = extract_placeholders_from_docx_bytes(data)
                union_placeholders |= placeholders
            except Exception as e:
                st.error(f"{name} analiz edilemedi: {e}")
    
    st.markdown("---")
    
    # Ses kaydı bölümü
    st.subheader("🎤 Ses Kaydı ve Analiz")
    
    col_mic, col_btn = st.columns([3, 1])
    
    with col_mic:
        audio_bytes = render_audio_recorder_ui()
    
    with col_btn:
        if st.button("🧠 Analiz Et", use_container_width=True, type="primary"):
            effective_key = (api_key_input or st.session_state.get("api_key", "")).strip()
            
            if not template_items:
                st.warning("Önce şablon seçin.")
                return
            if not union_placeholders:
                st.warning("Şablonlarda placeholder bulunamadı.")
                return
            if not effective_key:
                st.warning("OpenAI API anahtarı girin.")
                return

            existing_transcript = (st.session_state.get("current_transcript", "") or "").strip()
            merged_transcript = ""

            if audio_bytes:
                with st.spinner("Ses metne çevriliyor..."):
                    text = transcribe_audio_bytes(audio_bytes, effective_key)
                if not text:
                    st.error("Ses metne çevrilemedi.")
                    return
                merged_transcript = (existing_transcript + " " + text.strip()).strip() if existing_transcript else text.strip()
                st.session_state["current_transcript"] = merged_transcript
                sm.update_session_transcript(current_session_id, merged_transcript)
            elif existing_transcript:
                merged_transcript = existing_transcript
            else:
                st.warning("Ses kaydı yapın veya mevcut transkript bulunmuyor.")
                return

            with st.spinner("Bilgiler çıkarılıyor..."):
                ctx = aggregate_contexts_across_templates(template_items, union_placeholders)
                suggested = infer_placeholder_values(
                    merged_transcript,
                    union_placeholders,
                    ctx,
                    effective_key,
                )
                
                # Mevcut verilerle birleştir
                existing_data = st.session_state.get("current_mapping", {})
                conflicts = detect_conflicts(existing_data, suggested)
                
                if conflicts:
                    st.warning(f"⚠️ {len(conflicts)} çakışma tespit edildi: {', '.join(conflicts)}")
                
                merged_data = merge_extracted_data(existing_data, suggested)
                st.session_state["current_mapping"] = merged_data
                
                # Session'a kaydet
                try:
                    if sm.update_session_data(current_session_id, suggested, merge=True):
                        filled_count = len([v for v in suggested.values() if v.strip()])
                        st.success(f"✅ {filled_count} yeni bilgi eklendi ve kaydedildi!")
                        
                        # Session ismini güncelle
                        updated_session = sm.get_session(current_session_id)
                        if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                            st.session_state["current_session_name"] = updated_session['session_name']
                            st.info("📝 Session ismi güncellendi!")
                except Exception as e:
                    st.warning(f"Veriler çıkarıldı ama kaydetme sırasında hata: {e}")
                
                st.rerun()
    
    # Transkript gösterimi
    if st.session_state.get("current_transcript"):
        col_transcript, col_clear = st.columns([4, 1])
        
        with col_transcript:
            st.text_area(
                "📜 Birleşik Transkript",
                value=st.session_state.get("current_transcript", ""),
                height=120,
                disabled=True,
                help="Bu transkript session bazında saklanır ve tüm Ek formlarında kullanılabilir"
            )
        
        with col_clear:
            st.write("")
            if st.button("🗑️ Temizle"):
                st.session_state["current_transcript"] = ""
                # Session'dan da transcript'i temizle
                sm.update_session_transcript(current_session_id, "")
                st.rerun()
    
    # Placeholder değerleri
    if union_placeholders:
        st.markdown("---")
        st.subheader("✏️ Bilgi Düzenleme")
        
        col_apply, col_clear = st.columns([2, 1])
        with col_apply:
            if st.button("🔄 Session Verilerini Uygula"):
                session_data = sm.get_session(current_session_id)
                if session_data and session_data.get('extracted_data'):
                    current_mapping = st.session_state.get("current_mapping", {})
                    applied_count = 0
                    
                    for ph in union_placeholders:
                        if ph in session_data['extracted_data'] and session_data['extracted_data'][ph]:
                            if ph not in current_mapping or not current_mapping.get(ph):
                                current_mapping[ph] = session_data['extracted_data'][ph]
                                applied_count += 1
                    
                    st.session_state["current_mapping"] = current_mapping
                    if applied_count > 0:
                        st.success(f"✅ {applied_count} alan dolduruldu!")
                        st.rerun()
        
        with col_clear:
            if st.button("🧹 Temizle"):
                st.session_state["current_mapping"] = {}
                st.rerun()
        
        # Placeholder düzenleme
        edit_cols = st.columns(2)
        for idx, ph in enumerate(sorted(list(union_placeholders))):
            with edit_cols[idx % 2]:
                display_name = ph.replace("{", "").replace("}", "")
                st.markdown(f"**{display_name}**")
                
                cur_val = st.session_state.get("current_mapping", {}).get(ph, "")
                new_val = st.text_input(
                    "Değer", 
                    value=cur_val, 
                    key=f"edit_{idx}_{ph}",
                    placeholder="Değer girin...",
                    label_visibility="collapsed"
                )
                
                if new_val != cur_val:
                    st.session_state["current_mapping"][ph] = new_val
                    
                    # Session'a kaydet
                    sm.update_session_data(current_session_id, {ph: new_val}, merge=True)
                    
                    # Session ismini kontrol et
                    updated_session = sm.get_session(current_session_id)
                    if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                        st.session_state["current_session_name"] = updated_session['session_name']
                    
                    st.rerun()
                
                st.markdown("---")
    
    # Seçilen şablonların önizlemesi
    if template_items:
        st.markdown("---")
        st.subheader("👁️ Seçilen Şablonların Önizlemesi")

        for template_name, template_data in template_items:
            # Basit dropdown (expander) ile tam içerik önizleme
            with st.expander(f"📄 {template_name}", expanded=False):
                try:
                    # Word belgesinin tam metnini çıkar
                    doc = Document(io.BytesIO(template_data))
                    parts = []

                    # Paragraflar
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            parts.append(paragraph.text.strip())

                    # Tablolar
                    for table in doc.tables:
                        for row in table.rows:
                            cells = []
                            for cell in row.cells:
                                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                                if cell_text:
                                    cells.append(cell_text)
                            if cells:
                                parts.append(" | ".join(cells))

                    # Header/Footer
                    for section in doc.sections:
                        if section.header:
                            for p in section.header.paragraphs:
                                if p.text.strip():
                                    parts.insert(0, f"[BAŞLIK: {p.text.strip()}]")
                        if section.footer:
                            for p in section.footer.paragraphs:
                                if p.text.strip():
                                    parts.append(f"[ALT BİLGİ: {p.text.strip()}]")

                    full_text = "\n\n".join(parts).strip()

                    if not full_text:
                        st.info("Bu şablonda görüntülenebilir metin bulunamadı.")
                    else:
                        # Tüm placeholder'ları doğrudan metin üzerinden regex ile işle
                        try:
                            current_mapping = st.session_state.get("current_mapping", {}) or {}
                            # Önizlemede işbu alanlarını da bugünün değeriyle doldur
                            mapping_with_isbu = {
                                **current_mapping,
                                **today_isbu(datetime.now(IST))
                            }
                            import html as _html
                            pattern = re.compile(r"\{[^}]+\}")

                            def _replace_placeholder(match: re.Match) -> str:
                                ph = match.group(0)
                                # Hem tam eşleşme hem de kıvrıksız anahtar ile eşleşmeyi dene
                                raw_val = str(mapping_with_isbu.get(ph, "")).strip()
                                if not raw_val:
                                    key_nobraces = ph.strip('{}')
                                    # Önce {key} biçimindeki varyantları tara
                                    for k, v in mapping_with_isbu.items():
                                        if isinstance(k, str) and k.strip('{}').lower() == key_nobraces.lower():
                                            raw_val = str(v).strip()
                                            if raw_val:
                                                break
                                if raw_val:
                                    return _html.escape(raw_val)
                                return f"<span style=\"color:#ff4d4f;font-weight:700;\">{_html.escape(ph)}</span>"

                            highlighted_text = pattern.sub(_replace_placeholder, full_text)
                        except Exception:
                            highlighted_text = full_text

                        # Sade, tam genişlikte metin (inline renk ile)
                        st.markdown(
                            f"""
                            <div style="white-space: pre-wrap; word-wrap: break-word; line-height: 1.75; font-size: 16px; font-weight: 500; color: #374151;">{highlighted_text}</div>
                            """,
                            unsafe_allow_html=True
                        )
                except Exception as e:
                    st.error(f"Şablon önizlemesi oluşturulamadı: {e}")
    
    # Belge oluşturma
    if template_items:
        st.subheader("📄 Belge Oluşturma")
        
        if st.button("📄 Tüm Belgeleri Oluştur", type="primary", use_container_width=True):
            if not st.session_state.get("current_mapping"):
                st.warning("Önce bilgileri doldurun.")
            else:
                try:
                    results = []
                    current_mapping = st.session_state["current_mapping"]
                    
                    for idx, (name, data) in enumerate(template_items):
                        doc = Document(io.BytesIO(data))
                        mapping = {k: v for k, v in current_mapping.items() if str(v).strip()}
                        # İşbu alanlarını belge oluşturma anının tarihi/saatine sabitle
                        mapping = {
                            **mapping,
                            **today_isbu(datetime.now(IST))
                        }
                        replaced = replace_placeholders_in_document(doc, mapping)
                        
                        buf = io.BytesIO()
                        doc.save(buf)
                        out_bytes = buf.getvalue()
                        
                        # Dosya adı
                        safe_session_name = re.sub(r'[^\w\s-]', '', current_session_name).strip()[:20]
                        out_name = f"{safe_session_name}_{os.path.splitext(name)[0]}.docx"
                        
                        results.append({
                            "name": out_name,
                            "replaced": replaced,
                            "data": out_bytes,
                            "key": f"dl_{idx}_{out_name}",
                        })
                    
                    st.session_state["results"] = results
                    st.success("✅ Belgeler hazırlandı!")
                except Exception as e:
                    st.error(f"Belge oluşturma hatası: {e}")
        
        # İndirme butonları
        if st.session_state.get("results"):
            st.markdown("---")
            st.subheader("📥 İndirilecek Belgeler")
            
            for r in st.session_state["results"]:
                col_info, col_download = st.columns([3, 1])
                
                with col_info:
                    st.write(f"**{r['name']}** → {r['replaced']} değişiklik")
                
                with col_download:
                    st.download_button(
                        label="📥 İndir",
                        data=r["data"],
                        file_name=r["name"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=r["key"],
                        use_container_width=True
                    )

if __name__ == "__main__":
    main()

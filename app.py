# app.py - Sesli Belge Doldurma Sistemi
# Voice.py ve app2.py projelerinin en iyi özelliklerini birleştiren gelişmiş sistem

import io
import os
import re
import json
import sys
import tempfile
import uuid
import traceback
from typing import Dict, List, Optional, Set, Tuple
from datetime import datetime, date, time
from zoneinfo import ZoneInfo

import streamlit as st
import importlib
from docx import Document
import dateparser

# Local session management import
from local_session_manager import get_local_session_manager, merge_extracted_data, detect_conflicts
# User management import
from user_manager import get_user_manager
# Feedback management import
from feedback_manager import get_feedback_manager

# Özel form davranışları (Ek bazlı özel prompt ve alan kısıtlama)
# Burada Ek 15 için, uzun metni 4 parçaya ayırma talimatını tanımlayabilirsiniz.
# expected_placeholders: Şablondaki TAM placeholder anahtarları (örn: "{metin1}")
# placeholder_explanations: Her anahtar için kısa açıklama (opsiyonel)
SPECIAL_FORMS: Dict[str, Dict[str, object]] = {
    "Ek 15": {
        "expected_placeholders": [
            "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}",
            "{iddilar_hakkinda_ne_diyorsunuz}",
            "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}",
            "{iddia_nedir}",
        ],
        "placeholder_explanations": {
            "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}": "Şüpheli olayla ilgili eklemek istediği ek açıklamalar (kısa, öz).",
            "{iddilar_hakkinda_ne_diyorsunuz}": "Öğrencinin iddialara cevabı/ifadesi; kabul/inkâr ve gerekçe.",
            "{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}": "Tutanağa dair ekleme/çıkarma isteği; varsa düzeltmeler.",
            "{iddia_nedir}": "Hakkındaki iddianın özeti; kopya iddiasının kısa tanımı.",
        },
        "custom_instructions": """
ÖZEL TALİMATLAR - EK 15 İFADE ANALİZİ:

Bu sistem Ek-15 ifade formu için 4 soru alanını öğrencinin ağzından doldurur. Öğrencinin verdiği uzun ifade metnini analiz ederek, sanki öğrenci bu 4 soruya teker teker cevap veriyormuş gibi doldur.

TEMEL GÖREV:
- Sorular önceden bellidir, sen sadece öğrencinin bu sorulara vereceği CEVAPLARI üret
- Her cevap öğrencinin ağzından, birinci şahıs olarak yazılacak
- Öğrencinin ifadesindeki pişmanlık/pişman değil durumunu doğru yansıt
- Verilen metin hangi senaryoyu içeriyorsa ona uygun cevaplar üret

GENEL İLKELER:
- Tüm cevaplar doğal, insancıl ve öğrencinin kendi sesiyle olacak
- Robotik, yapay ifadeler kullanma
- Cevaplar ayrıntılı ve kapsamlı olacak (minimum 2-3 cümle)
- Açıkça geçmeyen bilgiyi uydurma; emin değilsen o alanı boş bırak
- ÖNEMLI: Sadece CEVAPLAR üret, soruları tekrar yazma

ÖZEL DURUMLAR:
1. SUÇ KABULÜ: Eğer öğrenci ifadesinde suçunu kabul etmiş ise (açık veya örtülü), bu kabul tüm ilgili cevaplara yansıtılacak
2. PİŞMANLIK: Eğer öğrenci pişmanlık belirtmiş ise, bunu özellikle "{konu_hk_eklemek_istediginiz_bir_sey_var_mi}" alanında vurgula

ALAN BAZLI TALİMATLAR:

{iddia_nedir}: 
- Bu alanda öğrencinin "Hakkınızda yöneltilen iddia nedir? Açıklar mısınız?" sorusuna vereceği cevabı yaz
- Öğrenci ağzından, kendisine yöneltilen iddiayı kabul edip etmediğini açıklayacak
- Olayın ne olduğunu öğrencinin kendi ifadesiyle anlat
- Birinci şahıs ("ben", "benim") kullanarak öğrencinin ağzından yaz

{iddilar_hakkinda_ne_diyorsunuz}:
- Bu alanda öğrencinin "İddia/iddialar hakkında ne diyorsunuz. Ayrıntılı açıklayarak anlatınız?" sorusuna vereceği cevabı yaz
- Öğrencinin olayı kendi ağzından detaylı şekilde anlatmasını sağla
- Olay nasıl gelişti, ne yaptı, neden yaptı gibi ayrıntıları dahil et
- Birinci şahıs anlatımla öğrencinin perspektifini yansıt

{konu_hk_eklemek_istediginiz_bir_sey_var_mi}:
- Bu alanda öğrencinin "Konu hakkında eklemek istediğiniz başka bir şey var mı?" sorusuna vereceği cevabı yaz
- Pişmanlık ifadeleri, özür beyanları, ek açıklamaları öğrencinin ağzından ekle
- Öğrencinin duygusal durumu ve gelecek planlarını birinci şahıs olarak ifade et
- Kendisinin söylemek istediği ek noktaları dahil et

{tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}:
- Bu alanda öğrencinin "Tutanağı okuyunuz, eklenmesini, düzeltilmesini istediğiniz bir bölüm var mı?" sorusuna vereceği cevabı yaz
- Öğrencinin ifadesinde tutanakla ilgili açık bir düzeltme/ekleme talebi varsa o talebi birinci şahıs olarak belirt
- Eğer öğrencinin ifadesinde tutanak hakkında herhangi bir şikayet/düzeltme talebi GEÇMİYORSA, şu tarzda standart cevap ver:
  "Tutanakta eklenmesini veya çıkarılmasını istediğim bir bölüm bulunmamaktadır" veya
  "Eklemek veya çıkarmak istediğim bir şey yoktur" veya
  "Tutanağın bu şekilde kalmasında sakınca görmüyorum"
- Birinci şahıs ifadeyle ve doğal dilde yaz

CEVAP UZUNLUK VE KALİTE KURALLARI:
- Her cevap minimum 2-3 cümle uzunluğunda olmalı
- Tek kelime veya kısa cevaplar kabul edilemez
- Cevaplar ayrıntılı, anlamlı ve kapsamlı olacak
- Pişmanlık durumunda açık pişmanlık ifadeleri kullanılacak
- Pişman değilse olay yumuşatılarak ama net şekilde aktarılacak

Bu format ve ton kullanılarak, öğrencinin gerçek ifadesinden benzer cevaplar üret.

Yalnızca şu anahtarlar için çıktı ver: {konu_hk_eklemek_istediginiz_bir_sey_var_mi}, {iddilar_hakkinda_ne_diyorsunuz}, {tutanagi_okuyun_eklemek_cikarmak_istediginiz_yer_var_mi}, {iddia_nedir}.
""",
    }
}

# OpenAI import
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Mikrofon kütüphaneleri - birincil ve alternatifler
MIC_IMPORT_ERROR: Optional[str] = None
mic_recorder = None  # from streamlit_mic_recorder
audio_recorder_fn = None  # from audio_recorder_streamlit
st_audiorec_fn = None  # from st_audiorec

try:
    _mic_module = importlib.import_module("streamlit_mic_recorder")
    mic_recorder = getattr(_mic_module, "mic_recorder", None)
except Exception as e:
    MIC_IMPORT_ERROR = f"streamlit-mic-recorder yüklenemedi: {e}"

try:
    _ar_module = importlib.import_module("audio_recorder_streamlit")
    audio_recorder_fn = getattr(_ar_module, "audio_recorder", None)
except Exception:
    pass

try:
    _sar_module = importlib.import_module("st_audiorec")
    st_audiorec_fn = getattr(_sar_module, "st_audiorec", None)
except Exception:
    pass

def render_audio_recorder_ui() -> Optional[bytes]:
    """Tarayıcıdan ses kaydı al ve bytes döndür (mevcut bileşene göre)."""
    # Öncelik: streamlit-mic-recorder
    if mic_recorder is not None:
        st.write("**Mikrofon ile Kayıt**")
        rec_val = mic_recorder(
            start_prompt="🎙️ Kaydı Başlat",
            stop_prompt="⏹️ Kaydı Durdur",
            just_once=False,
            use_container_width=True,
            key="unified_mic_recorder",
        )
        if isinstance(rec_val, dict) and rec_val.get("error"):
            st.error(f"Mikrofon hatası: {rec_val['error']}")
            return None
        return bytes_from_mic_return(rec_val)

    # Alternatif 1: audio-recorder-streamlit
    if audio_recorder_fn is not None:
        st.write("**Mikrofon ile Kayıt (alternatif)**")
        rec_val = audio_recorder_fn()
        return bytes_from_mic_return(rec_val) if rec_val else None

    # Alternatif 2: streamlit-audiorec
    if st_audiorec_fn is not None:
        st.write("**Mikrofon ile Kayıt (alternatif)**")
        rec_val = st_audiorec_fn()
        return bytes_from_mic_return(rec_val) if rec_val else None

    st.error("Mikrofon kütüphanesi mevcut değil.")
    if MIC_IMPORT_ERROR:
        st.error(MIC_IMPORT_ERROR)
    st.info("Lütfen 'streamlit-mic-recorder' veya 'audio-recorder-streamlit' paketini kurun.")
    return None

# Zaman dilimi
IST = ZoneInfo("Europe/Istanbul")
TR_DAYS = {0:"Pazartesi", 1:"Salı", 2:"Çarşamba", 3:"Perşembe", 4:"Cuma", 5:"Cumartesi", 6:"Pazar"}

# ================== Yardımcı Fonksiyonlar ==================

def ensure_utf8_encoding():
    """Sistem encoding'ini kontrol et ve gerekirse UTF-8'e ayarla"""
    import locale
    try:
        # Sistem locale'ini kontrol et
        current_encoding = locale.getpreferredencoding()
        if 'utf-8' not in current_encoding.lower() and 'cp65001' not in current_encoding.lower():
            # Windows'ta UTF-8 desteği için
            import os
            os.environ['PYTHONIOENCODING'] = 'utf-8'
    except Exception:
        pass  # Encoding ayarlanamadıysa sessizce devam et

def safe_str(obj) -> str:
    """Herhangi bir objeyi güvenli şekilde string'e çevir"""
    try:
        return str(obj)
    except UnicodeError:
        try:
            return str(obj).encode('utf-8', errors='replace').decode('utf-8')
        except:
            return repr(obj)

def bytes_from_mic_return(value) -> Optional[bytes]:
    """Mikrofon dönüş değerini normalize et"""
    if value is None:
        return None
    if isinstance(value, dict) and "bytes" in value:
        return value["bytes"]
    if isinstance(value, (bytes, bytearray)):
        return bytes(value)
    return None

def transcribe_audio_bytes(audio_bytes: bytes, api_key: str, lang: str = "tr") -> Optional[str]:
    """Ses dosyasını metne çevir"""
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil. 'openai' paketini kurun.")
        return None

    tmp_path = None
    try:
        # API key'i güvenli şekilde işle
        safe_api_key = api_key.strip() if api_key else ""
        if not safe_api_key:
            st.error("API key boş veya geçersiz")
            return None
            
        client = OpenAI(api_key=safe_api_key)
        
        # Güvenli temp dosya oluştur - Unicode güvenli
        safe_filename = f"audio_{uuid.uuid4().hex}.wav"
        tmp_dir = tempfile.gettempdir()
        tmp_path = os.path.join(tmp_dir, safe_filename)
        
        # Dosya yolunun Unicode karakterler içerip içermediğini kontrol et
        try:
            tmp_path.encode('ascii')
        except UnicodeEncodeError:
            # ASCII olmayan karakterler varsa, farklı bir yol kullan
            tmp_path = os.path.join(tempfile.gettempdir(), f"temp_audio_{uuid.uuid4().hex[:8]}.wav")
        
        with open(tmp_path, "wb") as f:
            f.write(audio_bytes)

        with open(tmp_path, "rb") as f:
            resp = client.audio.transcriptions.create(
                model="whisper-1", 
                file=f,
                language=lang,
                response_format="text"
            )

        # Response'u güvenli şekilde işle
        if isinstance(resp, str):
            # Türkçe karakterleri güvenli şekilde handle et
            try:
                return resp
            except UnicodeError:
                return resp.encode('utf-8', errors='replace').decode('utf-8')
        else:
            text_result = getattr(resp, "text", None) or (resp.get("text") if isinstance(resp, dict) else None)
            if text_result:
                try:
                    return str(text_result)
                except UnicodeError:
                    return str(text_result).encode('utf-8', errors='replace').decode('utf-8')
            return None
        
    except Exception as e:
        # Unicode karakterleri tamamen güvenli şekilde işle
        try:
            # Exception mesajını güvenli şekilde al
            error_msg = str(e)
            # ASCII dışı karakterleri kontrol et ve temizle
            error_msg.encode('ascii')
        except (UnicodeError, UnicodeEncodeError, UnicodeDecodeError):
            # Unicode sorunu varsa, güvenli bir mesaj kullan
            try:
                error_msg = repr(str(e))  # repr() kullanarak güvenli gösterim
            except:
                error_msg = "Ses işleme sırasında Unicode karakter hatası oluştu"
        
        # Streamlit error mesajını da güvenli şekilde göster
        safe_error_msg = safe_str(error_msg)
        st.error(f"Ses metne çevrilemedi: {safe_error_msg}")
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass

def extract_placeholders_from_docx_bytes(file_bytes: bytes) -> Tuple[Set[str], str]:
    """Word dosyasından placeholder'ları çıkar"""
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    
    # Tüm metinleri topla
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + " "
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text += p.text + " "
    
    # Header/Footer
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    text += p.text + " "
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    text += p.text + " "
    
    placeholders = set(re.findall(r"\{[^}]+\}", text))
    return placeholders, text

def replace_placeholders_in_document(doc: Document, placeholder_values: Dict[str, str]) -> int:
    """Word belgesindeki placeholder'ları değiştir"""
    def replace_in_paragraph(paragraph):
        if not paragraph.runs:
            return 0
        original_text = "".join(run.text for run in paragraph.runs)
        replaced_text = original_text
        total_replacements = 0
        
        for placeholder, value in placeholder_values.items():
            if value is None:
                continue
            count = replaced_text.count(placeholder)
            if count:
                replaced_text = replaced_text.replace(placeholder, str(value))
                total_replacements += count
        
        if replaced_text != original_text:
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            paragraph.add_run(replaced_text)
        
        return total_replacements

    replacements_made = 0
    
    # Body paragraphs
    for p in doc.paragraphs:
        replacements_made += replace_in_paragraph(p)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replacements_made += replace_in_paragraph(p)
    
    # Headers/Footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replacements_made += replace_in_paragraph(p)
        if section.footer:
            for p in section.footer.paragraphs:
                replacements_made += replace_in_paragraph(p)
    
    return replacements_made

def parse_tr_date(text: str) -> Optional[datetime]:
    """Türkçe tarih formatlarını çözümle"""
    if not text:
        return None
    return dateparser.parse(text, languages=["tr"])

def split_date(dt: datetime) -> Dict[str, str]:
    """Datetime'ı gün/ay/yıl olarak böl"""
    return {
        "gun": f"{dt.day:02d}",
        "ay": f"{dt.month:02d}",
        "yil": f"{dt.year}"
    }

def today_isbu(dt: datetime = None) -> Dict[str, str]:
    """İşbu tarih/saat bilgilerini al"""
    now = dt or datetime.now(IST)
    return {
        "isbu_gun": f"{now.day:02d}",
        "isbu_ay": f"{now.month:02d}",
        "isbu_yil": f"{now.year}",
        "isbu_saat": now.strftime("%H:%M")
    }

# ================== Etiket Biçimlendirme (Kullanıcı Dostu) ==================

def _turkish_capitalize(word: str) -> str:
    """Türkçeye uygun büyük harfe çevirme (i→İ, ı→I)."""
    if not word:
        return ""
    first = word[0]
    rest = word[1:]
    if first == "i":
        first_u = "İ"
    elif first == "ı":
        first_u = "I"
    else:
        first_u = first.upper()
    return first_u + rest

def format_placeholder_label(placeholder: str) -> str:
    """{ogrenci_adi} → Öğrenci Adı gibi kullanıcı dostu etiket üretir."""
    try:
        key = str(placeholder or "")
        if key.startswith("{") and key.endswith("}"):
            key = key[1:-1]
        key = key.replace("-", " ").replace("_", " ")
        key = re.sub(r"\s+", " ", key).strip()

        special_map = {
            "tc": "T.C.",
            "t.c": "T.C.",
            "ogr": "Öğr.",
            "ogrenci": "Öğrenci",
            "öğrenci": "Öğrenci",
            "ad": "Ad",
            "adi": "Adı",
            "soyad": "Soyad",
            "soyadi": "Soyadı",
            "adsoyad": "Ad Soyad",
            "adi soyadi": "Adı Soyadı",
            "bolum": "Bölüm",
            "bölüm": "Bölüm",
            "no": "No",
            "numara": "Numarası",
            "kimlik": "Kimlik",
            "sinav": "Sınav",
            "sınav": "Sınav",
            "tarih": "Tarih",
            "saat": "Saat",
            "blok": "Blok",
            "ders": "Ders",
            "gozetmen": "Gözetmen",
            "gözetmen": "Gözetmen",
            "imza": "İmza",
            "yer": "Yer",
            "salon": "Salon",
            "derslik": "Derslik",
            "fakulte": "Fakülte",
            "fakülte": "Fakülte",
            "unvan": "Unvan",
            "unvani": "Unvanı",
            "aciklama": "Açıklama",
            "açıklama": "Açıklama",
        }

        words = [w for w in key.split(" ") if w]
        pretty_words = []
        for w in words:
            lw = w.lower()
            if lw in special_map:
                pretty_words.append(special_map[lw])
            elif len(w) <= 2 and w.isalpha():
                pretty_words.append(w.upper())
            else:
                pretty_words.append(_turkish_capitalize(w))

        label = " ".join(pretty_words)
        label = label.replace("Ogrenci", "Öğrenci").replace("Ogrencinin", "Öğrencinin")
        return label
    except Exception:
        return str(placeholder).strip("{}")

# ================== AI Analiz Fonksiyonları ==================

def extract_placeholder_contexts_from_docx_bytes(file_bytes: bytes, placeholders: Set[str], window: int = 120) -> Dict[str, List[str]]:
    """Placeholder'ların bağlamlarını çıkar"""
    doc = Document(io.BytesIO(file_bytes))
    blocks = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        blocks.append(p.text)
    
    contexts = {ph: [] for ph in placeholders}
    
    for block in blocks:
        for ph in placeholders:
            pattern = re.escape(ph)
            for m in re.finditer(pattern, block):
                start, end = m.start(), m.end()
                before = block[max(0, start - window): start]
                after = block[end: end + window]
                snippet = f"{before}{ph}{after}"
                if len(contexts[ph]) < 3:
                    contexts[ph].append(snippet)
    
    return contexts

def aggregate_contexts_across_templates(templates: List[Tuple[str, bytes]], placeholders: Set[str]) -> Dict[str, List[str]]:
    """Tüm şablonlardan bağlamları topla"""
    combined = {ph: [] for ph in placeholders}
    
    for _, data in templates:
        try:
            local_ctx = extract_placeholder_contexts_from_docx_bytes(data, placeholders)
            for ph, lst in local_ctx.items():
                for s in lst:
                    if len(combined[ph]) < 5:
                        combined[ph].append(s)
        except Exception:
            continue
    
    return combined

def parse_json_loose(s: str) -> Dict[str, str]:
    """JSON'u esnek şekilde parse et"""
    try:
        return json.loads(s)
    except Exception:
        pass
    
    try:
        m = re.search(r"\{[\s\S]*\}", s)
        if m:
            return json.loads(m.group(0))
    except Exception:
        pass
    
    return {}

# Regex tabanlı fallback kaldırıldı (tüm çıkarım LLM ile yapılacak)

def infer_placeholder_values(
    transcript: str,
    placeholders: Set[str],
    contexts: Dict[str, List[str]],
    api_key: str,
    model: str = "gpt-4o-mini",
    only_placeholders: Optional[Set[str]] = None,
    extra_instructions: Optional[str] = None,
    placeholder_explanations: Optional[Dict[str, str]] = None,
) -> Dict[str, str]:
    """AI ile placeholder değerlerini çıkar"""
    if OpenAI is None:
        st.error("OpenAI SDK mevcut değil.")
        return {}
    
    client = OpenAI(api_key=api_key)
    ph_list = sorted(list(placeholders))
    # Eğer sadece belirli placeholder'lar isteniyorsa, set'i daralt
    if only_placeholders:
        wanted = {p.strip("{}").lower() for p in only_placeholders}
        filtered: List[str] = []
        for ph in ph_list:
            key_nb = ph.strip("{}").lower()
            if key_nb in wanted:
                filtered.append(ph)
        ph_list = filtered or ph_list
    
    # Mevcut (kullanıcı tarafından girilmiş) değerleri al ve prompt'a ekle (değiştirme)
    existing_values = {}
    try:
        existing_values = {
            k: v for k, v in (st.session_state.get("current_mapping", {}) or {}).items()
            if k in placeholders and str(v).strip()
        }
    except Exception:
        existing_values = {}

    # Gelişmiş prompt
    # İsim ve bölüm alanları için kuralları dinamik üret
    ph_lower_list = [ph.lower() for ph in placeholders]
    has_fullname_key = any(("ogrenci" in p and ("adi_soyadi" in p or "ad_soyad" in p)) for p in ph_lower_list)
    has_name_key = any(("ogrenci" in p and ("ad" in p or "adi" in p or "isim" in p) and "soyad" not in p) for p in ph_lower_list)
    has_surname_key = any(("ogrenci" in p and ("soyad" in p or "soyadi" in p)) for p in ph_lower_list)
    has_department_key = any(("bolum" in p) or ("bölüm" in p) for p in ph_lower_list)

    name_rules_lines: List[str] = []
    if has_fullname_key:
        name_rules_lines.append("- {ogrenci_adi_soyadi} alanı için öğrencinin tam adını 'Ad Soyad' formatında ver (örn: 'Ecem Nalbantoğlu').")
    if has_name_key and has_surname_key:
        name_rules_lines.append("- Ayrı alanlar varsa {ogrenci_ad}/{ogrenci_adi} ve {ogrenci_soyad}/{ogrenci_soyadi} alanlarını ayrı ayrı doldur (örn: 'Ecem' ve 'Nalbantoğlu').")
    elif has_name_key or has_surname_key:
        name_rules_lines.append("- Öğrenci adı/soyadı alanları varsa transkriptte geçtiği şekliyle doldur.")

    department_rules_lines: List[str] = []
    if has_department_key:
        department_rules_lines.append("- {bolum}/{bolum_adi} gibi bölüm alanlarında tek bir bölüm adı ver. Birden fazla bölüm yazma (örn: 'İşletme' veya 'Bilgisayar Mühendisliği'; 'İşletme Muhasebe' yazma).")

    prompt_text = f"""
SES TRANSKRİPTİ:
"{transcript}"

TEMPLATE PLACEHOLDER'LARI VE BAĞLAMLARI:
"""
    
    for ph in ph_list:
        if ph in contexts and contexts[ph]:
            context_examples = "\n".join([f"  • {ctx[:200]}" for ctx in contexts[ph][:3]])
            prompt_text += f"\n{ph}:\n{context_examples}\n"
        else:
            prompt_text += f"\n{ph}: (Bağlam bulunamadı)\n"
    
    prompt_text += """

GÖREV:
1. Ses transkriptini analiz et
2. Her placeholder için template bağlamını incele
3. Bağlama uygun değerleri ses transkriptinden çıkar
4. Çıkaramadığın bilgiler için boş string ("") bırak
5. SADECE JSON formatında cevap ver

MEVCUT DEĞERLER (DEĞİŞTİRME):
""" + json.dumps(existing_values, ensure_ascii=False, indent=2) + """

ÇIKTI KURALLARI:
- JSON anahtarları, placeholder stringleriyle birebir aynı olmalı (örnek: {ogrenci_no})
- Mevcut dolu alanları DEĞİŞTİRME; sadece boş olanları doldur
- Tarih ve saat alanları bağlama uygun normalize edilmeli (tarih: YYYY-MM-DD veya {gun,ay,yil}; saat: HH:MM)
- Sayısal alanlar (no, tc vb.) sadece rakam içersin
- İsim alanlarında gereksiz ekleri çıkar; açıklama alanlarında öğrenci ismi geçmesin
 - {blok} alanı sadece tek büyük harf (A-Z) olmalı (ör: A, B, C). Tahmin etme; transkriptte yoksa boş bırak.
 - {ogrenci_adi_soyadi} alanına öğrencinin tam adı ve soyadı gelmeli (örn: "Emre Yılmaz").

EK ÖZEL KURALLAR:
"""
    # Dinamik ek kuralları prompt'a ekle
    if name_rules_lines or department_rules_lines:
        extra_rules = "\n".join(name_rules_lines + department_rules_lines)
        prompt_text += extra_rules + "\n"
    prompt_text += """

ÖZEL İSTEK:
- Açıklama alanlarında sadece olayın kendisini yaz
- Öğrencinin adı ve soyadını açıklama alanlarına ekleme
- Sadece ne olduğunu objektif şekilde açıkla
 - Verilmeyen bilgileri uydurma; emin değilsen boş string ver
 - ÖNEMLİ: Eğer bir alan "açıklama" niteliğindeyse (anahtar isminde "aciklama"/"açıklama" geçiyorsa), ürettiğin cümleyi şablondaki bu placeholder'ın ÖNÜNDE ve ARDINDA geçen kelime/ifadelere dilbilgisel olarak UYDUR. Örn: "... hakkında {aciklama}" kalıbında "... hakkında"dan sonra doğal akışla devam edecek bir ifade kur.
 - "Açıklama" üretirken, bağlam parçalarında (context) placeholder'ı çevreleyen 1-2 kelimeye özellikle dikkat et; gerektiğinde giriş/bağlaç ekleyerek (ör. "hakkında", "ile ilgili", "bu kapsamda", "bu doğrultuda") akıcı hale getir.

JSON formatı örneği:
""" + "{" + ", ".join([f'"{ph}": "değer_veya_boş_string"' for ph in ph_list[:3]]) + "...}"

    # Ek özel talimatlar/anahtar açıklamaları
    if placeholder_explanations:
        try:
            lines = []
            for ph in ph_list:
                desc = placeholder_explanations.get(ph) or placeholder_explanations.get(ph.strip("{}"))
                if desc:
                    lines.append(f"- {ph}: {desc}")
            if lines:
                prompt_text += "\nANAHTAR AÇIKLAMALARI:\n" + "\n".join(lines)
        except Exception:
            pass
    if extra_instructions:
        prompt_text += "\nEK TALİMATLAR:\n" + str(extra_instructions) + "\n"

    messages = [
        {"role": "system", "content": "Uzman bir bilgi çıkarım asistanısın. Kullanıcı transkriptini ve template bağlamlarını analiz ederek, placeholder anahtarlarıyla birebir eşleşen JSON üretirsin. Mevcut dolu değerleri asla değiştirme; sadece eksik (boş) alanları doldur. Tarih/saat ve sayısal alanları normalize et. Açıklama alanında öğrenci ismi geçmesin. Sadece transkriptte açıkça geçen bilgileri kullan; emin olmadığın durumda boş string ver. Sadece JSON döndür."},
        {"role": "user", "content": prompt_text},
    ]
    
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.0,
        )
        content = resp.choices[0].message.content if resp and resp.choices else "{}"
        data = parse_json_loose(content or "{}")
        
        result = {}
        for ph in placeholders:
            key_lower = ph.lower()
            raw_val = str(data.get(ph, "")).strip() if isinstance(data, dict) else ""
            val = raw_val
            # Özel kural: {blok} sadece tek harf (A-Z)
            if "blok" in key_lower:
                import re as _re
                only_letters = "".join(ch for ch in val if ch.isalpha())
                val = only_letters[:1].upper() if only_letters else ""
            # Özel kural: bölüm alanlarında tek bölüm adı döndür
            if ("bolum" in key_lower) or ("bölüm" in key_lower):
                lowered = val.lower()
                # Önce ayırıcılarla kes
                for sep in [",", "/", "&", "|", ";"]:
                    if sep in val:
                        val = val.split(sep)[0]
                # Bağlaçlara göre kes (ve/veya)
                for conj in [" ve ", " veya "]:
                    if conj in lowered:
                        idx = lowered.index(conj)
                        val = val[:idx]
                        break
                val = str(val).strip()
            # Genel: güvenli string
            if val:
                try:
                    result[ph] = str(val).strip()
                except UnicodeEncodeError:
                    result[ph] = val.encode('utf-8', errors='replace').decode('utf-8')
            else:
                result[ph] = ""

        # İsim alanları için ek post-processing: fullname <-> ad/soyad senkronizasyonu
        try:
            # Anahtarların küçük harf normalize edilmiş haritasını oluştur
            keys_by_lower = {k.lower(): k for k in result.keys()}
            # Varyantları bul
            fullname_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("adi_soyadi" in lk or "ad_soyad" in lk))), None)
            name_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("ad" in lk or "adi" in lk or "isim" in lk) and "soyad" not in lk)), None)
            surname_key = next((k for lk, k in keys_by_lower.items() if ("ogrenci" in lk and ("soyad" in lk or "soyadi" in lk))), None)

            # Eğer fullname boş ama ad ve soyad doluysa, birleştir
            if fullname_key and (not result.get(fullname_key)) and name_key and surname_key and result.get(name_key) and result.get(surname_key):
                combined = f"{str(result.get(name_key)).strip()} {str(result.get(surname_key)).strip()}".strip()
                result[fullname_key] = combined
            # Eğer ad/soyad boş ama fullname doluysa, basit böl
            if fullname_key and result.get(fullname_key) and ((name_key and not result.get(name_key)) or (surname_key and not result.get(surname_key))):
                fullname_val = str(result.get(fullname_key)).strip()
                parts = [p for p in fullname_val.split() if p]
                if len(parts) >= 2:
                    first = " ".join(parts[:-1])
                    last = parts[-1]
                    if name_key and not result.get(name_key):
                        result[name_key] = first
                    if surname_key and not result.get(surname_key):
                        result[surname_key] = last
        except Exception:
            pass
        
        return result
    except Exception as e:
        # Unicode güvenli hata mesajı
        try:
            error_msg = str(e)
        except UnicodeEncodeError:
            error_msg = "AI analizi sırasında karakter kodlama hatası"
        st.error(f"Ses analizi başarısız: {error_msg}")
        return {ph: "" for ph in placeholders}

# ================== Öğrenci Yönetimi ==================

def extract_student_info(session_data):
    """Session'dan öğrenci bilgilerini çıkar"""
    if not session_data or not session_data.get('extracted_data'):
        return None, None
    
    extracted = session_data['extracted_data']
    student_no = None
    student_name = None
    
    for key, value in extracted.items():
        if value and str(value).strip():
            key_lower = key.lower().replace('{', '').replace('}', '')
            value_str = str(value).strip()
            
            # Öğrenci dışı kişi alanlarını atla
            skip_person_keywords = [
                'gozetmen', 'gözetmen', 'ogretim', 'öğretim', 'elemani', 'elemanı',
                'gorevli', 'görevli', 'bolum_baskanligi', 'bölüm başkanlığı', 'baskan', 'başkan',
                'danisman', 'danışman', 'sifre', 'şifre', 'yetkili', 'imza'
            ]
            if any(k in key_lower for k in skip_person_keywords):
                continue

            # Öğrenci numarası
            if 'ogrencino' in key_lower or 'ogrenci_no' in key_lower:
                student_no = value_str
            elif 'no' in key_lower and not student_no:
                student_no = value_str
            
            # Öğrenci adı
            elif ('ogrenci' in key_lower) and ('ad' in key_lower or 'adi' in key_lower or 'isim' in key_lower) and 'soyad' not in key_lower:
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{value_str} {student_name}"
                    else:
                        student_name = value_str
            elif ('ogrenci' in key_lower) and ('soyad' in key_lower or 'soyadi' in key_lower):
                if not value_str.isdigit():
                    if student_name:
                        student_name = f"{student_name} {value_str}"
                    else:
                        student_name = value_str
            elif 'ogrenci' in key_lower and any(keyword in key_lower for keyword in ['adi_soyadi', 'ad_soyad']):
                if not value_str.isdigit():
                    student_name = value_str
    
    return student_no, student_name

def update_session_name_if_needed(session_id, session_data):
    """Öğrenci bilgileri varsa session ismini standart formata güncelle."""
    try:
        def _format_standard(no: str, name: str) -> str:
            safe_no = (no or "").strip()
            safe_name = " ".join((name or "").split())
            return f"{safe_no} - {safe_name}" if safe_no and safe_name else ""
        
        def _is_already_standard_format(name: str) -> bool:
            """Session adının zaten standart formatta olup olmadığını kontrol et"""
            # Standart format: "numara - isim" şeklinde
            import re
            pattern = r'^\d+\s*-\s*.+$'
            return bool(re.match(pattern, name.strip()))

        student_no, student_name = extract_student_info(session_data)
        current_name = session_data.get('session_name', '')

        # Eğer session adı zaten standart formattaysa güncelleme
        if _is_already_standard_format(current_name):
            return False

        # Sadece her ikisi de varsa ve henüz standart format değilse isim uygula
        if student_no and student_name:
            new_name = _format_standard(student_no, student_name)
            if new_name and new_name != current_name:
                sm = get_local_session_manager()
                session_data['session_name'] = new_name
                return sm.save_session(session_id, session_data)
        return False
    except Exception as e:
        st.error(f"Session ismi güncellenirken hata: {e}")
        return False

# ================== Kimlik Doğrulama Sayfaları ==================

def show_login():
    """Giriş sayfası"""
    # Başlık - merkezi ve güzel görünüm
    st.markdown("<div style='text-align: center;'>", unsafe_allow_html=True)
    st.title("🔐 Giriş Yap")
    st.caption("Sesli Belge Doldurma Sistemine Hoş Geldiniz")
    st.markdown("</div>", unsafe_allow_html=True)
    
    um = get_user_manager()
    users = um.get_all_users()
    
    if not users:
        st.warning("Henüz kayıtlı kullanıcı yok. Lütfen önce kayıt olun.")
        if st.button("📝 Kayıt Ol"):
            st.session_state["page"] = "register"
            st.rerun()
        return
    
    # Giriş formu - ortalanmış ve düzenli
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        with st.form("login_form"):
            st.markdown("### 📋 Giriş Bilgileri")
            
            username = st.text_input(
                "👤 Kullanıcı Adı", 
                placeholder="Kullanıcı adınızı girin",
                help="Kayıt olurken belirttiğiniz kullanıcı adı"
            )
            
            password = st.text_input(
                "🔒 Şifre", 
                type="password", 
                placeholder="Şifrenizi girin",
                help="Hesabınızın şifresi"
            )
            
            st.markdown("")  # Boşluk için
            
            col_login, col_forgot = st.columns([2, 1])
            
            with col_login:
                submit = st.form_submit_button("🚀 Giriş Yap", type="primary", use_container_width=True)
            
            with col_forgot:
                forgot_button = st.form_submit_button("🔑 Şifremi Unuttum", use_container_width=True)
            
            if submit:
                if not username or not password:
                    st.error("❌ Lütfen kullanıcı adı ve şifrenizi girin!")
                else:
                    user, message = um.authenticate_user(username, password)
                    if user and message == "success":
                        # Giriş başarılı
                        st.session_state["authenticated"] = True
                        st.session_state["current_user"] = user
                        st.session_state["user_role"] = user["role"]
                        st.session_state["page"] = "session_manager"
                        
                        # Son giriş zamanını güncelle
                        um.update_last_login(user["user_id"])
                        
                        st.success(f"✅ Hoş geldiniz, {user['display_name']}!")
                        st.rerun()
                    else:
                        st.error(f"❌ {message}")
            
            if forgot_button:
                st.session_state["page"] = "forgot_password"
                st.rerun()
        
        st.markdown("---")
        
        # Kayıt ol butonu
        if st.button("📝 Henüz hesabınız yok mu? Kayıt olun", use_container_width=True):
            st.session_state["page"] = "register"
            st.rerun()

def show_register():
    """Kayıt sayfası"""
    # Başlık - merkezi ve güzel görünüm
    st.markdown("<div style='text-align: center;'>", unsafe_allow_html=True)
    st.title("📝 Kayıt Ol")
    st.caption("Yeni kullanıcı hesabı oluşturun")
    st.markdown("</div>", unsafe_allow_html=True)
    
    um = get_user_manager()
    
    # Kayıt formu - ortalanmış ve düzenli
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        with st.form("register_form"):
            st.markdown("### 📋 Kayıt Bilgileri")
            
            # Alan alanları alt alta
            username = st.text_input(
                "👤 Kullanıcı Adı",
                placeholder="örn: ahmet_yilmaz",
                help="Benzersiz bir kullanıcı adı seçin"
            )
            
            email = st.text_input(
                "📧 E-posta Adresi",
                placeholder="örn: ahmet@example.com",
                help="Geçerli bir e-posta adresi girin"
            )
            
            password = st.text_input(
                "🔒 Şifre",
                type="password",
                placeholder="Güvenli bir şifre seçin",
                help="En az 4 karakter olmalı"
            )
            
            password_confirm = st.text_input(
                "🔒 Şifre Tekrar",
                type="password",
                placeholder="Şifreyi tekrar girin",
                help="Aynı şifreyi tekrar girin"
            )
            
            role = st.selectbox(
                "🎭 İstenen Rol",
                options=["level1", "level2", "admin"],
                format_func=lambda x: {
                    "admin": "👑 Yönetici (Tüm formlar + yönetim yetkisi)",
                    "level1": "📝 Seviye 1 (Sadece Ek 1-2-3 formları)",
                    "level2": "📄 Seviye 2 (Ek 4, 6, 8, 9, 11, 15 formları)"
                }[x],
                help="Admin onayından sonra bu role sahip olacaksınız"
            )
            
            st.markdown("---")
            st.markdown("### 🔐 Güvenlik Sorusu (Şifre sıfırlama için)")
            
            # Önceden tanımlanmış güvenlik soruları
            security_questions = [
                "İlk evcil hayvanınızın adı neydi?",
                "Doğduğunuz şehir neresidir?",
                "En sevdiğiniz yemeğin adı nedir?",
                "İlkokul öğretmeninizin soyadı neydi?",
                "En sevdiğiniz renk nedir?",
                "Anne kızlık soyadı nedir?",
                "İlk işyerinizin adı neydi?",
                "En sevdiğiniz film karakteri kimdir?"
            ]
            
            selected_question = st.selectbox("Güvenlik sorusu seçin:", security_questions)
            security_answer = st.text_input("Güvenlik sorusu cevabı:", help="Bu cevabı şifrenizi unuttuğunuzda kullanacaksınız.")
            
            st.markdown("")  # Boşluk için
            
            submit = st.form_submit_button("🚀 Kayıt Ol", type="primary", use_container_width=True)
            
            if submit:
                # Validasyon kontrolleri
                if not username or not email or not password:
                    st.error("❌ Lütfen tüm alanları doldurun!")
                elif len(password) < 4:
                    st.error("❌ Şifre en az 4 karakter olmalı!")
                elif password != password_confirm:
                    st.error("❌ Şifreler eşleşmiyor!")
                elif "@" not in email or "." not in email:
                    st.error("❌ Geçerli bir e-posta adresi girin!")
                elif um.get_user_by_username(username):
                    st.error("❌ Bu kullanıcı adı zaten kullanılıyor!")
                elif not security_answer.strip():
                    st.error("❌ Güvenlik sorusu cevabı boş olamaz!")
                else:
                    # Kullanıcıyı kaydet
                    user = um.register_user(username, email, role, password, selected_question, security_answer)
                    if user:
                        st.success(f"✅ Kayıt başarılı! {username}")
                        st.info("⏳ **Hesabınız admin onayı bekliyor.** Admin onayladıktan sonra giriş yapabileceksiniz.")
                        st.balloons()
                        
                        # Session state'e başarılı kayıt durumunu işaretle
                        st.session_state["registration_success"] = True
                    else:
                        st.error("❌ Kayıt sırasında hata oluştu!")
    
        # Başarılı kayıt sonrası kontrol (form dışında)
        if st.session_state.get("registration_success", False):
            st.markdown("---")
            if st.button("🔙 Giriş Sayfasına Git", type="primary", use_container_width=True):
                st.session_state["registration_success"] = False  # Reset flag
                st.session_state["page"] = "login"
                st.rerun()
        else:
            st.markdown("---")
            
            if st.button("🔙 Zaten hesabınız var mı? Giriş yapın", use_container_width=True):
                st.session_state["page"] = "login"
                st.rerun()
            
            # Güvenlik bilgisi
            st.info("🔐 **Güvenlik Notu:** Tüm bilgileriniz güvenli olarak şifrelenerek saklanır.")

def show_forgot_password():
    """Şifremi unuttum sayfası"""
    # Başlık - merkezi ve güzel görünüm
    st.markdown("<div style='text-align: center;'>", unsafe_allow_html=True)
    st.title("🔑 Şifremi Unuttum")
    st.caption("Güvenlik sorunuzla şifrenizi sıfırlayın")
    st.markdown("</div>", unsafe_allow_html=True)
    
    um = get_user_manager()
    
    # Step tracking için session state kullan
    if 'forgot_step' not in st.session_state:
        st.session_state['forgot_step'] = 1
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        if st.session_state['forgot_step'] == 1:
            # Adım 1: Kullanıcı adı gir
            st.markdown("### 👤 Kullanıcı Adınızı Girin")
            
            with st.form("username_form"):
                username = st.text_input(
                    "👤 Kullanıcı Adı",
                    placeholder="Kullanıcı adınızı girin",
                    help="Kayıt olurken kullandığınız kullanıcı adı"
                )
                
                st.markdown("")
                
                col_continue, col_back = st.columns([1, 1])
                
                with col_continue:
                    continue_button = st.form_submit_button("Devam Et ➜", type="primary", use_container_width=True)
                
                with col_back:
                    back_button = st.form_submit_button("↩️ Giriş Sayfası", use_container_width=True)
                
                if continue_button:
                    if not username.strip():
                        st.error("❌ Lütfen kullanıcı adınızı girin!")
                    else:
                        # Kullanıcıyı ve güvenlik sorusunu kontrol et
                        security_question = um.get_security_question(username)
                        if security_question:
                            st.session_state['forgot_username'] = username
                            st.session_state['forgot_security_question'] = security_question
                            st.session_state['forgot_step'] = 2
                            st.rerun()
                        else:
                            st.error("❌ Kullanıcı bulunamadı veya güvenlik sorusu tanımlanmamış!")
                
                if back_button:
                    # Reset forgot state
                    for key in ['forgot_step', 'forgot_username', 'forgot_security_question']:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.session_state["page"] = "login"
                    st.rerun()
        
        elif st.session_state['forgot_step'] == 2:
            # Adım 2: Güvenlik sorusunu cevapla
            username = st.session_state.get('forgot_username', '')
            security_question = st.session_state.get('forgot_security_question', '')
            
            st.markdown("### 🔐 Güvenlik Sorusu")
            st.write(f"**Kullanıcı:** {username}")
            st.info(f"**Soru:** {security_question}")
            
            with st.form("security_form"):
                security_answer = st.text_input(
                    "🔑 Cevabınız",
                    placeholder="Güvenlik sorusu cevabınızı girin",
                    help="Kayıt olurken verdiğiniz cevabı girin"
                )
                
                st.markdown("")
                
                col_verify, col_back = st.columns([1, 1])
                
                with col_verify:
                    verify_button = st.form_submit_button("Doğrula ✓", type="primary", use_container_width=True)
                
                with col_back:
                    back_button = st.form_submit_button("↩️ Geri", use_container_width=True)
                
                if verify_button:
                    if not security_answer.strip():
                        st.error("❌ Lütfen güvenlik sorusu cevabınızı girin!")
                    else:
                        if um.verify_security_answer(username, security_answer):
                            st.session_state['forgot_step'] = 3
                            st.rerun()
                        else:
                            st.error("❌ Güvenlik sorusu cevabı yanlış!")
                
                if back_button:
                    st.session_state['forgot_step'] = 1
                    st.rerun()
        
        elif st.session_state['forgot_step'] == 3:
            # Adım 3: Yeni şifre belirle
            username = st.session_state.get('forgot_username', '')
            
            st.markdown("### 🔒 Yeni Şifre Belirleyin")
            st.success("✅ Güvenlik sorusu doğrulandı!")
            st.write(f"**Kullanıcı:** {username}")
            
            with st.form("password_reset_form"):
                new_password = st.text_input(
                    "🔒 Yeni Şifre",
                    type="password",
                    placeholder="Yeni şifrenizi girin",
                    help="En az 4 karakter olmalı"
                )
                
                confirm_password = st.text_input(
                    "🔒 Yeni Şifre Tekrar",
                    type="password",
                    placeholder="Yeni şifrenizi tekrar girin",
                    help="Aynı şifreyi tekrar girin"
                )
                
                st.markdown("")
                
                col_reset, col_cancel = st.columns([1, 1])
                
                with col_reset:
                    reset_button = st.form_submit_button("🔄 Şifreyi Sıfırla", type="primary", use_container_width=True)
                
                with col_cancel:
                    cancel_button = st.form_submit_button("❌ İptal", use_container_width=True)
                
                if reset_button:
                    if not new_password or not confirm_password:
                        st.error("❌ Lütfen tüm alanları doldurun!")
                    elif len(new_password) < 4:
                        st.error("❌ Şifre en az 4 karakter olmalı!")
                    elif new_password != confirm_password:
                        st.error("❌ Şifreler eşleşmiyor!")
                    else:
                        success, message = um.reset_password(username, new_password)
                        if success:
                            st.success("🎉 Şifreniz başarıyla sıfırlandı!")
                            st.info("Artık yeni şifrenizle giriş yapabilirsiniz.")
                            st.balloons()
                            
                            # Reset all forgot password states
                            for key in ['forgot_step', 'forgot_username', 'forgot_security_question']:
                                if key in st.session_state:
                                    del st.session_state[key]
                            
                            # Kısa bir bekleme sonrası giriş sayfasına yönlendir
                            import time
                            time.sleep(2)
                            st.session_state["page"] = "login"
                            st.rerun()
                        else:
                            st.error(f"❌ {message}")
                
                if cancel_button:
                    # Reset forgot state
                    for key in ['forgot_step', 'forgot_username', 'forgot_security_question']:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.session_state["page"] = "login"
                    st.rerun()

def show_admin_approvals():
    """Admin kullanıcı onay sayfası"""
    current_user = st.session_state.get("current_user")
    
    # Sadece "admin" kullanıcısı kontrolü
    if not current_user or current_user.get("username") != "admin":
        st.error("❌ Bu sayfaya erişim yetkiniz yok!")
        st.warning("🔒 Bu sayfa sadece sistem yöneticisi için erişilebilirdir.")
        if st.button("🏠 Ana Sayfaya Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return
    
    # Header
    col_title, col_back = st.columns([3, 1])
    with col_title:
        st.title("👑 Kullanıcı Onay Merkezi")
        st.caption("Bekleyen kullanıcı kayıtlarını onaylayın veya reddedin")
    with col_back:
        if st.button("🏠 Ana Sayfa"):
            st.session_state["page"] = "session_manager"
            st.rerun()
    
    st.markdown("---")
    
    um = get_user_manager()
    pending_users = um.get_pending_users()
    all_users = um.get_all_users()
    approved_users = [u for u in all_users if u.get("status") == "approved"]
    
    # İstatistikler
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("⏳ Bekleyen", len(pending_users))
    with col2:
        st.metric("✅ Onaylı", len(approved_users))
    with col3:
        st.metric("👥 Toplam", len(all_users))
    
    st.markdown("---")
    
    if not pending_users:
        st.info("🎉 **Harika!** Şu anda bekleyen kullanıcı onayı yok.")
    
    # Onaylı kullanıcıları göster ve yönet
    if approved_users:
        st.markdown("---")
        st.subheader("✅ Kayıtlı Kullanıcılar")
        
        for user in approved_users:
            with st.container():
                col_info, col_role, col_actions = st.columns([2, 1, 1])
                
                with col_info:
                    role_icon = "👑" if user["role"] == "admin" else "📝" if user["role"] == "level1" else "📄"
                    role_name = "Yönetici" if user["role"] == "admin" else "Seviye 1" if user["role"] == "level1" else "Seviye 2"
                    
                    st.write(f"**{role_icon} {user['display_name']}** ({user['username']})")
                    st.caption(f"{role_name} • Kayıt: {user['created_at'][:10]}")
                
                with col_role:
                    # Admin kullanıcısının rolü değiştirilemez
                    if user.get("username") != "admin":
                        current_role = user["role"]
                        role_options = ["level1", "level2", "admin"]
                        role_labels = {
                            "level1": "📝 Seviye 1",
                            "level2": "📄 Seviye 2", 
                            "admin": "👑 Yönetici"
                        }
                        
                        new_role = st.selectbox(
                            "Rol:",
                            options=role_options,
                            index=role_options.index(current_role),
                            format_func=lambda x: role_labels[x],
                            key=f"role_{user['user_id']}"
                        )
                        
                        # Rol değiştirme butonu
                        if new_role != current_role:
                            if st.button("🔄 Değiştir", key=f"change_role_{user['user_id']}", use_container_width=True):
                                st.session_state[f"confirm_role_change_{user['user_id']}"] = new_role
                                st.rerun()
                        
                        # Rol değiştirme onayı
                        if st.session_state.get(f"confirm_role_change_{user['user_id']}"):
                            new_role_confirm = st.session_state[f"confirm_role_change_{user['user_id']}"]
                            role_name_new = role_labels[new_role_confirm]
                            
                            st.warning(f"⚠️ **{user['display_name']}** kullanıcısının rolünü **{role_name_new}** olarak değiştirmek istediğinizden emin misiniz?")
                            col_yes, col_no = st.columns(2)
                            
                            with col_yes:
                                if st.button("✅ Evet", key=f"confirm_yes_role_{user['user_id']}"):
                                    if um.change_user_role(user['user_id'], new_role_confirm, current_user['user_id']):
                                        st.success(f"🔄 {user['display_name']} rolü güncellendi!")
                                        del st.session_state[f"confirm_role_change_{user['user_id']}"]
                                        st.rerun()
                                    else:
                                        st.error("❌ Rol değiştirme hatası!")
                            
                            with col_no:
                                if st.button("❌ İptal", key=f"confirm_no_role_{user['user_id']}"):
                                    del st.session_state[f"confirm_role_change_{user['user_id']}"]
                                    st.rerun()
                    else:
                        st.write("🔒 **Korumalı**")
                        st.caption("Admin rolü")
                
                with col_actions:
                    # Admin kullanıcısını silemez
                    if user.get("username") != "admin":
                        if st.button("🗑️", 
                                   key=f"delete_{user['user_id']}", 
                                   help="Kullanıcıyı sil",
                                   use_container_width=True):
                            # Onay modalı için session state kullan
                            st.session_state[f"confirm_delete_user_{user['user_id']}"] = True
                            st.rerun()
                        
                        # Silme onayı
                        if st.session_state.get(f"confirm_delete_user_{user['user_id']}", False):
                            st.warning(f"⚠️ **{user['display_name']}** kullanıcısını silmek istediğinizden emin misiniz?")
                            col_yes, col_no = st.columns(2)
                            
                            with col_yes:
                                if st.button("✅ Evet", key=f"confirm_yes_user_{user['user_id']}"):
                                    if um.delete_user(user['user_id']):
                                        st.success(f"🗑️ {user['display_name']} silindi!")
                                        del st.session_state[f"confirm_delete_user_{user['user_id']}"]
                                        st.rerun()
                                    else:
                                        st.error("❌ Silme hatası!")
                            
                            with col_no:
                                if st.button("❌ İptal", key=f"confirm_no_user_{user['user_id']}"):
                                    del st.session_state[f"confirm_delete_user_{user['user_id']}"]
                                    st.rerun()
                    else:
                        st.write("🔒")
                        st.caption("Korumalı")
                
                st.markdown("---")
    
    # Bekleyen kullanıcılar bölümü
    if pending_users:
        st.subheader(f"⏳ Onay Bekleyen Kullanıcılar ({len(pending_users)})")
        
        for i, pending_user in enumerate(pending_users):
            with st.container():
                st.markdown(f"### 👤 {pending_user['display_name']}")
                
                col_info, col_actions = st.columns([2, 1])
                
                with col_info:
                    role_icon = "👑" if pending_user["role"] == "admin" else "📝" if pending_user["role"] == "level1" else "📄"
                    role_name = "Yönetici" if pending_user["role"] == "admin" else "Seviye 1" if pending_user["role"] == "level1" else "Seviye 2"
                    
                    st.write(f"**👤 Kullanıcı Adı:** {pending_user['username']}")
                    st.write(f"**📧 E-posta:** {pending_user.get('email', 'Belirtilmemiş')}")
                    st.write(f"**🎭 İstenen Rol:** {role_icon} {role_name}")
                    st.write(f"**📅 Kayıt Tarihi:** {pending_user['created_at'][:19].replace('T', ' ')}")
                    
                    # Rol açıklaması
                    if pending_user["role"] == "admin":
                        st.warning("⚠️ **Dikkat:** Yönetici rolü isteniyor!")
                    elif pending_user["role"] == "level1":
                        st.info("📝 Sadece Ek 1-2-3 formlarına erişim")
                    else:
                        st.info("📄 Ek 4, 6, 8, 9, 11, 15 formlarına erişim")
                
                with col_actions:
                    st.write("**Karar Verin:**")
                    
                    col_approve, col_reject = st.columns(2)
                    
                    with col_approve:
                        if st.button("✅ Onayla", 
                                   key=f"approve_{pending_user['user_id']}", 
                                   type="primary",
                                   use_container_width=True):
                            if um.approve_user(pending_user['user_id'], current_user['user_id']):
                                st.success(f"✅ {pending_user['display_name']} onaylandı!")
                                st.rerun()
                            else:
                                st.error("❌ Onay hatası!")
                    
                    with col_reject:
                        if st.button("❌ Reddet", 
                                   key=f"reject_{pending_user['user_id']}", 
                                   use_container_width=True):
                            if um.reject_user(pending_user['user_id'], current_user['user_id']):
                                st.success(f"🗑️ {pending_user['display_name']} reddedildi ve silindi!")
                                st.rerun()
                            else:
                                st.error("❌ Red hatası!")
                
                if i < len(pending_users) - 1:  # Son eleman değilse ayraç ekle
                    st.markdown("---")

# ================== Ana Uygulama ==================

def main():
    # UTF-8 encoding'i kontrol et ve ayarla
    ensure_utf8_encoding()
    
    st.set_page_config(
        page_title="🎯 Sesli Belge Doldurma Sistemi", 
        page_icon="🎯", 
        layout="wide"
    )

    # Authentication state initialization
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    if "current_user" not in st.session_state:
        st.session_state["current_user"] = None
    if "user_role" not in st.session_state:
        st.session_state["user_role"] = None
    
    # Session state initialization
    if "page" not in st.session_state:
        # Start with login page if not authenticated
        st.session_state["page"] = "login" if not st.session_state["authenticated"] else "session_manager"
    if "current_session_id" not in st.session_state:
        st.session_state["current_session_id"] = None
    if "current_session_name" not in st.session_state:
        st.session_state["current_session_name"] = ""
    if "api_key" not in st.session_state:
        st.session_state["api_key"] = ""
    # Form selection related state
    if "selected_form_group" not in st.session_state:
        st.session_state["selected_form_group"] = None  # Örn: "Ek 1-2-3", "Ek 4", "Ek 6", "Ek 8"
    if "form_group_applied" not in st.session_state:
        st.session_state["form_group_applied"] = None
    if "templates_initialized_for" not in st.session_state:
        st.session_state["templates_initialized_for"] = None
    if "selected_templates" not in st.session_state:
        st.session_state["selected_templates"] = []

    # İlk admin kullanıcısını oluştur (eğer hiç kullanıcı yoksa) - sessizce
    um = get_user_manager()
    um.create_initial_admin_if_needed()

    # Page routing
    if not st.session_state["authenticated"]:
        # Authentication required pages
        if st.session_state["page"] == "login":
            show_login()
        elif st.session_state["page"] == "register":
            show_register()
        elif st.session_state["page"] == "forgot_password":
            show_forgot_password()
        else:
            st.session_state["page"] = "login"
            st.rerun()
    else:
        # Authenticated pages
        if st.session_state["page"] == "session_manager":
            show_session_manager()
        elif st.session_state["page"] == "form_selector":
            show_form_selector()
        elif st.session_state["page"] == "voice_app":
            show_voice_app()
        elif st.session_state["page"] == "admin_approvals":
            show_admin_approvals()
        elif st.session_state["page"] == "feedback_panel":
            show_feedback_panel()
        else:
            st.session_state["page"] = "session_manager"
            st.rerun()

def show_session_manager():
    """Session yönetim arayüzü"""
    # Kullanıcı bilgisi ve çıkış butonu
    current_user = st.session_state.get("current_user")
    if current_user:
        col_title, col_user = st.columns([3, 1])
        with col_title:
            st.title("🎯 Sesli Belge Doldurma Sistemi")
            st.caption("Ses girdi ile Word şablonlarını otomatik dolduran akıllı sistem")
        with col_user:
            role_icon = "👑" if current_user["role"] == "admin" else "📝" if current_user["role"] == "level1" else "📄"
            st.write(f"{role_icon} **{current_user['display_name']}**")
            st.caption(f"Rol: {current_user['role']}")
            if st.button("🚪 Çıkış Yap"):
                st.session_state["authenticated"] = False
                st.session_state["current_user"] = None
                st.session_state["user_role"] = None
                st.session_state["page"] = "login"
                st.rerun()
    else:
        st.title("🎯 Sesli Belge Doldurma Sistemi")
        st.caption("Ses girdi ile Word şablonlarını otomatik dolduran akıllı sistem")
    
    sm = get_local_session_manager()
    fbm = get_feedback_manager()
    
    # Arama çubuğu (yalnızca öğrenci adı veya numarasına göre)
    search_term = st.text_input("🔍 Öğrenci Ara", placeholder="Öğrenci adı veya öğrenci numarası...")
    
    # Session listesi
    sessions = sm.get_all_sessions()
    
    # Arama filtresi (yalnızca ad veya numara)
    if search_term:
        q = search_term.lower().strip()
        filtered_sessions = []
        for session in sessions:
            student_no, student_name = extract_student_info(session)
            match_no = bool(student_no and q in student_no.lower())
            match_name = bool(student_name and q in student_name.lower())
            if match_no or match_name:
                filtered_sessions.append(session)
        sessions = filtered_sessions
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.subheader("📁 Öğrenci Session'ları")
        
        if not sessions:
            if search_term:
                st.info("Arama kriterinize uygun session bulunamadı.")
            else:
                st.info("Henüz session oluşturulmamış. 'Yeni Session' butonuna tıklayın.")
        else:
            for session in sessions:
                student_no, student_name = extract_student_info(session)
                
                if student_no and student_name:
                    display_title = f"👤 {student_no} - {student_name}"
                else:
                    # Standart dışı isim varsa da aynı formatla göster
                    fallback_name = session.get('session_name', '')
                    display_title = f"👤 {fallback_name}"
                
                with st.expander(display_title, expanded=False):
                    col_info, col_actions = st.columns([2, 1])
                    
                    with col_info:
                        st.write(f"**Oluşturma:** {session['created_date'][:10]}")
                        if student_no:
                            st.write(f"**Öğrenci No:** {student_no}")
                        if student_name:
                            st.write(f"**Öğrenci Adı:** {student_name}")
                        
                        data_count = len([v for v in session.get('extracted_data', {}).values() if v])
                        st.write(f"**Dolu Alanlar:** {data_count}")
                    
                    with col_actions:
                        if st.button(f"🚀 Aç", key=f"open_{session['session_id']}"):
                            st.session_state["current_session_id"] = session['session_id']
                            st.session_state["current_session_name"] = session['session_name']
                            # Form seçim sayfasına yönlendir ve önceki seçimleri sıfırla
                            st.session_state["selected_form_group"] = None
                            st.session_state["form_group_applied"] = None
                            st.session_state["templates_initialized_for"] = None
                            st.session_state["selected_templates"] = []
                            st.session_state["page"] = "form_selector"
                            st.rerun()
                        
                        # Sadece "admin" kullanıcısı session silebilir
                        if current_user and current_user.get("username") == "admin":
                            if st.button(f"🗑️ Sil", key=f"delete_{session['session_id']}"):
                                st.session_state[f"confirm_delete_{session['session_id']}"] = True
                                st.rerun()
                            
                            if st.session_state.get(f"confirm_delete_{session['session_id']}", False):
                                st.warning("⚠️ Silmek istediğinizden emin misiniz?")
                                col_yes, col_no = st.columns(2)
                                
                                with col_yes:
                                    if st.button("✅ Evet", key=f"confirm_yes_{session['session_id']}"):
                                        if sm.delete_session(session['session_id']):
                                            st.success("Session silindi!")
                                            del st.session_state[f"confirm_delete_{session['session_id']}"]
                                            st.rerun()
                                
                                with col_no:
                                    if st.button("❌ İptal", key=f"confirm_no_{session['session_id']}"):
                                        del st.session_state[f"confirm_delete_{session['session_id']}"]
                                        st.rerun()
    
    with col2:
        # Sadece "admin" kullanıcısı için paneller
        if current_user and current_user.get("username") == "admin":
            um = get_user_manager()
            pending_users = um.get_pending_users()
            pending_feedbacks = get_feedback_manager().get_pending_count()
            
            st.subheader("👑 Admin Panel")
            if pending_users:
                st.write(f"⏳ **{len(pending_users)} kullanıcı onay bekliyor**")
                if st.button("🔍 Kullanıcı Onaylarını Yönet", type="primary", use_container_width=True):
                    st.session_state["page"] = "admin_approvals"
                    st.rerun()
            else:
                st.write("✅ **Bekleyen onay yok**")
                if st.button("👥 Kullanıcı Yönetimi", use_container_width=True):
                    st.session_state["page"] = "admin_approvals"
                    st.rerun()
            
            st.markdown("---")
            st.write(f"💬 Bekleyen geri bildirim: **{pending_feedbacks}**")
            if st.button("💬 Geri Bildirim Paneli", use_container_width=True):
                st.session_state["page"] = "feedback_panel"
                st.rerun()
        
        # Geri Bildirim (Yeni Session bölümü gibi sağ sütunda)
        st.subheader("💬 Geri Bildirim")
        st.write("Hata, istek veya önerinizi iletin.")
        feedback_text_right = st.text_area("Mesajınız", key="feedback_text_main", placeholder="Örn: Ek 6 şablonunda bir alan çalışmıyor...", height=120)
        if st.button("📨 Gönder", key="send_feedback_main", use_container_width=True):
            if not feedback_text_right or not feedback_text_right.strip():
                st.warning("Lütfen bir mesaj yazın.")
            else:
                fb_id = fbm.submit_feedback(current_user or {}, feedback_text_right)
                if fb_id:
                    st.success("Teşekkürler! Geri bildiriminiz admin'e iletildi.")
                else:
                    st.error("Geri bildirim kaydedilemedi.")
        
        st.subheader("🚀 Yeni Session")
        st.write("Yeni bir öğrenci için session başlatın.")
        
        if st.button("📝 Yeni Session Başlat", type="primary", use_container_width=True):
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            session_name = f"Yeni Session - {timestamp}"
            
            # Kullanıcı bilgilerini session'a ekle
            session_id = sm.create_session(session_name, current_user)
            if session_id:
                st.success("Yeni session başlatıldı!")
                st.session_state["current_session_id"] = session_id
                st.session_state["current_session_name"] = session_name
                # Yeni session sonrası form seçim ekranına git
                st.session_state["page"] = "form_selector"
                # Yeni session'da transkript ve mapping boşlansın
                st.session_state["current_transcript"] = ""
                st.session_state["transcript_loaded_for"] = session_id
                st.session_state["current_mapping"] = {}
                st.session_state["mapping_loaded_for"] = session_id
                st.session_state["results"] = None
                # Form seçim state'leri
                st.session_state["selected_form_group"] = None
                st.session_state["form_group_applied"] = None
                st.session_state["templates_initialized_for"] = None
                st.session_state["selected_templates"] = []
                st.rerun()
            else:
                st.error("Session oluşturulamadı!")
        
        st.info("💡 **İpucu:** Session başlattıktan sonra öğrenci bilgilerini sesli girdi ile kaydedin.")

def show_feedback_panel():
    """Admin geri bildirim yönetim sayfası"""
    current_user = st.session_state.get("current_user")
    if not current_user or current_user.get("username") != "admin":
        st.error("❌ Bu sayfaya erişim yetkiniz yok!")
        if st.button("🏠 Ana Sayfaya Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return

    fbm = get_feedback_manager()
    feedbacks = fbm.get_all_feedbacks()

    col_title, col_back = st.columns([3, 1])
    with col_title:
        st.title("💬 Geri Bildirim Paneli")
        st.caption("Kullanıcı geri bildirimlerini inceleyin ve durum atayın")
    with col_back:
        if st.button("🏠 Ana Sayfa"):
            st.session_state["page"] = "session_manager"
            st.rerun()

    st.markdown("---")

    if not feedbacks:
        st.info("Şu anda geri bildirim yok.")
        return

    status_labels = {
        "pending": "⏳ Beklemede",
        "in_progress": "🔧 İşlemde",
        "resolved": "✅ Çözüldü",
    }

    for fb in feedbacks:
        with st.container():
            col_info, col_actions = st.columns([3, 1])
            with col_info:
                st.write(f"**📅 {fb.get('created_at','')[:19].replace('T',' ')}** • {status_labels.get(fb.get('status'), fb.get('status'))}")
                user = fb.get("submitted_by", {}) or {}
                st.write(f"Gönderen: {user.get('display_name') or user.get('username') or 'Bilinmiyor'} ({user.get('role') or '-'})")
                st.write(f"Mesaj:")
                st.code(fb.get("message", ""))

            with col_actions:
                new_status = st.selectbox(
                    "Durum",
                    options=["pending", "in_progress", "resolved"],
                    index=["pending", "in_progress", "resolved"].index(fb.get("status", "pending")),
                    format_func=lambda x: status_labels.get(x, x),
                    key=f"fb_status_{fb['feedback_id']}"
                )
                if st.button("💾 Kaydet", key=f"fb_save_{fb['feedback_id']}", use_container_width=True):
                    if fbm.set_status(fb['feedback_id'], new_status):
                        st.success("Durum güncellendi")
                        st.rerun()
                    else:
                        st.error("Güncelleme hatası")
                if st.button("🗑️ Sil", key=f"fb_delete_{fb['feedback_id']}", use_container_width=True):
                    if fbm.delete_feedback(fb['feedback_id']):
                        st.success("Geri bildirim silindi")
                        st.rerun()
                    else:
                        st.error("Silme hatası")

            st.markdown("---")

def show_form_selector():
    """Form (Ek) seçim ekranı"""
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    current_user = st.session_state.get("current_user")
    
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return

    st.title("🧩 Hangi Ek doldurulacak?")
    st.caption(f"{current_session_name}")

    # Kullanıcı rolüne göre form seçeneklerini filtrele
    if current_user:
        um = get_user_manager()
        permissions = um.get_user_permissions(current_user["role"])
        available_forms = permissions["available_forms"]
        
        # Kullanıcı rolü bilgisini göster
        role_icon = "👑" if current_user["role"] == "admin" else "📝" if current_user["role"] == "level1" else "📄"
        st.info(f"{role_icon} **{current_user['display_name']}** - Size açık formlar gösteriliyor")
    else:
        # Fallback: Tüm formları göster
        available_forms = ["Ek 1-2-3", "Ek 4", "Ek 6", "Ek 8", "Ek 9", "Ek 11", "Ek 15"]

    # Seçili form grubunu kontrol et
    current_selected = st.session_state.get("selected_form_group")
    default_idx = 0
    if current_selected and current_selected in available_forms:
        default_idx = available_forms.index(current_selected)
    
    selected = st.radio("Form seti", options=available_forms, index=default_idx, horizontal=True)

    col_go, col_back = st.columns([1, 1])
    with col_go:
        if st.button("Devam et ➜", type="primary", use_container_width=True):
            st.session_state["selected_form_group"] = selected
            st.session_state["form_group_applied"] = None  # Voice sayfasında yeniden uygula
            st.session_state["templates_initialized_for"] = None
            st.session_state["selected_templates"] = []
            st.session_state["page"] = "voice_app"
            st.rerun()
    with col_back:
        if st.button("↩️ Session listesine dön", use_container_width=True):
            st.session_state["page"] = "session_manager"
            st.rerun()

def show_voice_app():
    """Ana ses uygulama arayüzü"""
    current_session_id = st.session_state.get("current_session_id")
    current_session_name = st.session_state.get("current_session_name", "Bilinmeyen Session")
    current_user = st.session_state.get("current_user")
    
    if not current_session_id:
        st.error("Session bilgisi bulunamadı!")
        if st.button("🏠 Session Yöneticisine Dön"):
            st.session_state["page"] = "session_manager"
            st.rerun()
        return
    
    # Session verilerini yükle
    sm = get_local_session_manager()
    session_data = sm.get_session(current_session_id)
    
    if not session_data:
        st.error("Session verisi yüklenemedi!")
        return
    
    # Session state'leri initialize et (mapping ve transcript session bazlı yüklensin)
    if st.session_state.get("mapping_loaded_for") != current_session_id:
        st.session_state["current_mapping"] = {}
        st.session_state["mapping_loaded_for"] = current_session_id
        st.session_state["results"] = None
    # Transkript, session bazlı yüklensin (diğer session'dan taşınmasın)
    if st.session_state.get("transcript_loaded_for") != current_session_id:
        st.session_state["current_transcript"] = session_data.get('transcript', "")
        st.session_state["transcript_loaded_for"] = current_session_id
    
    # Header
    col_title, col_actions, col_user = st.columns([3, 1.5, 1.5])
    with col_title:
        st.title(f"🎯 {current_session_name}")
        st.caption(f"Session ID: {current_session_id[:12]}...")
        # Aktif form seti bilgisini kullanıcıya göstermeyelim
    
    with col_actions:
        if st.button("🧩 Form setini değiştir"):
            st.session_state["page"] = "form_selector"
            st.rerun()
        if st.button("🏠 Session listesi"):
            st.session_state["page"] = "session_manager"
            st.rerun()
    
    with col_user:
        if current_user:
            role_icon = "👑" if current_user["role"] == "admin" else "📝" if current_user["role"] == "level1" else "📄"
            st.write(f"{role_icon} **{current_user['display_name']}**")
            st.caption(f"Rol: {current_user['role']}")
            if st.button("🚪 Çıkış"):
                st.session_state["authenticated"] = False
                st.session_state["current_user"] = None
                st.session_state["user_role"] = None
                st.session_state["page"] = "login"
                st.rerun()
    
    st.markdown("---")
    
    # API Key
    col_api1, col_api2 = st.columns([3, 1])
    with col_api1:
        api_key_input = st.text_input(
            "🔑 OpenAI API Key",
            value=st.session_state.get("api_key", ""),
            type="password",
            help="Whisper ve AI analizi için gerekli"
        )
    with col_api2:
        if st.checkbox("Hatırla", value=bool(st.session_state.get("api_key"))):
            st.session_state["api_key"] = api_key_input
        else:
            st.session_state["api_key"] = ""
    
    st.markdown("---")
    
    # Şablon seçimi (arka planda otomatik)
    
    default_dir = os.path.join(os.getcwd(), "templates")
    selected_names = []
    available = []
    
    try:
        if os.path.isdir(default_dir):
            available = sorted([f for f in os.listdir(default_dir) if f.lower().endswith(".docx")])
            if available:
                # Form setine göre ön seçim hazırla (sadece ilk girişte uygula veya grup değiştiyse)
                group = st.session_state.get("selected_form_group")
                def _match_group_files(group_label: Optional[str], files: List[str]) -> List[str]:
                    if not group_label:
                        return []
                    prefixes_map = {
                        "Ek 1-2-3": ["Ek-1-", "Ek-2-", "Ek-3-"],
                        "Ek 4": ["Ek-4 "],
                        "Ek 6": ["Ek-6 "],
                        "Ek 8": ["Ek-8 "],
                        "Ek 9": ["Ek-9 "],
                        "Ek 11": ["Ek-11 "],
                        "Ek 15": ["Ek-15 "],
                    }
                    prefixes = prefixes_map.get(group_label, [])
                    matched_files = []
                    for f in files:
                        for pfx in prefixes:
                            if f.startswith(pfx):
                                matched_files.append(f)
                                break  # Bir dosya birden fazla prefix'e uymayacak
                    return matched_files

                should_apply_preselection = (
                    st.session_state.get("templates_initialized_for") != current_session_id or
                    st.session_state.get("form_group_applied") != group
                )
                if should_apply_preselection:
                    preselected = _match_group_files(group, available)
                    st.session_state["selected_templates"] = preselected
                    st.session_state["templates_initialized_for"] = current_session_id
                    st.session_state["form_group_applied"] = group
            else:
                # Klasörde .docx bulunmuyorsa sessiz geç; analiz adımında uyarılacak
                pass
        else:
            # Templates klasörü yoksa sessiz geç; analiz adımında uyarılacak
            pass
    except Exception as e:
        st.error(f"Templates klasörü okunamadı: {e}")
    
    template_items = []
    # Session state'den güncel seçimi al
    current_selected = st.session_state.get("selected_templates", [])
    for name in current_selected:
        try:
            full = os.path.join(default_dir, name)
            with open(full, "rb") as fh:
                template_items.append((name, fh.read()))
        except Exception as e:
            st.error(f"{name} okunamadı: {e}")
    
    # Placeholder'ları topla
    union_placeholders = set()
    if template_items:
        for name, data in template_items:
            try:
                placeholders, _ = extract_placeholders_from_docx_bytes(data)
                union_placeholders |= placeholders
            except Exception as e:
                st.error(f"{name} analiz edilemedi: {e}")
    
    st.markdown("---")
    
    # Ses kaydı bölümü
    st.subheader("🎤 Ses Kaydı ve Analiz")
    
    col_mic, col_btn = st.columns([3, 1])
    
    with col_mic:
        # Her zaman mikrofonu göster (genel transcript için)
        audio_bytes = render_audio_recorder_ui()
        # Ek 15 için ek olarak uzun metin alanını göster (transkripti değiştirmez)
        special_text_input = None
        if st.session_state.get("selected_form_group") == "Ek 15":
            special_text_input = st.text_area(
                "📝 Ek 15 İçerik (uzun metin)",
                value="",
                height=180,
                help="Bu metin yalnızca Ek 15'in 4 özel alanını doldurmak için kullanılır. Genel transkripti değiştirmez."
            )
    
    with col_btn:
        if st.button("🧠 Analiz Et", use_container_width=True, type="primary"):
            effective_key = (api_key_input or st.session_state.get("api_key", "")).strip()
            
            if not template_items:
                st.warning("Önce şablon seçin.")
                return
            if not union_placeholders:
                st.warning("Şablonlarda placeholder bulunamadı.")
                return
            if not effective_key:
                st.warning("OpenAI API anahtarı girin.")
                return

            existing_transcript = (st.session_state.get("current_transcript", "") or "").strip()
            merged_transcript = ""

            merged_transcript = existing_transcript
            if audio_bytes:
                with st.spinner("Ses metne çevriliyor..."):
                    text = transcribe_audio_bytes(audio_bytes, effective_key)
                if not text:
                    st.error("Ses metne çevrilemedi.")
                    return
                merged_transcript = (existing_transcript + " " + text.strip()).strip() if existing_transcript else text.strip()
                st.session_state["current_transcript"] = merged_transcript
                sm.update_session_transcript(current_session_id, merged_transcript)

            with st.spinner("Bilgiler çıkarılıyor..."):
                ctx = aggregate_contexts_across_templates(template_items, union_placeholders)
                # Genel ve özel (Ek 15) çıkarımları ayrı çalıştır ve birleştir
                suggested: Dict[str, str] = {}
                selected_group = st.session_state.get("selected_form_group")

                # Ek 15 özel seti
                ek15_conf = SPECIAL_FORMS.get("Ek 15", {}) if selected_group == "Ek 15" else {}
                ek15_set: Set[str] = set(ek15_conf.get("expected_placeholders", []) or [])

                # 1) Genel çıkarım: özel olmayan placeholder'lar (veya Ek 15 değilse tümü)
                general_placeholders: Set[str] = set(union_placeholders)
                if ek15_set:
                    general_placeholders = set(ph for ph in union_placeholders if ph not in ek15_set)
                if general_placeholders:
                    general_suggested = infer_placeholder_values(
                        merged_transcript,
                        general_placeholders,
                        ctx,
                        effective_key,
                    )
                    suggested.update(general_suggested or {})

                # 2) Ek 15 çıkarımı: sadece Ek 15 alanları, özel talimat ve gerekirse özel metin
                if ek15_set:
                    special_text = (special_text_input or "").strip()
                    if special_text:
                        ph_expl = ek15_conf.get("placeholder_explanations") or {}
                        extra_instr = ek15_conf.get("custom_instructions") or None
                        ek15_suggested = infer_placeholder_values(
                            special_text,
                            ek15_set,
                            ctx,
                            effective_key,
                            extra_instructions=extra_instr,
                            placeholder_explanations=ph_expl,
                        )
                        # Özel alanlar genel sonuçların üzerine yazsın
                        suggested.update(ek15_suggested or {})
                
                # Mevcut verilerle birleştir
                existing_data = st.session_state.get("current_mapping", {})
                conflicts = detect_conflicts(existing_data, suggested)
                
                if conflicts:
                    st.warning(f"⚠️ {len(conflicts)} çakışma tespit edildi: {', '.join(conflicts)}")
                
                merged_data = merge_extracted_data(existing_data, suggested)
                st.session_state["current_mapping"] = merged_data
                
                # Session'a kaydet
                try:
                    if sm.update_session_data(current_session_id, suggested, merge=True):
                        filled_count = len([v for v in suggested.values() if v.strip()])
                        st.success(f"✅ {filled_count} yeni bilgi eklendi ve kaydedildi!")
                        
                        # Session ismini güncelle
                        updated_session = sm.get_session(current_session_id)
                        if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                            st.session_state["current_session_name"] = updated_session['session_name']
                            st.info("📝 Session ismi güncellendi!")
                except Exception as e:
                    st.warning(f"Veriler çıkarıldı ama kaydetme sırasında hata: {e}")
                
                st.rerun()
    
    # Transkript gösterimi
    if st.session_state.get("current_transcript"):
        col_transcript, col_clear = st.columns([4, 1])
        
        with col_transcript:
            st.text_area(
                "📜 Birleşik Transkript",
                value=st.session_state.get("current_transcript", ""),
                height=120,
                disabled=True,
                help="Bu transkript session bazında saklanır ve tüm Ek formlarında kullanılabilir"
            )
        
        with col_clear:
            st.write("")
            if st.button("🗑️ Temizle"):
                st.session_state["confirm_clear_transcript"] = True
                st.rerun()

        # Temizleme onayı
        if st.session_state.get("confirm_clear_transcript", False):
            st.warning("⚠️ Birleşik transkripti silmek istediğinizden emin misiniz?")
            col_yes, col_no = st.columns(2)
            with col_yes:
                if st.button("✅ Evet", key="confirm_yes_clear_transcript"):
                    st.session_state["current_transcript"] = ""
                    sm.update_session_transcript(current_session_id, "")
                    st.session_state.pop("confirm_clear_transcript", None)
                    st.rerun()
            with col_no:
                if st.button("❌ İptal", key="confirm_no_clear_transcript"):
                    st.session_state.pop("confirm_clear_transcript", None)
                    st.rerun()
    
    # Placeholder değerleri
    if union_placeholders:
        st.markdown("---")
        st.subheader("✏️ Bilgi Düzenleme")
        
        col_apply, col_clear = st.columns([2, 1])
        with col_apply:
            if st.button("🔄 Session Verilerini Uygula"):
                session_data = sm.get_session(current_session_id)
                if session_data and session_data.get('extracted_data'):
                    current_mapping = st.session_state.get("current_mapping", {})
                    applied_count = 0
                    
                    for ph in union_placeholders:
                        if ph in session_data['extracted_data'] and session_data['extracted_data'][ph]:
                            if ph not in current_mapping or not current_mapping.get(ph):
                                current_mapping[ph] = session_data['extracted_data'][ph]
                                applied_count += 1
                    
                    st.session_state["current_mapping"] = current_mapping
                    if applied_count > 0:
                        st.success(f"✅ {applied_count} alan dolduruldu!")
                        st.rerun()
        
        with col_clear:
            if st.button("🧹 Temizle"):
                st.session_state["current_mapping"] = {}
                st.rerun()
        
        # Placeholder düzenleme
        edit_cols = st.columns(2)
        for idx, ph in enumerate(sorted(list(union_placeholders))):
            with edit_cols[idx % 2]:
                display_name = format_placeholder_label(ph)
                st.markdown(f"**{display_name}**")
                
                cur_val = st.session_state.get("current_mapping", {}).get(ph, "")
                new_val = st.text_input(
                    "Değer", 
                    value=cur_val, 
                    key=f"edit_{idx}_{ph}",
                    placeholder="Değer girin...",
                    label_visibility="collapsed"
                )
                
                if new_val != cur_val:
                    st.session_state["current_mapping"][ph] = new_val
                    
                    # Session'a kaydet
                    sm.update_session_data(current_session_id, {ph: new_val}, merge=True)
                    
                    # Session ismini kontrol et
                    updated_session = sm.get_session(current_session_id)
                    if updated_session and update_session_name_if_needed(current_session_id, updated_session):
                        st.session_state["current_session_name"] = updated_session['session_name']
                    
                    st.rerun()
                
                st.markdown("---")
    
    # Seçilen şablonların önizlemesi
    if template_items:
        st.markdown("---")
        st.subheader("👁️ Seçilen Şablonların Önizlemesi")

        for template_name, template_data in template_items:
            # Basit dropdown (expander) ile tam içerik önizleme
            with st.expander(f"📄 {template_name}", expanded=False):
                try:
                    # Word belgesinin tam metnini çıkar
                    doc = Document(io.BytesIO(template_data))
                    parts = []

                    # Paragraflar
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            parts.append(paragraph.text.strip())

                    # Tablolar
                    for table in doc.tables:
                        for row in table.rows:
                            cells = []
                            for cell in row.cells:
                                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                                if cell_text:
                                    cells.append(cell_text)
                            if cells:
                                parts.append(" | ".join(cells))

                    # Header/Footer
                    for section in doc.sections:
                        if section.header:
                            for p in section.header.paragraphs:
                                if p.text.strip():
                                    parts.insert(0, f"[BAŞLIK: {p.text.strip()}]")
                        if section.footer:
                            for p in section.footer.paragraphs:
                                if p.text.strip():
                                    parts.append(f"[ALT BİLGİ: {p.text.strip()}]")

                    full_text = "\n\n".join(parts).strip()

                    if not full_text:
                        st.info("Bu şablonda görüntülenebilir metin bulunamadı.")
                    else:
                        # Tüm placeholder'ları doğrudan metin üzerinden regex ile işle
                        try:
                            current_mapping = st.session_state.get("current_mapping", {}) or {}
                            # Önizlemede işbu alanlarını da bugünün değeriyle doldur
                            mapping_with_isbu = {
                                **current_mapping,
                                **today_isbu(datetime.now(IST))
                            }
                            import html as _html
                            pattern = re.compile(r"\{[^}]+\}")

                            def _replace_placeholder(match: re.Match) -> str:
                                ph = match.group(0)
                                # Hem tam eşleşme hem de kıvrıksız anahtar ile eşleşmeyi dene
                                raw_val = str(mapping_with_isbu.get(ph, "")).strip()
                                if not raw_val:
                                    key_nobraces = ph.strip('{}')
                                    # Önce {key} biçimindeki varyantları tara
                                    for k, v in mapping_with_isbu.items():
                                        if isinstance(k, str) and k.strip('{}').lower() == key_nobraces.lower():
                                            raw_val = str(v).strip()
                                            if raw_val:
                                                break
                                if raw_val:
                                    return _html.escape(raw_val)
                                return f"<span style=\"color:#ff4d4f;font-weight:700;\">{_html.escape(ph)}</span>"

                            highlighted_text = pattern.sub(_replace_placeholder, full_text)
                        except Exception:
                            highlighted_text = full_text

                        # Sade, tam genişlikte metin (inline renk ile)
                        st.markdown(
                            f"""
                            <div style="white-space: pre-wrap; word-wrap: break-word; line-height: 1.75; font-size: 16px; font-weight: 500; color: #374151;">{highlighted_text}</div>
                            """,
                            unsafe_allow_html=True
                        )
                except Exception as e:
                    st.error(f"Şablon önizlemesi oluşturulamadı: {e}")
    
    # Belge oluşturma
    if template_items:
        st.subheader("📄 Belge Oluşturma")
        
        if st.button("📄 Tüm Belgeleri Oluştur", type="primary", use_container_width=True):
            if not st.session_state.get("current_mapping"):
                st.warning("Önce bilgileri doldurun.")
            else:
                try:
                    results = []
                    current_mapping = st.session_state["current_mapping"]
                    
                    for idx, (name, data) in enumerate(template_items):
                        doc = Document(io.BytesIO(data))
                        mapping = {k: v for k, v in current_mapping.items() if str(v).strip()}
                        # İşbu alanlarını belge oluşturma anının tarihi/saatine sabitle
                        mapping = {
                            **mapping,
                            **today_isbu(datetime.now(IST))
                        }
                        replaced = replace_placeholders_in_document(doc, mapping)
                        
                        buf = io.BytesIO()
                        doc.save(buf)
                        out_bytes = buf.getvalue()
                        
                        # Dosya adı
                        safe_session_name = re.sub(r'[^\w\s-]', '', current_session_name).strip()[:20]
                        out_name = f"{safe_session_name}_{os.path.splitext(name)[0]}.docx"
                        
                        results.append({
                            "name": out_name,
                            "replaced": replaced,
                            "data": out_bytes,
                            "key": f"dl_{idx}_{out_name}",
                        })
                    
                    st.session_state["results"] = results
                    st.success("✅ Belgeler hazırlandı!")
                except Exception as e:
                    st.error(f"Belge oluşturma hatası: {e}")
        
        # İndirme butonları
        if st.session_state.get("results"):
            st.markdown("---")
            st.subheader("📥 İndirilecek Belgeler")
            
            for r in st.session_state["results"]:
                col_info, col_download = st.columns([3, 1])
                
                with col_info:
                    st.write(f"**{r['name']}** → {r['replaced']} değişiklik")
                
                with col_download:
                    st.download_button(
                        label="📥 İndir",
                        data=r["data"],
                        file_name=r["name"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=r["key"],
                        use_container_width=True
                    )

if __name__ == "__main__":
    main()
